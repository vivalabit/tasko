import base64
import hashlib
import json
import re
from datetime import datetime, timedelta
from io import BytesIO
from uuid import uuid4

from docx import Document
from sqlalchemy import delete, func, select
from sqlalchemy.exc import IntegrityError, SQLAlchemyError
from sqlalchemy.orm import Session, selectinload

from fastapi import APIRouter, Depends, HTTPException, Query, Response, status

from app.core.database import get_db
from app.core.settings import Settings, get_settings
from app.models.applications import StoredApplicationRecord
from app.models.assistant import (
    AssistantApplicationContext,
    AssistantCandidateConfirmation,
    AssistantJobContext,
    AssistantSourceDocument,
)
from app.models.documents import (
    DocumentAttachRequest,
    DocumentAttachmentRecord,
    DocumentCreateRequest,
    DocumentFileRecord,
    DocumentGenerationProvenanceRecord,
    DocumentPackItemRequest,
    DocumentPackJobRecord,
    DocumentPackPayload,
    DocumentPackRequest,
    DocumentPackStagePayload,
    DocumentPackValidationRequest,
    DocumentPackValidationPayload,
    DocumentPayload,
    DocumentRecord,
    DocumentRestoreRequest,
    DocumentTemplateCreateRequest,
    DocumentTemplatePayload,
    DocumentTemplatePreflightPayload,
    DocumentTemplatePreflightRequest,
    DocumentTemplateRecord,
    DocumentUpdateRequest,
    DocumentValidationArtifactRecord,
    DocumentVersionGenerationProvenanceRecord,
    DocumentVersionPagePayload,
    DocumentVersionPayload,
    DocumentVersionRecord,
    DocumentVersionValidationRecord,
    WorkspaceSourceDocumentCreateRequest,
    WorkspaceSourceDocumentPayload,
    WorkspaceSourceDocumentRecord,
    utc_now,
)
from app.models.profile import ProfilePayload
from app.services.assistant import (
    analyze_openclaw_assistant_context,
    build_source_document_context,
)
from app.services.document_export import build_document_docx, build_document_from_template
from app.services.generation_context import (
    AuthoritativeGenerationContext,
    GenerationContextError,
    load_authoritative_generation_context,
)
from app.services.document_security import DocumentSecurityError, validate_docx_package
from app.services.document_validation import (
    DocumentValidationError,
    validate_generated_document,
)
from app.services.document_preflight import analyze_document_template

router = APIRouter()

DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MAX_TEMPLATE_BYTES = 10_000_000
VALIDATION_ARTIFACT_TTL = timedelta(minutes=30)
PACK_JOB_TTL = timedelta(days=7)
DOCUMENT_VERSION_PAGE_SIZE = 20


@router.get("", response_model=list[DocumentPayload])
def list_documents(
    job_id: str | None = Query(default=None, alias="jobId"),
    application_id: str | None = Query(default=None, alias="applicationId"),
    db: Session = Depends(get_db),
) -> list[DocumentPayload]:
    try:
        statement = (
            select(DocumentRecord)
            .options(
                selectinload(DocumentRecord.versions),
                selectinload(DocumentRecord.attachments),
                selectinload(DocumentRecord.generation_provenance),
                selectinload(DocumentRecord.version_validations),
                selectinload(DocumentRecord.files),
            )
            .order_by(DocumentRecord.updated_at.desc())
        )
        if job_id is not None:
            statement = statement.where(DocumentRecord.job_id == job_id)
        if application_id is not None:
            statement = statement.join(DocumentAttachmentRecord).where(
                DocumentAttachmentRecord.application_id == application_id
            )
        records = db.scalars(statement).unique().all()
        return [
            document_payload(
                record,
                version_limit=DOCUMENT_VERSION_PAGE_SIZE,
                current_generation_fingerprint=authoritative_current_generation_fingerprint(
                    db,
                    record,
                    application_id=application_id,
                ),
            )
            for record in records
        ]
    except SQLAlchemyError as exc:
        raise database_unavailable(exc) from exc


@router.post("", response_model=DocumentPayload, status_code=status.HTTP_201_CREATED)
def create_document(
    request: DocumentCreateRequest,
    db: Session = Depends(get_db),
) -> DocumentPayload:
    try:
        has_provenance = bool(request.generation_model and request.generation_model.strip())
        generation_context = None
        generation_provenance = None
        if has_provenance:
            if not request.application_id:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                    detail="Application ID is required for generated documents",
                )
            if not request.template_id:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                    detail="Document template is required for generated documents",
                )
            generation_context = load_generation_context_or_http(
                db,
                application_id=request.application_id,
                template_id=request.template_id,
                document_type=request.type,
                expected_job_id=request.job_id,
            )
            generation_provenance = generation_context.provenance()
        now = utc_now()
        document_id = str(uuid4())
        record = DocumentRecord(
            id=document_id,
            type=request.type,
            title=request.title.strip(),
            job_id=generation_context.job_id if generation_context else request.job_id,
            current_version=1,
            created_at=now,
            updated_at=now,
        )
        record.versions.append(
            DocumentVersionRecord(
                id=str(uuid4()),
                document_id=document_id,
                version=1,
                content=request.content,
                created_at=now,
            )
        )
        if has_provenance:
            assert generation_provenance is not None
            set_current_generation_provenance(
                record,
                generation_provenance.generation_fingerprint,
                request.generation_model or "",
                generation_provenance.input_versions,
                now,
            )
            append_version_generation_provenance(
                record,
                1,
                generation_provenance.generation_fingerprint,
                request.generation_model or "",
                generation_provenance.input_versions,
                now,
            )
        db.add(record)
        db.flush()
        if request.template_id:
            template = (
                generation_context.template
                if generation_context
                else require_template(db, request.template_id)
            )
            if template.type != request.type:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail="Document template type does not match document type",
                )
            try:
                rendered_content = build_document_from_template(
                    template_content=template.content,
                    content=request.content,
                    document_type=request.type,
                )
            except ValueError as exc:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail=str(exc),
                ) from exc
            if has_provenance:
                assert generation_context is not None
                validation = validate_document_or_422(
                    template_content=template.content,
                    rendered_content=rendered_content,
                    generated_content=request.content,
                    document_type=request.type,
                    evidence=generation_context.validation_evidence(),
                )
                append_document_validation(record, 1, validation, now)
            db.add(
                DocumentFileRecord(
                    id=str(uuid4()),
                    document_id=document_id,
                    version=1,
                    template_id=template.id,
                    content=rendered_content,
                    created_at=now,
                )
            )
        if request.application_id:
            attach_document_record(db, record, request.application_id)
        db.commit()
        return document_payload(
            require_document(db, document_id),
            current_generation_fingerprint=(
                generation_provenance.generation_fingerprint
                if generation_provenance
                else None
            ),
        )
    except HTTPException:
        db.rollback()
        raise
    except SQLAlchemyError as exc:
        db.rollback()
        raise database_unavailable(exc) from exc


@router.post(
    "/packs/validate-resume",
    response_model=DocumentPackValidationPayload,
)
def validate_resume_pack_item(
    request: DocumentPackValidationRequest,
    db: Session = Depends(get_db),
) -> DocumentPackValidationPayload:
    try:
        cleanup_expired_pack_storage(db)
        require_pack_document_ownership(
            db,
            request.resume,
            document_type="tailored_resume",
            application_id=request.application_id,
        )
        context = load_authoritative_generation_context(
            db,
            application_id=request.application_id,
            template_id=request.resume.template_id,
            document_type="tailored_resume",
        )
        rendered_content, validation = prepare_pack_document(
            request.resume,
            "tailored_resume",
            context=context,
        )
        artifact = create_validation_artifact(
            application_id=request.application_id,
            item=request.resume,
            context=context,
            rendered_content=rendered_content,
            validation=validation,
        )
        db.add(artifact)
        db.commit()
        return DocumentPackValidationPayload(
            validation=validation,
            validation_artifact_id=artifact.id,
            expires_at=artifact.expires_at,
        )
    except GenerationContextError as exc:
        db.rollback()
        raise pack_validation_failed("resume_validation", exc) from exc
    except DocumentValidationError as exc:
        db.rollback()
        raise pack_validation_failed("resume_validation", exc) from exc
    except ValueError as exc:
        db.rollback()
        raise pack_validation_failed("resume_validation", exc) from exc
    except SQLAlchemyError as exc:
        db.rollback()
        raise database_unavailable(exc) from exc


@router.post(
    "/packs",
    response_model=DocumentPackPayload,
    status_code=status.HTTP_201_CREATED,
)
def create_document_pack(
    request: DocumentPackRequest,
    db: Session = Depends(get_db),
) -> DocumentPackPayload:
    try:
        cleanup_expired_pack_storage(db)
        if request.persistence_mode == "atomic" and request.cover_letter is None:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail={
                    "stage": "cover_letter_generation",
                    "status": "rolled_back",
                    "message": "Atomic application packs require a cover letter",
                },
            )

        require_pack_document_ownership(
            db,
            request.resume,
            document_type="tailored_resume",
            application_id=request.application_id,
        )
        if request.cover_letter is not None:
            require_pack_document_ownership(
                db,
                request.cover_letter,
                document_type="cover_letter",
                application_id=request.application_id,
            )

        try:
            resume_context = load_authoritative_generation_context(
                db,
                application_id=request.application_id,
                template_id=request.resume.template_id,
                document_type="tailored_resume",
                expected_job_id=request.job_id,
            )
        except (DocumentValidationError, GenerationContextError, ValueError) as exc:
            raise pack_validation_failed("resume_validation", exc) from exc

        cover_prepared: (
            tuple[
                AuthoritativeGenerationContext,
                bytes,
                dict[str, object],
            ]
            | None
        ) = None
        cover_failure = ""
        cover_context = None
        if request.cover_letter is not None:
            try:
                cover_context = load_authoritative_generation_context(
                    db,
                    application_id=request.application_id,
                    template_id=request.cover_letter.template_id,
                    document_type="cover_letter",
                    expected_job_id=request.job_id,
                )
            except (DocumentValidationError, GenerationContextError, ValueError) as exc:
                cover_failure = str(exc)
                if request.persistence_mode == "atomic":
                    raise pack_validation_failed("cover_letter_validation", exc) from exc
        else:
            cover_failure = request.partial_reason or "Cover letter generation did not complete"

        request_fingerprint = document_pack_request_fingerprint(
            request,
            resume_context=resume_context,
            cover_context=cover_context,
        )
        existing_job = db.get(DocumentPackJobRecord, request.pack_job_id)
        if existing_job:
            if existing_job.request_fingerprint != request_fingerprint:
                raise HTTPException(
                    status_code=status.HTTP_409_CONFLICT,
                    detail="Pack job ID was already used for a different request",
                )
            return document_pack_payload(existing_job, db)

        try:
            if request.resume.validation_artifact_id:
                resume_content, resume_validation = consume_validation_artifact(
                    db,
                    artifact_id=request.resume.validation_artifact_id,
                    item=request.resume,
                    context=resume_context,
                )
            else:
                resume_content, resume_validation = prepare_pack_document(
                    request.resume,
                    "tailored_resume",
                    context=resume_context,
                )
        except (DocumentValidationError, GenerationContextError, ValueError) as exc:
            raise pack_validation_failed("resume_validation", exc) from exc

        if request.cover_letter is not None and cover_context is not None:
            try:
                cover_content, cover_validation = prepare_pack_document(
                    request.cover_letter,
                    "cover_letter",
                    context=cover_context,
                )
                cover_prepared = cover_context, cover_content, cover_validation
            except (DocumentValidationError, GenerationContextError, ValueError) as exc:
                cover_failure = str(exc)
                if request.persistence_mode == "atomic":
                    raise pack_validation_failed("cover_letter_validation", exc) from exc

        now = utc_now()
        document_ids = [
            persist_pack_document(
                db=db,
                item=request.resume,
                document_type="tailored_resume",
                context=resume_context,
                rendered_content=resume_content,
                validation=resume_validation,
                created_at=now,
            )
        ]
        stages = [
            {
                "id": "resume_validation",
                "status": "completed",
                "message": "CV passed factual and visual validation",
            }
        ]
        pack_status = "partial"
        message = cover_failure
        if request.cover_letter is not None and cover_prepared is not None:
            cover_context, cover_content, cover_validation = cover_prepared
            document_ids.append(
                persist_pack_document(
                    db=db,
                    item=request.cover_letter,
                    document_type="cover_letter",
                    context=cover_context,
                    rendered_content=cover_content,
                    validation=cover_validation,
                    created_at=now,
                )
            )
            pack_status = "completed"
            message = "Application pack saved atomically"
            stages.append(
                {
                    "id": "cover_letter_validation",
                    "status": "completed",
                    "message": "Cover letter passed factual and visual validation",
                }
            )
        else:
            stages.append(
                {
                    "id": "cover_letter_validation",
                    "status": "failed" if request.cover_letter else "skipped",
                    "message": cover_failure,
                }
            )
        stages.append(
            {
                "id": "saving",
                "status": "completed",
                "message": (
                    "Both documents committed in one transaction"
                    if pack_status == "completed"
                    else "Validated CV saved as an explicit partial pack"
                ),
            }
        )
        job = DocumentPackJobRecord(
            id=request.pack_job_id,
            request_fingerprint=request_fingerprint,
            application_id=request.application_id,
            persistence_mode=request.persistence_mode,
            status=pack_status,
            document_ids=document_ids,
            stages=stages,
            message=message,
            created_at=now,
            updated_at=now,
            expires_at=now + PACK_JOB_TTL,
        )
        db.add(job)
        db.commit()
        return document_pack_payload(job, db)
    except HTTPException:
        db.rollback()
        raise
    except SQLAlchemyError as exc:
        db.rollback()
        raise database_unavailable(exc) from exc


@router.get(
    "/packs/{pack_job_id}",
    response_model=DocumentPackPayload,
)
def get_document_pack_status(
    pack_job_id: str,
    application_id: str = Query(min_length=1, max_length=160, alias="applicationId"),
    db: Session = Depends(get_db),
) -> DocumentPackPayload:
    try:
        cleanup_expired_pack_storage(db)
        job = db.get(DocumentPackJobRecord, pack_job_id)
        if not job:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Pack job not found",
            )
        if job.application_id != application_id:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail="Pack job is not attached to the application",
            )
        return document_pack_payload(job, db)
    except HTTPException:
        raise
    except SQLAlchemyError as exc:
        raise database_unavailable(exc) from exc


@router.get("/{document_id}", response_model=DocumentPayload)
def get_document(document_id: str, db: Session = Depends(get_db)) -> DocumentPayload:
    try:
        record = require_document(db, document_id)
        return document_payload(
            record,
            version_limit=DOCUMENT_VERSION_PAGE_SIZE,
            current_generation_fingerprint=authoritative_current_generation_fingerprint(db, record),
        )
    except HTTPException:
        raise
    except SQLAlchemyError as exc:
        raise database_unavailable(exc) from exc


@router.get("/{document_id}/versions", response_model=DocumentVersionPagePayload)
def list_document_versions(
    document_id: str,
    limit: int = Query(default=DOCUMENT_VERSION_PAGE_SIZE, ge=1, le=100),
    offset: int = Query(default=0, ge=0),
    db: Session = Depends(get_db),
) -> DocumentVersionPagePayload:
    try:
        if db.get(DocumentRecord, document_id) is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found",
            )
        total = db.scalar(
            select(func.count()).select_from(DocumentVersionRecord).where(
                DocumentVersionRecord.document_id == document_id
            )
        ) or 0
        versions = list(
            reversed(
                db.scalars(
                    select(DocumentVersionRecord)
                    .where(DocumentVersionRecord.document_id == document_id)
                    .order_by(DocumentVersionRecord.version.desc())
                    .offset(offset)
                    .limit(limit)
                ).all()
            )
        )
        version_numbers = [version.version for version in versions]
        rendered_versions = (
            set(
                db.scalars(
                    select(DocumentFileRecord.version).where(
                        DocumentFileRecord.document_id == document_id,
                        DocumentFileRecord.version.in_(version_numbers),
                    )
                ).all()
            )
            if version_numbers
            else set()
        )
        validations = (
            db.scalars(
                select(DocumentVersionValidationRecord).where(
                    DocumentVersionValidationRecord.document_id == document_id,
                    DocumentVersionValidationRecord.version.in_(version_numbers),
                )
            ).all()
            if version_numbers
            else []
        )
        return DocumentVersionPagePayload(
            items=document_version_payloads(
                versions,
                rendered_versions=rendered_versions,
                validations_by_version={
                    validation.version: validation for validation in validations
                },
            ),
            total=total,
            limit=limit,
            offset=offset,
        )
    except HTTPException:
        raise
    except SQLAlchemyError as exc:
        raise database_unavailable(exc) from exc


@router.patch("/{document_id}", response_model=DocumentPayload)
def update_document(
    document_id: str,
    request: DocumentUpdateRequest,
    db: Session = Depends(get_db),
) -> DocumentPayload:
    try:
        record = require_document(db, document_id)
        fields = request.model_fields_set
        has_provenance = bool(request.generation_model and request.generation_model.strip())
        if request.application_id:
            require_document_application_ownership(record, request.application_id)
        if (has_provenance or "template_id" in fields) and (
            "content" not in fields or request.content is None
        ):
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                detail="Generation provenance and template require document content",
            )
        if "title" in fields and request.title is not None:
            record.title = request.title.strip()
        if "job_id" in fields:
            record.job_id = request.job_id
        if "content" in fields and request.content is not None:
            current = current_version_record(record)
            if current.content != request.content or has_provenance or "template_id" in fields:
                now = utc_now()
                current_provenance = record.generation_provenance
                if current_provenance and not any(
                    provenance.version == record.current_version
                    for provenance in record.version_generation_provenance
                ):
                    append_version_generation_provenance(
                        record,
                        record.current_version,
                        current_provenance.generation_fingerprint,
                        current_provenance.generation_model,
                        current_provenance.input_versions,
                        now,
                    )
                next_version = record.current_version + 1
                record.versions.append(
                    DocumentVersionRecord(
                        id=str(uuid4()),
                        document_id=record.id,
                        version=next_version,
                        content=request.content,
                        created_at=now,
                    )
                )
                current_file = document_file_record(db, record.id, record.current_version)
                if "template_id" in fields:
                    template_id = request.template_id
                else:
                    template_id = current_file.template_id if current_file else None
                generation_context = None
                generation_provenance = None
                if has_provenance and not template_id:
                    raise HTTPException(
                        status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                        detail="Document template is required for generated documents",
                    )
                if template_id:
                    if has_provenance:
                        application_id = request.application_id or single_attachment_application_id(
                            record
                        )
                        if not application_id:
                            raise HTTPException(
                                status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                                detail="Application ID is required for generated documents",
                            )
                        generation_context = load_generation_context_or_http(
                            db,
                            application_id=application_id,
                            template_id=template_id,
                            document_type=record.type,
                            expected_job_id=(
                                request.job_id if "job_id" in fields else record.job_id
                            ),
                        )
                        generation_provenance = generation_context.provenance()
                    template = (
                        generation_context.template
                        if generation_context
                        else require_template(db, template_id)
                    )
                    if template.type != record.type:
                        raise HTTPException(
                            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
                            detail="Document template type does not match document type",
                        )
                    try:
                        rendered_content = build_document_from_template(
                            template_content=template.content,
                            content=request.content,
                            document_type=record.type,
                        )
                    except ValueError as exc:
                        raise HTTPException(
                            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                            detail=str(exc),
                        ) from exc
                    if has_provenance:
                        assert generation_context is not None
                        validation = validate_document_or_422(
                            template_content=template.content,
                            rendered_content=rendered_content,
                            generated_content=request.content,
                            document_type=record.type,
                            evidence=generation_context.validation_evidence(),
                        )
                        append_document_validation(
                            record,
                            next_version,
                            validation,
                            now,
                        )
                    db.add(
                        DocumentFileRecord(
                            id=str(uuid4()),
                            document_id=record.id,
                            version=next_version,
                            template_id=template.id,
                            content=rendered_content,
                            created_at=now,
                        )
                    )
                    if generation_context:
                        record.job_id = generation_context.job_id
                        attach_document_record(db, record, generation_context.application_id)
                if has_provenance:
                    assert generation_provenance is not None
                    set_current_generation_provenance(
                        record,
                        generation_provenance.generation_fingerprint,
                        request.generation_model or "",
                        generation_provenance.input_versions,
                        now,
                    )
                else:
                    record.generation_provenance = None
                current_provenance = record.generation_provenance
                if current_provenance:
                    append_version_generation_provenance(
                        record,
                        next_version,
                        current_provenance.generation_fingerprint,
                        current_provenance.generation_model,
                        current_provenance.input_versions,
                        now,
                    )
                record.current_version = next_version
        record.updated_at = utc_now()
        db.commit()
        updated = require_document(db, document_id)
        return document_payload(
            updated,
            current_generation_fingerprint=authoritative_current_generation_fingerprint(
                db,
                updated,
                application_id=request.application_id,
            ),
        )
    except HTTPException:
        db.rollback()
        raise
    except SQLAlchemyError as exc:
        db.rollback()
        raise database_unavailable(exc) from exc


@router.post("/{document_id}/restore", response_model=DocumentPayload)
def restore_document_version(
    document_id: str,
    request: DocumentRestoreRequest,
    db: Session = Depends(get_db),
) -> DocumentPayload:
    try:
        record = require_document(db, document_id)
        source = next(
            (version for version in record.versions if version.version == request.version),
            None,
        )
        if not source:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document version not found",
            )
        now = utc_now()
        next_version = record.current_version + 1
        record.versions.append(
            DocumentVersionRecord(
                id=str(uuid4()),
                document_id=record.id,
                version=next_version,
                content=source.content,
                created_at=now,
            )
        )
        source_file = document_file_record(db, record.id, source.version)
        if source_file:
            db.add(
                DocumentFileRecord(
                    id=str(uuid4()),
                    document_id=record.id,
                    version=next_version,
                    template_id=source_file.template_id,
                    content=source_file.content,
                    created_at=now,
                )
            )
        source_provenance = next(
            (
                provenance
                for provenance in record.version_generation_provenance
                if provenance.version == source.version
            ),
            None,
        )
        if source_provenance:
            set_current_generation_provenance(
                record,
                source_provenance.generation_fingerprint,
                source_provenance.generation_model,
                source_provenance.input_versions,
                now,
            )
            append_version_generation_provenance(
                record,
                next_version,
                source_provenance.generation_fingerprint,
                source_provenance.generation_model,
                source_provenance.input_versions,
                now,
            )
        else:
            record.generation_provenance = None
        source_validation = next(
            (
                validation
                for validation in record.version_validations
                if validation.version == source.version
            ),
            None,
        )
        if source_validation:
            append_document_validation(
                record,
                next_version,
                {
                    "factual": source_validation.factual_report,
                    "visual": source_validation.visual_report,
                    "diff": source_validation.diff_items,
                },
                now,
            )
        record.current_version = next_version
        record.updated_at = now
        db.commit()
        restored = require_document(db, document_id)
        return document_payload(
            restored,
            current_generation_fingerprint=authoritative_current_generation_fingerprint(
                db,
                restored,
            ),
        )
    except HTTPException:
        db.rollback()
        raise
    except SQLAlchemyError as exc:
        db.rollback()
        raise database_unavailable(exc) from exc


@router.post("/{document_id}/attachments", response_model=DocumentPayload)
def attach_document(
    document_id: str,
    request: DocumentAttachRequest,
    db: Session = Depends(get_db),
) -> DocumentPayload:
    try:
        record = require_document(db, document_id)
        attach_document_record(db, record, request.application_id)
        record.updated_at = utc_now()
        db.commit()
        attached = require_document(db, document_id)
        return document_payload(
            attached,
            current_generation_fingerprint=authoritative_current_generation_fingerprint(
                db,
                attached,
                application_id=request.application_id,
            ),
        )
    except HTTPException:
        db.rollback()
        raise
    except SQLAlchemyError as exc:
        db.rollback()
        raise database_unavailable(exc) from exc


@router.delete(
    "/{document_id}/attachments/{application_id}",
    status_code=status.HTTP_204_NO_CONTENT,
)
def detach_document(
    document_id: str,
    application_id: str,
    db: Session = Depends(get_db),
) -> None:
    try:
        record = require_document(db, document_id)
        attachment = next(
            (
                item
                for item in record.attachments
                if item.application_id == application_id
            ),
            None,
        )
        if not attachment:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document attachment not found",
            )
        db.delete(attachment)
        record.updated_at = utc_now()
        db.commit()
    except HTTPException:
        db.rollback()
        raise
    except SQLAlchemyError as exc:
        db.rollback()
        raise database_unavailable(exc) from exc


@router.get("/{document_id}/download")
def download_document(
    document_id: str,
    version: int | None = Query(default=None, ge=1),
    db: Session = Depends(get_db),
) -> Response:
    try:
        record = require_document(db, document_id)
        selected_version = (
            next((item for item in record.versions if item.version == version), None)
            if version is not None
            else current_version_record(record)
        )
        if not selected_version:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document version not found",
            )
        rendered_file = document_file_record(db, record.id, selected_version.version)
        content = rendered_file.content if rendered_file else build_document_docx(
            title=record.title,
            content=selected_version.content,
            document_type=record.type,
            version=selected_version.version,
        )
        filename = safe_filename(record.title)
        return Response(
            content=content,
            media_type=(
                DOCX_CONTENT_TYPE
            ),
            headers={
                "Content-Disposition": f'attachment; filename="{filename}-v{selected_version.version}.docx"'
            },
        )
    except HTTPException:
        raise
    except SQLAlchemyError as exc:
        raise database_unavailable(exc) from exc


@router.get("/templates/library", response_model=list[DocumentTemplatePayload])
def list_document_templates(db: Session = Depends(get_db)) -> list[DocumentTemplatePayload]:
    try:
        records = db.scalars(
            select(DocumentTemplateRecord).order_by(DocumentTemplateRecord.updated_at.desc())
        ).all()
        return [document_template_payload(record) for record in records]
    except SQLAlchemyError as exc:
        raise database_unavailable(exc) from exc


@router.post(
    "/templates/preflight",
    response_model=DocumentTemplatePreflightPayload,
)
def preflight_document_template(
    request: DocumentTemplatePreflightRequest,
    db: Session = Depends(get_db),
    settings: Settings = Depends(get_settings),
) -> DocumentTemplatePreflightPayload:
    if not request.file_name.lower().endswith(".docx"):
        return DocumentTemplatePreflightPayload(
            supported=False,
            editable_count=0,
            immutable_count=0,
            immutable_elements=[],
            rejected_elements=[
                {
                    "element": "fileType",
                    "description": "template must be a .docx file",
                }
            ],
        )
    _, content = decode_data_url(request.data_url)
    if not content or len(content) > MAX_TEMPLATE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_CONTENT_TOO_LARGE,
            detail="Template must be a non-empty DOCX file under 10 MB",
        )

    report = analyze_document_template(content, request.type)
    if not report["supported"]:
        return DocumentTemplatePreflightPayload(
            supported=False,
            editable_count=report["editableCount"],
            immutable_count=report["immutableCount"],
            immutable_elements=report["immutableElements"],
            rejected_elements=report["rejectedElements"],
        )

    ai_context = None
    warnings: list[str] = []
    if request.application_id:
        try:
            now = utc_now()
            transient_template = DocumentTemplateRecord(
                id=str(uuid4()),
                type=request.type,
                name=request.name.strip(),
                file_name=safe_upload_filename(request.file_name),
                content_type=DOCX_CONTENT_TYPE,
                content_sha256=hashlib.sha256(content).hexdigest(),
                content=content,
                extracted_text="",
                created_at=now,
                updated_at=now,
            )
            generation_context = load_authoritative_generation_context(
                db,
                application_id=request.application_id,
                template_id=transient_template.id,
                document_type=request.type,
                template_override=transient_template,
            )
            profile, job, application, source, confirmations = (
                assistant_preflight_inputs(generation_context)
            )
            ai_context = analyze_openclaw_assistant_context(
                message_characters=request.prompt_characters,
                context_kind="application",
                profile=profile,
                job=job,
                application=application,
                source_documents=[source],
                candidate_confirmations=confirmations,
                max_prompt_chars=settings.openclaw_assistant_max_prompt_chars,
            )
            source_context = dict(report["sourceContext"])
            actual_source_context = build_source_document_context([source])[0]
            source_elements = actual_source_context.get(
                "blocks" if request.type == "tailored_resume" else "paragraphs",
                [],
            )
            actual_source_characters = len(
                json.dumps(
                    actual_source_context,
                    ensure_ascii=False,
                    separators=(",", ":"),
                )
            )
            metadata_characters = max(
                0,
                actual_source_characters - source_context["includedCharacters"],
            )
            source_context["includedElements"] = len(source_elements)
            source_context["omittedElements"] = max(
                0,
                source_context["totalElements"] - len(source_elements),
            )
            source_context["includedCharacters"] = actual_source_characters
            source_context["estimatedCharacters"] += metadata_characters
            source_context["truncated"] = source_context["omittedElements"] > 0
            omitted_source_characters = max(
                0,
                source_context["estimatedCharacters"]
                - source_context["includedCharacters"],
            )
            ai_context["estimatedCharacters"] += omitted_source_characters
            ai_context["source"] = source_context
            ai_context["truncated"] = bool(
                ai_context["truncated"] or source_context["truncated"]
            )
        except GenerationContextError as exc:
            warnings.append(f"AI context could not be estimated: {exc}")

    return DocumentTemplatePreflightPayload(
        supported=True,
        template=None,
        editable_count=report["editableCount"],
        immutable_count=report["immutableCount"],
        immutable_elements=report["immutableElements"],
        rejected_elements=report["rejectedElements"],
        ai_context=ai_context,
        warnings=warnings,
    )


@router.get(
    "/workspace-sources/library",
    response_model=list[WorkspaceSourceDocumentPayload],
)
def list_workspace_source_documents(
    application_id: str = Query(min_length=1, max_length=160, alias="applicationId"),
    db: Session = Depends(get_db),
) -> list[WorkspaceSourceDocumentPayload]:
    try:
        require_stored_application(db, application_id)
        records = db.scalars(
            select(WorkspaceSourceDocumentRecord)
            .where(WorkspaceSourceDocumentRecord.application_id == application_id)
            .order_by(WorkspaceSourceDocumentRecord.updated_at.desc())
        ).all()
        return [workspace_source_document_payload(record) for record in records]
    except HTTPException:
        raise
    except SQLAlchemyError as exc:
        raise database_unavailable(exc) from exc


@router.post(
    "/workspace-sources",
    response_model=WorkspaceSourceDocumentPayload,
    status_code=status.HTTP_201_CREATED,
)
def create_workspace_source_document(
    request: WorkspaceSourceDocumentCreateRequest,
    db: Session = Depends(get_db),
) -> WorkspaceSourceDocumentPayload:
    if not request.file_name.lower().endswith(".docx"):
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Workspace source must be a .docx file",
        )
    _, content = decode_data_url(request.data_url)
    if not content or len(content) > MAX_TEMPLATE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_CONTENT_TOO_LARGE,
            detail="Workspace source must be a non-empty DOCX file under 10 MB",
        )
    try:
        validate_docx_package(content)
    except DocumentSecurityError as exc:
        raise HTTPException(
            status_code=(
                status.HTTP_413_CONTENT_TOO_LARGE
                if exc.limit_exceeded
                else status.HTTP_422_UNPROCESSABLE_ENTITY
            ),
            detail=str(exc),
        ) from exc

    now = utc_now()
    record = WorkspaceSourceDocumentRecord(
        id=str(uuid4()),
        application_id=request.application_id,
        category=request.category,
        title=request.title.strip(),
        language=request.language.strip(),
        file_name=safe_upload_filename(request.file_name),
        content_type=DOCX_CONTENT_TYPE,
        content=content,
        created_at=now,
        updated_at=now,
    )
    try:
        require_stored_application(db, request.application_id)
        db.add(record)
        db.commit()
        db.refresh(record)
        return workspace_source_document_payload(record)
    except HTTPException:
        db.rollback()
        raise
    except SQLAlchemyError as exc:
        db.rollback()
        raise database_unavailable(exc) from exc


@router.delete(
    "/workspace-sources/{source_id}",
    status_code=status.HTTP_204_NO_CONTENT,
)
def delete_workspace_source_document(
    source_id: str,
    application_id: str = Query(min_length=1, max_length=160, alias="applicationId"),
    db: Session = Depends(get_db),
) -> None:
    try:
        record = db.scalar(
            select(WorkspaceSourceDocumentRecord).where(
                WorkspaceSourceDocumentRecord.id == source_id,
                WorkspaceSourceDocumentRecord.application_id == application_id,
            )
        )
        if record is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Workspace source not found",
            )
        db.delete(record)
        db.commit()
    except HTTPException:
        db.rollback()
        raise
    except SQLAlchemyError as exc:
        db.rollback()
        raise database_unavailable(exc) from exc


@router.post(
    "/templates",
    response_model=DocumentTemplatePayload,
    status_code=status.HTTP_201_CREATED,
)
def create_document_template(
    request: DocumentTemplateCreateRequest,
    db: Session = Depends(get_db),
) -> DocumentTemplatePayload:
    if not request.file_name.lower().endswith(".docx"):
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Template must be a .docx file",
        )
    content_type, content = decode_data_url(request.data_url)
    if not content or len(content) > MAX_TEMPLATE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_CONTENT_TOO_LARGE,
            detail="Template must be a non-empty DOCX file under 10 MB",
        )
    try:
        validate_docx_package(content)
    except DocumentSecurityError as exc:
        raise HTTPException(
            status_code=(
                status.HTTP_413_CONTENT_TOO_LARGE
                if exc.limit_exceeded
                else status.HTTP_422_UNPROCESSABLE_ENTITY
            ),
            detail=str(exc),
        ) from exc

    content_sha256 = hashlib.sha256(content).hexdigest()
    try:
        duplicate = db.scalar(
            select(DocumentTemplateRecord).where(
                DocumentTemplateRecord.type == request.type,
                DocumentTemplateRecord.content_sha256 == content_sha256,
            )
        )
        if duplicate is not None:
            return document_template_payload(duplicate)
    except SQLAlchemyError as exc:
        raise database_unavailable(exc) from exc

    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Template is not a valid DOCX file",
        ) from exc

    extracted_text = "\n".join(
        paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
    )
    now = utc_now()
    record = DocumentTemplateRecord(
        id=str(uuid4()),
        type=request.type,
        name=request.name.strip(),
        file_name=safe_upload_filename(request.file_name),
        content_type=content_type or DOCX_CONTENT_TYPE,
        content_sha256=content_sha256,
        content=content,
        extracted_text=extracted_text,
        created_at=now,
        updated_at=now,
    )
    try:
        db.add(record)
        db.commit()
        db.refresh(record)
        return document_template_payload(record)
    except IntegrityError as exc:
        db.rollback()
        duplicate = db.scalar(
            select(DocumentTemplateRecord).where(
                DocumentTemplateRecord.type == request.type,
                DocumentTemplateRecord.content_sha256 == content_sha256,
            )
        )
        if duplicate is not None:
            return document_template_payload(duplicate)
        raise database_unavailable(exc) from exc
    except SQLAlchemyError as exc:
        db.rollback()
        raise database_unavailable(exc) from exc


@router.delete("/templates/{template_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_document_template(template_id: str, db: Session = Depends(get_db)) -> None:
    try:
        template = require_template(db, template_id)
        db.delete(template)
        db.commit()
    except HTTPException:
        db.rollback()
        raise
    except SQLAlchemyError as exc:
        db.rollback()
        raise database_unavailable(exc) from exc


@router.delete("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_document(document_id: str, db: Session = Depends(get_db)) -> None:
    try:
        record = require_document(db, document_id)
        db.delete(record)
        db.commit()
    except HTTPException:
        db.rollback()
        raise
    except SQLAlchemyError as exc:
        db.rollback()
        raise database_unavailable(exc) from exc


def set_current_generation_provenance(
    record: DocumentRecord,
    generation_fingerprint: str,
    generation_model: str,
    input_versions: dict[str, object],
    created_at: datetime,
) -> None:
    provenance = record.generation_provenance
    if provenance is None:
        provenance = DocumentGenerationProvenanceRecord(document_id=record.id)
        record.generation_provenance = provenance
    provenance.generation_fingerprint = generation_fingerprint
    provenance.generation_model = generation_model.strip()
    provenance.input_versions = input_versions
    provenance.created_at = created_at


def append_version_generation_provenance(
    record: DocumentRecord,
    version: int,
    generation_fingerprint: str,
    generation_model: str,
    input_versions: dict[str, object],
    created_at: datetime,
) -> None:
    record.version_generation_provenance.append(
        DocumentVersionGenerationProvenanceRecord(
            document_id=record.id,
            version=version,
            generation_fingerprint=generation_fingerprint,
            generation_model=generation_model.strip(),
            input_versions=input_versions,
            created_at=created_at,
        )
    )


def validate_document_or_422(
    *,
    template_content: bytes,
    rendered_content: bytes,
    generated_content: str,
    document_type: str,
    evidence: dict[str, object],
) -> dict[str, object]:
    try:
        return validate_generated_document(
            template_content=template_content,
            rendered_content=rendered_content,
            generated_content=generated_content,
            document_type=document_type,
            evidence=evidence,
        )
    except DocumentValidationError as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail=str(exc),
        ) from exc


def append_document_validation(
    record: DocumentRecord,
    version: int,
    validation: dict[str, object],
    created_at: datetime,
) -> None:
    record.version_validations.append(
        DocumentVersionValidationRecord(
            document_id=record.id,
            version=version,
            factual_report=validation.get("factual", {}),
            visual_report=validation.get("visual", {}),
            diff_items=validation.get("diff", []),
            created_at=created_at,
        )
    )


def prepare_pack_document(
    item: DocumentPackItemRequest,
    document_type: str,
    *,
    context: AuthoritativeGenerationContext,
) -> tuple[bytes, dict[str, object]]:
    rendered_content = build_document_from_template(
        template_content=context.template.content,
        content=item.content,
        document_type=document_type,
    )
    validation = validate_generated_document(
        template_content=context.template.content,
        rendered_content=rendered_content,
        generated_content=item.content,
        document_type=document_type,
        evidence=context.validation_evidence(),
    )
    return rendered_content, validation


def create_validation_artifact(
    *,
    application_id: str,
    item: DocumentPackItemRequest,
    context: AuthoritativeGenerationContext,
    rendered_content: bytes,
    validation: dict[str, object],
) -> DocumentValidationArtifactRecord:
    now = utc_now()
    template_hash, result_hash, evidence_hash = validation_artifact_hashes(
        item,
        context,
    )
    return DocumentValidationArtifactRecord(
        id=str(uuid4()),
        application_id=application_id,
        document_type="tailored_resume",
        template_id=context.template.id,
        template_hash=template_hash,
        result_hash=result_hash,
        evidence_hash=evidence_hash,
        rendered_hash=hashlib.sha256(rendered_content).hexdigest(),
        rendered_content=rendered_content,
        validation_report=validation,
        consumed_at=None,
        expires_at=now + VALIDATION_ARTIFACT_TTL,
        created_at=now,
    )


def consume_validation_artifact(
    db: Session,
    *,
    artifact_id: str,
    item: DocumentPackItemRequest,
    context: AuthoritativeGenerationContext,
) -> tuple[bytes, dict[str, object]]:
    artifact = db.scalar(
        select(DocumentValidationArtifactRecord)
        .where(DocumentValidationArtifactRecord.id == artifact_id)
        .with_for_update()
    )
    if artifact is None:
        raise ValueError("Resume validation artifact was not found")
    if artifact.consumed_at is not None:
        raise ValueError("Resume validation artifact has already been used")
    expires_at = artifact.expires_at
    if expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=utc_now().tzinfo)
    if expires_at <= utc_now():
        raise ValueError("Resume validation artifact has expired")
    if (
        artifact.application_id != context.application_id
        or artifact.document_type != "tailored_resume"
        or artifact.template_id != context.template.id
    ):
        raise ValueError("Resume validation artifact does not match the application or template")

    expected_hashes = validation_artifact_hashes(item, context)
    artifact_hashes = (
        artifact.template_hash,
        artifact.result_hash,
        artifact.evidence_hash,
    )
    if artifact_hashes != expected_hashes:
        raise ValueError(
            "Resume validation artifact hashes do not match the current template, result, or evidence"
        )
    if hashlib.sha256(artifact.rendered_content).hexdigest() != artifact.rendered_hash:
        raise ValueError("Resume validation artifact content is corrupted")
    artifact.consumed_at = utc_now()
    return artifact.rendered_content, dict(artifact.validation_report)


def validation_artifact_hashes(
    item: DocumentPackItemRequest,
    context: AuthoritativeGenerationContext,
) -> tuple[str, str, str]:
    return (
        hashlib.sha256(context.template.content).hexdigest(),
        hashlib.sha256(item.content.encode()).hexdigest(),
        canonical_json_hash(context.validation_evidence()),
    )


def canonical_json_hash(value: object) -> str:
    canonical = json.dumps(
        value,
        ensure_ascii=False,
        separators=(",", ":"),
        sort_keys=True,
    )
    return hashlib.sha256(canonical.encode()).hexdigest()


def persist_pack_document(
    *,
    db: Session,
    item: DocumentPackItemRequest,
    document_type: str,
    context: AuthoritativeGenerationContext,
    rendered_content: bytes,
    validation: dict[str, object],
    created_at: datetime,
) -> str:
    provenance = context.provenance()
    if item.document_id:
        record = require_pack_document_ownership(
            db,
            item,
            document_type=document_type,
            application_id=context.application_id,
        )
        assert record is not None
        next_version = record.current_version + 1
        record.title = item.title.strip()
        record.job_id = context.job_id
    else:
        document_id = str(uuid4())
        record = DocumentRecord(
            id=document_id,
            type=document_type,
            title=item.title.strip(),
            job_id=context.job_id,
            current_version=1,
            created_at=created_at,
            updated_at=created_at,
        )
        db.add(record)
        next_version = 1

    record.versions.append(
        DocumentVersionRecord(
            id=str(uuid4()),
            document_id=record.id,
            version=next_version,
            content=item.content,
            created_at=created_at,
        )
    )
    set_current_generation_provenance(
        record,
        provenance.generation_fingerprint,
        item.generation_model,
        provenance.input_versions,
        created_at,
    )
    append_version_generation_provenance(
        record,
        next_version,
        provenance.generation_fingerprint,
        item.generation_model,
        provenance.input_versions,
        created_at,
    )
    append_document_validation(record, next_version, validation, created_at)
    db.add(
        DocumentFileRecord(
            id=str(uuid4()),
            document_id=record.id,
            version=next_version,
            template_id=context.template.id,
            content=rendered_content,
            created_at=created_at,
        )
    )
    attach_document_record(db, record, context.application_id)
    record.current_version = next_version
    record.updated_at = created_at
    return record.id


def require_pack_document_ownership(
    db: Session,
    item: DocumentPackItemRequest,
    *,
    document_type: str,
    application_id: str,
) -> DocumentRecord | None:
    if not item.document_id:
        return None
    record = require_document(db, item.document_id)
    if record.type != document_type:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail="Existing document type does not match pack item type",
        )
    require_document_application_ownership(record, application_id)
    return record


def document_pack_request_fingerprint(
    request: DocumentPackRequest,
    *,
    resume_context: AuthoritativeGenerationContext,
    cover_context: AuthoritativeGenerationContext | None,
) -> str:
    payload = {
        "request": request.model_dump(mode="json", by_alias=True),
        "resumeGenerationFingerprint": resume_context.provenance().generation_fingerprint,
        "coverGenerationFingerprint": (
            cover_context.provenance().generation_fingerprint if cover_context else None
        ),
    }
    canonical = json.dumps(
        payload,
        ensure_ascii=True,
        separators=(",", ":"),
        sort_keys=True,
    )
    return hashlib.sha256(canonical.encode()).hexdigest()


def document_pack_payload(job: DocumentPackJobRecord, db: Session) -> DocumentPackPayload:
    documents = []
    for document_id in job.document_ids:
        record = require_document(db, document_id)
        documents.append(
            document_payload(
                record,
                current_generation_fingerprint=authoritative_current_generation_fingerprint(
                    db,
                    record,
                    application_id=job.application_id,
                ),
            )
        )
    return DocumentPackPayload(
        pack_job_id=job.id,
        status=job.status,
        persistence_mode=job.persistence_mode,
        documents=documents,
        stages=[DocumentPackStagePayload.model_validate(stage) for stage in job.stages],
        message=job.message,
        expires_at=job.expires_at,
    )


def cleanup_expired_pack_storage(db: Session) -> None:
    now = utc_now()
    db.execute(
        delete(DocumentValidationArtifactRecord).where(
            DocumentValidationArtifactRecord.expires_at <= now
        )
    )
    db.execute(
        delete(DocumentPackJobRecord).where(DocumentPackJobRecord.expires_at <= now)
    )
    db.commit()


def pack_validation_failed(stage: str, exc: Exception) -> HTTPException:
    return HTTPException(
        status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
        detail={
            "stage": stage,
            "status": "rolled_back",
            "message": str(exc),
        },
    )


def require_document(db: Session, document_id: str) -> DocumentRecord:
    record = db.scalar(
        select(DocumentRecord)
        .where(DocumentRecord.id == document_id)
        .options(
            selectinload(DocumentRecord.versions),
            selectinload(DocumentRecord.attachments),
            selectinload(DocumentRecord.generation_provenance),
            selectinload(DocumentRecord.version_generation_provenance),
            selectinload(DocumentRecord.version_validations),
            selectinload(DocumentRecord.files),
        )
    )
    if not record:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found",
        )
    return record


def require_template(db: Session, template_id: str) -> DocumentTemplateRecord:
    record = db.get(DocumentTemplateRecord, template_id)
    if not record:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document template not found",
        )
    return record


def load_generation_context_or_http(
    db: Session,
    *,
    application_id: str,
    template_id: str,
    document_type: str,
    expected_job_id: str | None = None,
) -> AuthoritativeGenerationContext:
    try:
        return load_authoritative_generation_context(
            db,
            application_id=application_id,
            template_id=template_id,
            document_type=document_type,
            expected_job_id=expected_job_id,
        )
    except GenerationContextError as exc:
        raise HTTPException(status_code=exc.status_code, detail=str(exc)) from exc


def single_attachment_application_id(record: DocumentRecord) -> str | None:
    application_ids = {attachment.application_id for attachment in record.attachments}
    if len(application_ids) != 1:
        return None
    return next(iter(application_ids))


def require_document_application_ownership(
    record: DocumentRecord,
    application_id: str,
) -> None:
    if any(
        attachment.application_id == application_id for attachment in record.attachments
    ):
        return
    raise HTTPException(
        status_code=status.HTTP_409_CONFLICT,
        detail="Existing document is not attached to the application",
    )


def document_file_record(
    db: Session,
    document_id: str,
    version: int,
) -> DocumentFileRecord | None:
    return db.scalar(
        select(DocumentFileRecord).where(
            DocumentFileRecord.document_id == document_id,
            DocumentFileRecord.version == version,
        )
    )


def attach_document_record(
    db: Session,
    document: DocumentRecord,
    application_id: str,
) -> None:
    if not db.get(StoredApplicationRecord, application_id):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Application not found",
        )
    if any(item.application_id == application_id for item in document.attachments):
        return
    document.attachments.append(
        DocumentAttachmentRecord(
            id=str(uuid4()),
            document_id=document.id,
            application_id=application_id,
            created_at=utc_now(),
        )
    )


def require_stored_application(db: Session, application_id: str) -> StoredApplicationRecord:
    application = db.get(StoredApplicationRecord, application_id)
    if application is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Application not found",
        )
    return application


def current_version_record(record: DocumentRecord) -> DocumentVersionRecord:
    current = next(
        (version for version in record.versions if version.version == record.current_version),
        None,
    )
    if not current:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Document has no current version",
        )
    return current


def authoritative_current_generation_fingerprint(
    db: Session,
    record: DocumentRecord,
    *,
    application_id: str | None = None,
) -> str | None:
    if record.generation_provenance is None:
        return None
    current_file = document_file_record(db, record.id, record.current_version)
    if current_file is None or not current_file.template_id:
        return None
    resolved_application_id = application_id or single_attachment_application_id(record)
    if not resolved_application_id:
        return None
    try:
        context = load_authoritative_generation_context(
            db,
            application_id=resolved_application_id,
            template_id=current_file.template_id,
            document_type=record.type,
            expected_job_id=record.job_id,
        )
    except GenerationContextError:
        return None
    return context.provenance().generation_fingerprint


def document_payload(
    record: DocumentRecord,
    *,
    current_generation_fingerprint: str | None = None,
    version_limit: int | None = None,
    version_offset: int = 0,
) -> DocumentPayload:
    provenance = record.generation_provenance
    rendered_versions = {file.version for file in record.files}
    validations_by_version = {
        validation.version: validation for validation in record.version_validations
    }
    ordered_versions = sorted(record.versions, key=lambda version: version.version, reverse=True)
    paged_versions = ordered_versions[version_offset:]
    if version_limit is not None:
        paged_versions = paged_versions[:version_limit]
    paged_versions = list(reversed(paged_versions))
    return DocumentPayload(
        id=record.id,
        type=record.type,
        title=record.title,
        job_id=record.job_id,
        application_ids=[item.application_id for item in record.attachments],
        current_version=record.current_version,
        created_at=record.created_at,
        updated_at=record.updated_at,
        generation_fingerprint=provenance.generation_fingerprint if provenance else None,
        current_generation_fingerprint=current_generation_fingerprint,
        generation_model=provenance.generation_model if provenance else None,
        input_versions=provenance.input_versions if provenance else {},
        versions=document_version_payloads(
            paged_versions,
            rendered_versions=rendered_versions,
            validations_by_version=validations_by_version,
        ),
        versions_total=len(record.versions),
        versions_has_more=version_offset + len(paged_versions) < len(record.versions),
    )


def document_version_payloads(
    versions: list[DocumentVersionRecord],
    *,
    rendered_versions: set[int],
    validations_by_version: dict[int, DocumentVersionValidationRecord],
) -> list[DocumentVersionPayload]:
    return [
        DocumentVersionPayload(
            id=version.id,
            version=version.version,
            content=version.content,
            created_at=version.created_at,
            has_rendered_docx=version.version in rendered_versions,
            factual_validation=(
                validations_by_version[version.version].factual_report
                if version.version in validations_by_version
                else {}
            ),
            visual_validation=(
                validations_by_version[version.version].visual_report
                if version.version in validations_by_version
                else {}
            ),
            diff=(
                validations_by_version[version.version].diff_items
                if version.version in validations_by_version
                else []
            ),
        )
        for version in versions
    ]


def safe_filename(value: str) -> str:
    normalized = re.sub(r"[^A-Za-z0-9._-]+", "-", value.strip()).strip("-._")
    return normalized[:120] or "tasko-document"


def safe_upload_filename(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9._ -]+", "-", value.strip())[:240] or "template.docx"


def decode_data_url(value: str) -> tuple[str, bytes]:
    header, separator, payload = value.partition(",")
    if not separator:
        return "", b""
    content_type = header.removeprefix("data:").split(";")[0]
    try:
        if ";base64" in header:
            return content_type, base64.b64decode(payload, validate=True)
        return content_type, payload.encode()
    except (ValueError, TypeError):
        return content_type, b""


def document_template_payload(record: DocumentTemplateRecord) -> DocumentTemplatePayload:
    return DocumentTemplatePayload(
        id=record.id,
        type=record.type,
        name=record.name,
        file_name=record.file_name,
        created_at=record.created_at,
        updated_at=record.updated_at,
    )


def workspace_source_document_payload(
    record: WorkspaceSourceDocumentRecord,
) -> WorkspaceSourceDocumentPayload:
    encoded_content = base64.b64encode(record.content).decode()
    return WorkspaceSourceDocumentPayload(
        id=record.id,
        application_id=record.application_id,
        category=record.category,
        title=record.title,
        language=record.language,
        file_name=record.file_name,
        file_size=f"{max(1, round(len(record.content) / 1024))} KB",
        file_type=record.content_type,
        uploaded_at=record.updated_at,
        data_url=f"data:{record.content_type};base64,{encoded_content}",
    )


def assistant_preflight_inputs(
    context: AuthoritativeGenerationContext,
) -> tuple[
    ProfilePayload,
    AssistantJobContext,
    AssistantApplicationContext,
    AssistantSourceDocument,
    list[AssistantCandidateConfirmation],
]:
    profile = ProfilePayload.model_validate(context.profile)
    job_data = dict(context.vacancy)
    job_data["aiMatch"] = {"applicationGuide": context.application_guide}
    job = AssistantJobContext.model_validate(job_data)
    application_data = dict(context.application)
    application_data["job"] = job.model_dump(by_alias=True)
    application = AssistantApplicationContext.model_validate(application_data)
    confirmations = [
        AssistantCandidateConfirmation(
            questionId=confirmation.question_id,
            requirement=confirmation.requirement,
            question=confirmation.question,
            answer=(
                confirmation.response.upper()
                + (
                    f": {confirmation.example_text}"
                    if confirmation.example_text
                    else ""
                )
            ),
        )
        for confirmation in context.confirmations
    ]
    encoded_template = base64.b64encode(context.template.content).decode("ascii")
    source = AssistantSourceDocument(
        id=context.template.id,
        title=context.template.name,
        category=(
            "Cover Letter"
            if context.template.type == "cover_letter"
            else "CV / Resume"
        ),
        fileName=context.template.file_name,
        dataUrl=(
            f"data:{context.template.content_type};base64,{encoded_template}"
        ),
    )
    return profile, job, application, source, confirmations


def database_unavailable(exc: Exception) -> HTTPException:
    return HTTPException(
        status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
        detail="Documents database is unavailable",
    )
