import json
import re
import zipfile
from copy import deepcopy
from io import BytesIO

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

from app.services.cover_letter_blocks import (
    parse_cover_letter_blocks,
    replace_cover_letter_text_span,
)
from app.services.resume_blocks import (
    parse_resume_blocks,
    replace_resume_text_span,
    set_text_node_value,
)


BODY_FONT = "Calibri"
HEADING_BLUE = RGBColor(0x2E, 0x74, 0xB5)
HEADING_DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x66, 0x66)


def build_document_docx(
    *,
    title: str,
    content: str,
    document_type: str,
    version: int,
) -> bytes:
    document = Document()
    configure_document(document)
    add_title_block(document, title, document_type)
    add_document_content(document, content)
    add_footer(document, document_type, version)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_document_from_template(
    *,
    template_content: bytes,
    content: str,
    document_type: str,
) -> bytes:
    """Replace text inside a copied DOCX while retaining its visual structure."""
    if document_type == "tailored_resume":
        return build_resume_from_template_package(template_content, content)
    if is_structured_cover_letter_content(content):
        return build_cover_letter_from_template_package(template_content, content)

    document = Document(BytesIO(template_content))
    replace_cover_letter_text(document, content)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_cover_letter_from_template_package(template_content: bytes, content: str) -> bytes:
    """Patch editable cover-letter spans without rebuilding the DOCX package."""
    with zipfile.ZipFile(BytesIO(template_content)) as source:
        document_xml = source.read("word/document.xml")
        root = etree.fromstring(document_xml)
        body = root.find(qn("w:body"))
        if body is None:
            raise ValueError("Cover-letter template has no document body")
        replace_cover_letter_spans(body, content)
        rendered_xml = etree.tostring(
            root,
            encoding="UTF-8",
            xml_declaration=True,
            standalone=True,
        )

        output = BytesIO()
        with zipfile.ZipFile(output, "w") as target:
            for item in source.infolist():
                target.writestr(
                    item,
                    rendered_xml if item.filename == "word/document.xml" else source.read(item.filename),
                )
    return output.getvalue()


def build_resume_from_template_package(template_content: bytes, content: str) -> bytes:
    """Patch only document.xml so every other DOCX package part remains untouched."""
    with zipfile.ZipFile(BytesIO(template_content)) as source:
        document_xml = source.read("word/document.xml")
        root = etree.fromstring(document_xml)
        body = root.find(qn("w:body"))
        if body is None:
            raise ValueError("Resume template has no document body")
        replace_resume_text(body, content)
        rendered_xml = etree.tostring(
            root,
            encoding="UTF-8",
            xml_declaration=True,
            standalone=True,
        )

        output = BytesIO()
        with zipfile.ZipFile(output, "w") as target:
            for item in source.infolist():
                target.writestr(
                    item,
                    rendered_xml if item.filename == "word/document.xml" else source.read(item.filename),
                )
    return output.getvalue()


def replace_cover_letter_spans(body, content: str) -> None:
    paragraphs = parse_cover_letter_blocks(body)
    paragraphs_by_id = {paragraph.paragraph_id: paragraph for paragraph in paragraphs}
    seen_span_ids: set[str] = set()
    for replacement in parse_cover_letter_replacements(content):
        paragraph_id = replacement["paragraphId"]
        paragraph = paragraphs_by_id.get(paragraph_id)
        if paragraph is None:
            raise ValueError(f"Unknown cover-letter paragraph: {paragraph_id}")
        span_id = replacement["spanId"]
        if span_id in seen_span_ids:
            raise ValueError(f"Duplicate cover-letter replacement for {span_id}")
        seen_span_ids.add(span_id)
        span = next(
            (candidate for candidate in paragraph.spans if candidate.span_id == span_id),
            None,
        )
        if span is None:
            raise ValueError(f"Unknown cover-letter span for {paragraph_id}: {span_id}")
        if replacement["original"] != span.original:
            raise ValueError(f"Cover-letter span original does not match template: {span_id}")
        replace_cover_letter_text_span(span, replacement["replacement"])


def replace_cover_letter_text(document: Document, content: str) -> None:
    paragraphs = document.paragraphs
    generated = normalized_content_paragraphs(content)
    if not generated:
        raise ValueError("Generated document content is empty")

    marker_index = next(
        (
            index
            for index, paragraph in enumerate(paragraphs)
            if "{{cover_letter_body}}" in paragraph.text.lower()
            or "{{content}}" in paragraph.text.lower()
        ),
        None,
    )
    if marker_index is not None:
        replace_paragraph_elements(
            paragraphs[marker_index : marker_index + 1],
            generated,
        )
    else:
        greeting_index = next(
            (
                index
                for index, paragraph in enumerate(paragraphs)
                if re.match(r"^\s*(?:dear|hello|hi|to\b)", paragraph.text, re.IGNORECASE)
            ),
            None,
        )
        closing_index = next(
            (
                index
                for index, paragraph in enumerate(paragraphs)
                if greeting_index is not None
                and index > greeting_index
                and re.match(
                    r"^\s*(?:sincerely|best|kind regards|warm regards|regards|thank you)",
                    paragraph.text,
                    re.IGNORECASE,
                )
            ),
            None,
        )
        if greeting_index is None or closing_index is None or closing_index <= greeting_index:
            raise ValueError(
                "Template needs {{cover_letter_body}} or recognizable greeting and closing paragraphs"
            )

        generated_greeting, generated_body, generated_closing = split_generated_letter(generated)
        if generated_greeting:
            set_paragraph_element_text(paragraphs[greeting_index]._p, generated_greeting)
        body_paragraphs = paragraphs[greeting_index + 1 : closing_index]
        if not body_paragraphs:
            body_paragraphs = [paragraphs[greeting_index]]
        replace_paragraph_elements(body_paragraphs, generated_body or generated)
        if generated_closing:
            set_paragraph_element_text(paragraphs[closing_index]._p, generated_closing)



def replace_resume_text(body, content: str) -> None:
    blocks = parse_resume_blocks(body)
    blocks_by_id = {block.block_id: block for block in blocks}
    seen_span_ids: set[str] = set()
    for replacement in parse_resume_replacements(content):
        block_id = replacement["blockId"]
        block = blocks_by_id.get(block_id)
        if block is None:
            raise ValueError(f"Unknown resume block: {block_id}")
        span_id = replacement["spanId"]
        if span_id in seen_span_ids:
            raise ValueError(f"Duplicate resume replacement for {span_id}")
        seen_span_ids.add(span_id)
        span = next((candidate for candidate in block.spans if candidate.span_id == span_id), None)
        if span is None:
            raise ValueError(f"Unknown resume span for {block_id}: {span_id}")
        if replacement["original"] != span.original:
            raise ValueError(f"Resume span original does not match template: {span_id}")
        replace_resume_text_span(span, replacement["replacement"])


def parse_resume_replacements(content: str) -> list[dict[str, object]]:
    cleaned = content.strip()
    fenced = re.fullmatch(r"```(?:json)?\s*(.*?)\s*```", cleaned, re.DOTALL | re.IGNORECASE)
    if fenced:
        cleaned = fenced.group(1).strip()
    try:
        payload = json.loads(cleaned)
    except json.JSONDecodeError as exc:
        raise ValueError("Generated resume must be structured JSON") from exc

    if isinstance(payload, dict) and isinstance(payload.get("replacements"), list):
        raw_replacements = payload["replacements"]
    elif isinstance(payload, list):
        raw_replacements = payload
    elif isinstance(payload, dict) and "blockId" in payload:
        raw_replacements = [payload]
    else:
        raise ValueError("Generated resume JSON must contain a replacements array")

    replacements: list[dict[str, object]] = []
    required_fields = ("blockId", "spanId", "original", "replacement", "reason")
    for index, raw_replacement in enumerate(raw_replacements):
        if not isinstance(raw_replacement, dict) or not all(
            isinstance(raw_replacement.get(field), str) for field in required_fields
        ):
            raise ValueError(
                f"Resume replacement {index + 1} must contain string fields: "
                "blockId, spanId, original, replacement, reason"
            )
        evidence_ids = raw_replacement.get("evidenceIds")
        if (
            not isinstance(evidence_ids, list)
            or not 1 <= len(evidence_ids) <= 20
            or not all(
                isinstance(evidence_id, str)
                and evidence_id == evidence_id.strip()
                and 1 <= len(evidence_id) <= 200
                for evidence_id in evidence_ids
            )
            or len(set(evidence_ids)) != len(evidence_ids)
        ):
            raise ValueError(
                f"Resume replacement {index + 1} must contain 1-20 unique "
                "evidenceIds strings"
            )
        replacements.append(
            {
                **{field: raw_replacement[field] for field in required_fields},
                "evidenceIds": evidence_ids,
            }
        )
    return replacements


def is_structured_cover_letter_content(content: str) -> bool:
    cleaned = content.strip()
    if cleaned.lower().startswith("```json"):
        return True
    if not cleaned.startswith(("{", "[")):
        return False
    try:
        payload = parse_json_content(cleaned)
    except ValueError:
        return True
    return isinstance(payload, (dict, list))


def parse_cover_letter_replacements(content: str) -> list[dict[str, object]]:
    payload = parse_json_content(
        content, error_message="Generated cover letter must be structured JSON"
    )
    if isinstance(payload, dict) and isinstance(payload.get("replacements"), list):
        raw_replacements = payload["replacements"]
    elif isinstance(payload, list):
        raw_replacements = payload
    elif isinstance(payload, dict) and "paragraphId" in payload:
        raw_replacements = [payload]
    else:
        raise ValueError("Generated cover-letter JSON must contain a replacements array")

    replacements: list[dict[str, object]] = []
    required_fields = ("paragraphId", "spanId", "original", "replacement", "reason")
    for index, raw_replacement in enumerate(raw_replacements):
        if not isinstance(raw_replacement, dict) or not all(
            isinstance(raw_replacement.get(field), str) for field in required_fields
        ):
            raise ValueError(
                f"Cover-letter replacement {index + 1} must contain string fields: "
                "paragraphId, spanId, original, replacement, reason"
            )
        evidence_ids = raw_replacement.get("evidenceIds")
        if (
            not isinstance(evidence_ids, list)
            or not 1 <= len(evidence_ids) <= 20
            or not all(
                isinstance(evidence_id, str)
                and evidence_id == evidence_id.strip()
                and 1 <= len(evidence_id) <= 200
                for evidence_id in evidence_ids
            )
            or len(set(evidence_ids)) != len(evidence_ids)
        ):
            raise ValueError(
                f"Cover-letter replacement {index + 1} must contain 1-20 unique evidenceIds strings"
            )
        replacements.append(
            {
                **{field: raw_replacement[field] for field in required_fields},
                "evidenceIds": evidence_ids,
            }
        )
    return replacements


def parse_json_content(content: str, *, error_message: str = "Generated content must be JSON"):
    cleaned = content.strip()
    fenced = re.fullmatch(r"```(?:json)?\s*(.*?)\s*```", cleaned, re.DOTALL | re.IGNORECASE)
    if fenced:
        cleaned = fenced.group(1).strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError as exc:
        raise ValueError(error_message) from exc


def normalized_content_paragraphs(content: str) -> list[str]:
    lines = [line.strip() for line in content.replace("\r\n", "\n").split("\n")]
    paragraphs: list[str] = []
    current: list[str] = []
    for line in lines:
        if not line:
            if current:
                paragraphs.append(" ".join(current))
                current = []
            continue
        cleaned = re.sub(r"^#{1,3}\s+", "", line)
        current.append(cleaned)
    if current:
        paragraphs.append(" ".join(current))
    return paragraphs


def split_generated_letter(paragraphs: list[str]) -> tuple[str, list[str], str]:
    greeting = paragraphs[0] if re.match(
        r"^(?:dear|hello|hi|to\b)", paragraphs[0], re.IGNORECASE
    ) else ""
    closing_index = next(
        (
            index
            for index, paragraph in enumerate(paragraphs)
            if re.match(
                r"^(?:sincerely|best|kind regards|warm regards|regards|thank you)",
                paragraph,
                re.IGNORECASE,
            )
        ),
        None,
    )
    body_start = 1 if greeting else 0
    body_end = closing_index if closing_index is not None else len(paragraphs)
    closing = paragraphs[closing_index] if closing_index is not None else ""
    return greeting, paragraphs[body_start:body_end], closing


def replace_paragraph_elements(paragraphs, replacement_texts: list[str]) -> None:
    if not paragraphs:
        raise ValueError("Template has no editable body paragraphs")
    elements = [getattr(paragraph, "_p", paragraph) for paragraph in paragraphs]
    prototype = elements[0]
    parent = prototype.getparent()
    insert_at = parent.index(prototype)
    for element in elements:
        parent.remove(element)
    for offset, replacement in enumerate(replacement_texts):
        clone = deepcopy(prototype)
        set_paragraph_element_text(clone, replacement)
        parent.insert(insert_at + offset, clone)


def set_paragraph_element_text(element, text: str) -> None:
    text_nodes = list(element.iter(qn("w:t")))
    if not text_nodes:
        run = OxmlElement("w:r")
        node = OxmlElement("w:t")
        run.append(node)
        element.append(run)
        text_nodes = [node]
    first, *remaining = text_nodes
    set_text_node_value(first, text)
    for node in remaining:
        set_text_node_value(node, "")


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.start_type = WD_SECTION.NEW_PAGE

    normal = document.styles["Normal"]
    set_style_font(normal, BODY_FONT, 11, RGBColor(0, 0, 0))
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        "Heading 1": (16, HEADING_BLUE, 16, 8),
        "Heading 2": (13, HEADING_BLUE, 12, 6),
        "Heading 3": (12, HEADING_DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        set_style_font(style, BODY_FONT, size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        set_style_font(style, BODY_FONT, 11, RGBColor(0, 0, 0))
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_title_block(document: Document, title: str, document_type: str) -> None:
    label = "Cover Letter" if document_type == "cover_letter" else "Tailored Resume"
    label_paragraph = document.add_paragraph()
    label_paragraph.paragraph_format.space_after = Pt(4)
    label_run = label_paragraph.add_run(label.upper())
    set_run_font(label_run, BODY_FONT, 9, MUTED, bold=True)

    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_after = Pt(14)
    title_paragraph.paragraph_format.keep_with_next = True
    title_run = title_paragraph.add_run(title)
    set_run_font(title_run, BODY_FONT, 22, RGBColor(0, 0, 0), bold=True)


def add_document_content(document: Document, content: str) -> None:
    for raw_line in content.replace("\r\n", "\n").split("\n"):
        line = raw_line.rstrip()
        if not line:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(3)
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading_match:
            document.add_paragraph(
                heading_match.group(2).strip(),
                style=f"Heading {len(heading_match.group(1))}",
            )
            continue

        bullet_match = re.match(r"^(?:[-*•])\s+(.+)$", line)
        if bullet_match:
            document.add_paragraph(bullet_match.group(1).strip(), style="List Bullet")
            continue

        numbered_match = re.match(r"^\d+[.)]\s+(.+)$", line)
        if numbered_match:
            document.add_paragraph(numbered_match.group(1).strip(), style="List Number")
            continue

        paragraph = document.add_paragraph(line)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_footer(document: Document, document_type: str, version: int) -> None:
    label = "Cover letter" if document_type == "cover_letter" else "Tailored resume"
    paragraph = document.sections[0].footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"Tasko · {label} · v{version} · ")
    set_run_font(run, BODY_FONT, 8, MUTED)
    append_page_field(paragraph)


def append_page_field(paragraph) -> None:
    run = paragraph.add_run()
    set_run_font(run, BODY_FONT, 8, MUTED)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend((field_begin, instruction, field_end))


def set_style_font(style, name: str, size: int, color: RGBColor, *, bold: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def set_run_font(
    run,
    name: str,
    size: int,
    color: RGBColor,
    *,
    bold: bool = False,
) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
