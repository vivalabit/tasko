import hashlib
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from collections import Counter
from datetime import date
from difflib import SequenceMatcher
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from lxml import etree
from pypdf import PdfReader

from app.services.cover_letter_blocks import extract_cover_letter_blocks_from_docx
from app.services.document_export import (
    is_structured_cover_letter_content,
    parse_cover_letter_replacements,
    parse_resume_replacements,
)
from app.services.resume_blocks import extract_resume_blocks_from_docx


TECHNOLOGIES = {
    "angular",
    "aws",
    "azure",
    "c#",
    "c++",
    "css",
    "django",
    "docker",
    "excel",
    "fastapi",
    "figma",
    "flask",
    "gcp",
    "git",
    "go",
    "graphql",
    "html",
    "java",
    "javascript",
    "kotlin",
    "kubernetes",
    "mongodb",
    "mysql",
    "next.js",
    "node.js",
    "postgresql",
    "power bi",
    "python",
    "react",
    "redis",
    "rest",
    "rust",
    "salesforce",
    "sap",
    "sketch",
    "sql",
    "swift",
    "tableau",
    "terraform",
    "typescript",
    "vue",
}
CLAIM_VERBS = {
    "achieved",
    "architected",
    "automated",
    "baute",
    "built",
    "created",
    "delivered",
    "designed",
    "developed",
    "drove",
    "entwickelte",
    "generated",
    "grew",
    "implemented",
    "implementierte",
    "improved",
    "increased",
    "launched",
    "led",
    "leitete",
    "lieferte",
    "managed",
    "mentored",
    "optimized",
    "optimierte",
    "owned",
    "reduced",
    "reduzierte",
    "saved",
    "scaled",
    "skalierte",
    "spearheaded",
    "steigerte",
    "transformed",
}
STOP_WORDS = {
    "about",
    "across",
    "and",
    "for",
    "from",
    "für",
    "into",
    "mit",
    "und",
    "that",
    "the",
    "their",
    "through",
    "using",
    "with",
}
EXPERIENCE_CLAIM_TYPES = {
    "employer",
    "title",
    "period",
    "technology",
    "achievement",
}
FACTUAL_ASSERTION_PATTERN = re.compile(
    r"\b(?:build|certified|creat(?:e|ed|ing)|deliver(?:ed|ing)?|deploy(?:ed|ing)?|"
    r"design(?:ed|ing)?|develop(?:ed|ing)?|experienced|expert(?:ise)?|held|"
    r"implement(?:ed|ing)?|improv(?:e|ed|ing)|know|led|maintained|manag(?:e|ed|ing)|"
    r"master(?:ed|y)|operated|optimi[sz](?:e|ed|ing)|proficient|reduc(?:e|ed|ing)|"
    r"responsible|scal(?:e|ed|ing)|served|skilled|speciali[sz](?:e|ed|ing|t)|"
    r"support(?:ed|ing)?|used|worked|"
    r"arbeitete|eingesetzt|erfahren|experte|verantwortlich|zertifiziert)\b",
    re.IGNORECASE,
)
CURRENT_PERIOD_PATTERN = re.compile(
    r"\b(?:current(?:ly)?|present(?:ly)?|ongoing|today|now|"
    r"aktuell|derzeit|gegenw[aä]rtig|heute|laufend|bis heute)\b",
    re.IGNORECASE,
)
PAST_PERIOD_PATTERN = re.compile(
    r"\b(?:former(?:ly)?|previously|past|ehemalig|fr[uü]her|zuvor)\b",
    re.IGNORECASE,
)
NEGATION_PATTERN = re.compile(
    r"\b(?:no|not|never|without|cannot|can't|couldn't|didn't|doesn't|don't|"
    r"hasn't|haven't|isn't|wasn't|weren't|"
    r"kein(?:e|en|em|er|es)?|nicht|nie|niemals|ohne|"
    r"ne|pas|jamais|sans|non|senza|nunca|sin)\b",
    re.IGNORECASE,
)
MODALITY_PATTERNS = {
    "capability": re.compile(
        r"\b(?:can|able to|capable of|kann|f[aä]hig|capable de|puede)\b",
        re.IGNORECASE,
    ),
    "hypothetical": re.compile(
        r"\b(?:could|may|might|would|perhaps|possibly|probably|"
        r"k[oö]nnte|w[uü]rde|m[oö]glicherweise|vermutlich|"
        r"pourrait|peut-être|podr[ií]a)\b",
        re.IGNORECASE,
    ),
    "obligation": re.compile(
        r"\b(?:must|should|required to|muss|sollte|doit|debe)\b",
        re.IGNORECASE,
    ),
    "future": re.compile(
        r"\b(?:will|going to|intend(?:s|ed)? to|werde|wird|va a)\b",
        re.IGNORECASE,
    ),
}
ATTRIBUTION_GLUE_WORDS = {
    "able",
    "can",
    "could",
    "currently",
    "did",
    "former",
    "formerly",
    "may",
    "might",
    "must",
    "never",
    "not",
    "presently",
    "previously",
    "should",
    "will",
    "worked",
    "would",
}
MIN_CLAIM_TOKEN_COVERAGE = 0.7
DATE_PATTERN = re.compile(
    r"\b(?:19|20)\d{2}\b|\b(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|"
    r"may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|"
    r"nov(?:ember)?|dec(?:ember)?|januar|februar|märz|mai|juni|juli|oktober|"
    r"dezember)\s+(?:19|20)\d{2}\b|\b\d{1,2}[./-]\d{1,2}[./-]\d{2,4}\b",
    re.IGNORECASE,
)
NUMBER_PATTERN = re.compile(
    r"(?<![\w-])(?:[$€£]\s*)?\d+(?:[.,]\d+)?(?:\s*(?:%|x|k|m|million|billion|users?|"
    r"clients?|projects?|people|members?|months?|years?))?(?![\w-])",
    re.IGNORECASE,
)
COMPANY_PATTERN = re.compile(
    r"(?:\b(?:at|with|bei)\s+|\b(?:worked|working|consulted|employed|arbeitete)\s+"
    r"(?:for|für)\s+)([A-ZÄÖÜ][\w&.-]+"
    r"(?:\s+[A-ZÄÖÜ][\w&.-]+){0,3})",
)
TITLE_PATTERN = re.compile(
    r"\b(?:as|als)\s+(?:an?\s+)?((?:[A-ZÄÖÜ][\w/-]*\s+){0,4}"
    r"(?:Engineer|Designer|Manager|Director|Lead|Consultant|Developer|Architect|Analyst|"
    r"Entwickler|Ingenieur|Berater|Leiter))\b|"
    r"\b((?:Senior|Lead|Principal|Staff|Junior)\s+(?:Software\s+)?"
    r"(?:Engineer|Designer|Manager|Director|Consultant|Developer|Architect|Analyst))"
    r"\s+(?:at|bei)\b",
)
RENDER_TIMEOUT_SECONDS = 90
GEOMETRY_POSITION_TOLERANCE = 0.03
GEOMETRY_SIZE_TOLERANCE = 0.25
PAGE_BOUNDARY_TOLERANCE = 1.5


class DocumentValidationError(ValueError):
    pass


def validate_generated_document(
    *,
    template_content: bytes,
    rendered_content: bytes,
    generated_content: str,
    document_type: str,
    evidence: dict[str, Any],
) -> dict[str, Any]:
    diff = build_document_diff(
        template_content,
        generated_content,
        document_type,
        rendered_content=rendered_content,
    )
    structured_document = document_type == "tailored_resume" or (
        document_type == "cover_letter" and is_structured_cover_letter_content(generated_content)
    )
    if structured_document:
        replacements = (
            parse_resume_replacements(generated_content)
            if document_type == "tailored_resume"
            else parse_cover_letter_replacements(generated_content)
        )
        evidence_catalog = build_authoritative_evidence_catalog(
            template_content,
            evidence,
            document_type=document_type,
        )
        factual_issues = [
            *validate_evidence_id_references(replacements, evidence_catalog),
            *validate_referenced_factual_changes(diff, evidence_catalog),
        ]
        cited_evidence_ids = {
            evidence_id
            for replacement in replacements
            for evidence_id in replacement["evidenceIds"]
            if evidence_id in evidence_catalog
        }
        checked_evidence_characters = sum(
            len(evidence_catalog[evidence_id]["text"])
            for evidence_id in cited_evidence_ids
        )
    else:
        source_text = extract_docx_text(template_content)
        evidence_text = "\n".join(flatten_evidence(evidence))
        allowed_text = f"{source_text}\n{evidence_text}"
        factual_issues = validate_factual_changes(diff, allowed_text)
        checked_evidence_characters = len(allowed_text)
    visual_report, visual_issues = validate_visual_output(
        template_content,
        rendered_content,
        allowed_removed_text="\n".join(str(change["original"]) for change in diff),
    )
    issues = [*factual_issues, *visual_issues]
    if issues:
        raise DocumentValidationError("Document validation failed: " + "; ".join(issues[:8]))
    return {
        "factual": {
            "status": "passed",
            "checkedChanges": len(diff),
            "checkedEvidenceCharacters": checked_evidence_characters,
        },
        "visual": visual_report,
        "diff": diff,
    }


def build_document_diff(
    template_content: bytes,
    generated_content: str,
    document_type: str,
    *,
    rendered_content: bytes | None = None,
) -> list[dict[str, Any]]:
    if document_type == "tailored_resume":
        blocks = {
            block["blockId"]: block
            for block in extract_resume_blocks_from_docx(template_content)
        }
        return [
            {
                "blockId": replacement["blockId"],
                "spanId": replacement["spanId"],
                "type": blocks.get(replacement["blockId"], {}).get("type", "block"),
                "original": replacement["original"],
                "replacement": replacement["replacement"],
                "reason": replacement["reason"],
                "evidenceIds": replacement["evidenceIds"],
            }
            for replacement in parse_resume_replacements(generated_content)
            if replacement["replacement"] != replacement["original"]
        ]
    if document_type == "cover_letter" and is_structured_cover_letter_content(generated_content):
        paragraphs = {
            paragraph["paragraphId"]: paragraph
            for paragraph in extract_cover_letter_blocks_from_docx(template_content)
        }
        return [
            {
                "blockId": replacement["paragraphId"],
                "paragraphId": replacement["paragraphId"],
                "spanId": replacement["spanId"],
                "type": paragraphs.get(replacement["paragraphId"], {}).get("type", "paragraph"),
                "original": replacement["original"],
                "replacement": replacement["replacement"],
                "reason": replacement["reason"],
                "evidenceIds": replacement["evidenceIds"],
            }
            for replacement in parse_cover_letter_replacements(generated_content)
            if replacement["replacement"] != replacement["original"]
        ]

    original_paragraphs = extract_docx_paragraphs(template_content)
    if rendered_content is not None:
        generated_paragraphs = extract_docx_paragraphs(rendered_content)
    else:
        generated_paragraphs = [
            paragraph.strip()
            for paragraph in re.split(r"\n\s*\n", generated_content.strip())
            if paragraph.strip()
        ]
    matcher = SequenceMatcher(a=original_paragraphs, b=generated_paragraphs, autojunk=False)
    diff: list[dict[str, Any]] = []
    for change_index, (operation, left_start, left_end, right_start, right_end) in enumerate(
        matcher.get_opcodes(),
        start=1,
    ):
        if operation == "equal":
            continue
        diff.append(
            {
                "blockId": f"paragraph-change-{change_index:04d}",
                "type": "paragraph",
                "original": "\n\n".join(original_paragraphs[left_start:left_end]),
                "replacement": "\n\n".join(generated_paragraphs[right_start:right_end]),
                "reason": "Generated cover-letter paragraph update",
            }
        )
    return diff


def validate_factual_changes(diff: list[dict[str, Any]], allowed_text: str) -> list[str]:
    issues: list[str] = []
    for change in diff:
        issues.extend(validate_factual_change_against_text(change, allowed_text))
    return issues


def build_authoritative_evidence_catalog(
    template_content: bytes,
    evidence: dict[str, Any],
    *,
    document_type: str = "tailored_resume",
) -> dict[str, dict[str, Any]]:
    catalog: dict[str, dict[str, Any]] = {}
    containers = (
        extract_resume_blocks_from_docx(template_content)
        if document_type == "tailored_resume"
        else extract_cover_letter_blocks_from_docx(template_content)
    )
    for container in containers:
        for span in container["spans"]:
            evidence_id = span.get("evidenceId")
            original = span.get("original")
            if isinstance(evidence_id, str) and isinstance(original, str) and original:
                catalog[evidence_id] = {"type": "source", "text": original}

    for item in evidence.get("evidenceCatalog", []):
        if not isinstance(item, dict):
            continue
        evidence_id = item.get("id")
        evidence_type = item.get("type")
        text = item.get("text")
        if (
            isinstance(evidence_id, str)
            and evidence_type in {"profile", "confirmation"}
            and isinstance(text, str)
            and text.strip()
            and evidence_id.startswith(f"{evidence_type}:")
        ):
            catalog_item: dict[str, Any] = {
                "type": evidence_type,
                "text": text.strip(),
            }
            claim_type = item.get("claimType")
            experience_id = item.get("experienceId")
            if (
                evidence_type == "profile"
                and evidence_id.startswith("profile:experience:")
                and claim_type in EXPERIENCE_CLAIM_TYPES
                and isinstance(experience_id, str)
                and experience_id.strip()
            ):
                catalog_item["claimType"] = claim_type
                catalog_item["experienceId"] = experience_id.strip()
            catalog[evidence_id] = catalog_item
    return catalog


def validate_evidence_id_references(
    replacements: list[dict[str, object]],
    evidence_catalog: dict[str, dict[str, Any]],
) -> list[str]:
    return [
        f'{replacement["spanId"]} references unknown evidence "{evidence_id}"'
        for replacement in replacements
        for evidence_id in replacement["evidenceIds"]
        if evidence_id not in evidence_catalog
    ]


def validate_referenced_factual_changes(
    diff: list[dict[str, Any]],
    evidence_catalog: dict[str, dict[str, Any]],
) -> list[str]:
    issues: list[str] = []
    for change in diff:
        referenced_items = [
            evidence_catalog[evidence_id]
            for evidence_id in change["evidenceIds"]
            if evidence_id in evidence_catalog
        ]
        referenced_text = "\n".join(item["text"] for item in referenced_items)
        if referenced_text:
            issues.extend(
                validate_factual_change_against_text(
                    change,
                    referenced_text,
                    referenced=True,
                    allowed_periods=[
                        item["text"]
                        for item in referenced_items
                        if item.get("claimType") == "period"
                    ],
                )
            )
            issues.extend(validate_atomic_experience_attribution(change, referenced_items))
    return issues


def validate_factual_change_against_text(
    change: dict[str, Any],
    allowed_text: str,
    *,
    referenced: bool = False,
    allowed_periods: list[str] | None = None,
) -> list[str]:
    issues: list[str] = []
    original = change["original"]
    replacement = change["replacement"]
    location = change.get("spanId", change["blockId"])
    allowed_normalized = normalize_fact(allowed_text)
    allowed_tokens = semantic_tokens(allowed_text)
    for label, extractor in (
        ("date", extract_dates),
        ("number", extract_numbers),
        ("technology", extract_technologies),
        ("company", extract_companies),
        ("job title", extract_titles),
    ):
        original_values = {normalize_fact(value) for value in extractor(original)}
        for value in extractor(replacement):
            normalized = normalize_fact(value)
            supported_by_period = (
                label == "date"
                or (
                    label == "number"
                    and normalized
                    in {normalize_fact(date_value) for date_value in extract_dates(replacement)}
                )
            ) and any(date_is_within_period(value, period) for period in (allowed_periods or []))
            if (
                normalized not in original_values
                and not contains_normalized_fact(allowed_normalized, normalized)
                and not supported_by_period
            ):
                if referenced:
                    issues.append(
                        f'{location} adds {label} "{value}" not supported by referenced evidence'
                    )
                else:
                    issues.append(f'{location} adds unsupported {label} "{value}"')

    original_tokens = semantic_tokens(original)
    original_negated = has_negation(original)
    original_modality = modality_signature(original)
    for sentence in split_sentences(replacement):
        sentence_tokens = semantic_tokens(sentence) - CLAIM_VERBS - ATTRIBUTION_GLUE_WORDS
        if not sentence_tokens or not is_factual_assertion(sentence):
            continue
        if (
            sentence_tokens <= original_tokens
            and has_negation(sentence) == original_negated
            and modality_signature(sentence) == original_modality
        ):
            continue
        supported = sentence_tokens & allowed_tokens
        if len(supported) / len(sentence_tokens) < MIN_CLAIM_TOKEN_COVERAGE:
            if referenced:
                issues.append(
                    f"{location} adds a claim not supported by referenced evidence "
                    f'"{sentence[:100]}"'
                )
            else:
                issues.append(f'{location} adds an unsupported claim "{sentence[:100]}"')
            continue
        semantic_evidence = best_supporting_sentence(sentence, allowed_text)
        if semantic_evidence is None:
            continue
        qualifier = "referenced evidence" if referenced else "evidence"
        if has_negation(sentence) != has_negation(semantic_evidence):
            issues.append(f"{location} changes negation relative to {qualifier}")
        if modality_signature(sentence) != modality_signature(semantic_evidence):
            issues.append(f"{location} changes modality relative to {qualifier}")
    return issues


def validate_atomic_experience_attribution(
    change: dict[str, Any],
    referenced_items: list[dict[str, Any]],
) -> list[str]:
    atomic_items = [
        item
        for item in referenced_items
        if item.get("claimType") in EXPERIENCE_CLAIM_TYPES
        and isinstance(item.get("experienceId"), str)
    ]
    if not atomic_items:
        return []

    non_atomic_text = "\n".join(
        item["text"] for item in referenced_items if item not in atomic_items
    )
    groups: dict[str, list[dict[str, Any]]] = {}
    for item in atomic_items:
        groups.setdefault(item["experienceId"], []).append(item)

    issues: list[str] = []
    location = change.get("spanId", change["blockId"])
    original = str(change["original"])
    for sentence in split_sentences(str(change["replacement"])):
        if not is_factual_assertion(sentence):
            continue
        if sentence_supported_by_text(sentence, original):
            continue
        if non_atomic_text and sentence_supported_by_text(sentence, non_atomic_text):
            continue

        required_types = required_experience_claim_types(sentence, atomic_items)
        if not required_types:
            issues.append(
                f'{location} adds an unknown factual assertion and fails closed "{sentence[:100]}"'
            )
            continue

        typed_candidates = [
            (experience_id, items)
            for experience_id, items in groups.items()
            if required_types <= {item["claimType"] for item in items}
        ]
        if not typed_candidates:
            cited_types = {item["claimType"] for item in atomic_items}
            if required_types <= cited_types:
                issues.append(
                    f"{location} mixes claims attributed to different experience records; "
                    f"{', '.join(sorted(required_types))} must share one experienceId"
                )
            else:
                missing = ", ".join(sorted(required_types - cited_types))
                issues.append(
                    f"{location} adds an unknown claim without cited atomic evidence "
                    f"for {missing} (fail-closed)"
                )
            continue

        chronological_candidates = [
            candidate
            for candidate in typed_candidates
            if group_supports_chronology(sentence, candidate[1])
        ]
        if not chronological_candidates:
            issues.append(
                f"{location} adds a chronology or employment-period claim not supported "
                f'by the cited experience "{sentence[:100]}"'
            )
            continue

        attributed_candidates = [
            candidate
            for candidate in chronological_candidates
            if group_supports_attributed_values(
                sentence,
                candidate[1],
                atomic_items,
            )
        ]
        if not attributed_candidates:
            issues.append(
                f"{location} attributes a fact to an unsupported employer or experience "
                f'"{sentence[:100]}"'
            )
            continue

        supported_candidates = [
            candidate
            for candidate in attributed_candidates
            if atomic_group_token_coverage(sentence, candidate[1])
            >= MIN_CLAIM_TOKEN_COVERAGE
        ]
        if not supported_candidates:
            issues.append(
                f"{location} adds an unknown assertion not supported by cited atomic "
                f'evidence (fail-closed) "{sentence[:100]}"'
            )
            continue

        _, supporting_items = max(
            supported_candidates,
            key=lambda candidate: atomic_group_token_coverage(sentence, candidate[1]),
        )
        semantic_evidence = best_semantic_atomic_claim(sentence, supporting_items)
        if semantic_evidence is None:
            continue
        evidence_text = semantic_evidence["text"]
        if has_negation(sentence) != has_negation(evidence_text):
            issues.append(
                f'{location} changes negation relative to cited evidence "{evidence_text[:80]}"'
            )
        if modality_signature(sentence) != modality_signature(evidence_text):
            issues.append(
                f'{location} changes modality relative to cited evidence "{evidence_text[:80]}"'
            )
    return issues


def required_experience_claim_types(
    sentence: str,
    atomic_items: list[dict[str, Any]],
) -> set[str]:
    normalized = normalize_fact(sentence)
    required: set[str] = set()
    for claim_type in ("employer", "title", "technology"):
        if any(
            item["claimType"] == claim_type
            and contains_normalized_fact(normalized, normalize_fact(item["text"]))
            for item in atomic_items
        ):
            required.add(claim_type)
    if extract_companies(sentence):
        required.add("employer")
    if extract_titles(sentence):
        required.add("title")
    if (
        extract_dates(sentence)
        or CURRENT_PERIOD_PATTERN.search(sentence)
        or PAST_PERIOD_PATTERN.search(sentence)
    ):
        required.add("period")
    if extract_technologies(sentence):
        required.add("technology")
    if requires_achievement_evidence(sentence):
        required.add("achievement")
    return required


def requires_achievement_evidence(sentence: str) -> bool:
    tokens = significant_tokens(sentence)
    if tokens & CLAIM_VERBS:
        return True
    match = FACTUAL_ASSERTION_PATTERN.search(sentence)
    return bool(
        match and match.group(0).casefold() not in {"worked", "arbeitete", "held", "served"}
    )


def group_supports_chronology(
    sentence: str,
    items: list[dict[str, Any]],
) -> bool:
    periods = [item["text"] for item in items if item["claimType"] == "period"]
    dates = extract_dates(sentence)
    if dates and not all(
        any(date_is_within_period(value, period) for period in periods) for value in dates
    ):
        return False
    if CURRENT_PERIOD_PATTERN.search(sentence) and not any(
        CURRENT_PERIOD_PATTERN.search(period) for period in periods
    ):
        return False
    if PAST_PERIOD_PATTERN.search(sentence) and any(
        CURRENT_PERIOD_PATTERN.search(period) for period in periods
    ):
        return False
    return True


def group_supports_attributed_values(
    sentence: str,
    items: list[dict[str, Any]],
    all_atomic_items: list[dict[str, Any]],
) -> bool:
    normalized = normalize_fact(sentence)
    claims_by_type = {
        claim_type: [
            normalize_fact(item["text"]) for item in items if item["claimType"] == claim_type
        ]
        for claim_type in EXPERIENCE_CLAIM_TYPES
    }
    for claim_type in ("employer", "title", "technology"):
        all_mentioned_claims = {
            normalize_fact(item["text"])
            for item in all_atomic_items
            if item["claimType"] == claim_type
            and contains_normalized_fact(normalized, normalize_fact(item["text"]))
        }
        if not all_mentioned_claims <= set(claims_by_type[claim_type]):
            return False
        mentioned_claims = [
            claim
            for claim in claims_by_type[claim_type]
            if contains_normalized_fact(normalized, claim)
        ]
        if mentioned_claims:
            continue
        extractor = {
            "employer": extract_companies,
            "title": extract_titles,
            "technology": extract_technologies,
        }[claim_type]
        values = extractor(sentence)
        if values and not all(
            any(
                contains_normalized_fact(claim, normalize_fact(value))
                or contains_normalized_fact(normalize_fact(value), claim)
                for claim in claims_by_type[claim_type]
            )
            for value in values
        ):
            return False

    group_text = "\n".join(item["text"] for item in items)
    group_normalized = normalize_fact(group_text)
    sentence_dates = {normalize_fact(value) for value in extract_dates(sentence)}
    for number in extract_numbers(sentence):
        normalized_number = normalize_fact(number)
        if contains_normalized_fact(group_normalized, normalized_number):
            continue
        if normalized_number in sentence_dates and any(
            date_is_within_period(number, period) for period in claims_by_type["period"]
        ):
            continue
        return False
    return True


def atomic_group_token_coverage(
    sentence: str,
    items: list[dict[str, Any]],
) -> float:
    sentence_tokens = semantic_tokens(sentence) - CLAIM_VERBS - ATTRIBUTION_GLUE_WORDS
    if not sentence_tokens:
        return 1.0
    evidence_tokens = semantic_tokens("\n".join(item["text"] for item in items))
    return len(sentence_tokens & evidence_tokens) / len(sentence_tokens)


def best_semantic_atomic_claim(
    sentence: str,
    items: list[dict[str, Any]],
) -> dict[str, Any] | None:
    candidates = [item for item in items if item["claimType"] == "achievement"] or items
    sentence_tokens = semantic_tokens(sentence) - CLAIM_VERBS - ATTRIBUTION_GLUE_WORDS
    if not candidates:
        return None
    return max(
        candidates,
        key=lambda item: len(sentence_tokens & semantic_tokens(item["text"])),
    )


def sentence_supported_by_text(sentence: str, text: str) -> bool:
    sentence_tokens = semantic_tokens(sentence) - CLAIM_VERBS - ATTRIBUTION_GLUE_WORDS
    if not sentence_tokens:
        return True
    text_tokens = semantic_tokens(text)
    if len(sentence_tokens & text_tokens) / len(sentence_tokens) < 0.85:
        return False
    if has_negation(sentence) != has_negation(text):
        return False
    return modality_signature(sentence) == modality_signature(text)


def best_supporting_sentence(sentence: str, evidence_text: str) -> str | None:
    sentence_tokens = semantic_tokens(sentence) - CLAIM_VERBS - ATTRIBUTION_GLUE_WORDS
    if not sentence_tokens:
        return None
    candidates = split_sentences(evidence_text)
    if not candidates:
        return None
    candidate = max(
        candidates,
        key=lambda value: len(sentence_tokens & semantic_tokens(value)),
    )
    if (
        len(sentence_tokens & semantic_tokens(candidate)) / len(sentence_tokens)
        < MIN_CLAIM_TOKEN_COVERAGE
    ):
        return None
    return candidate


def date_is_within_period(value: str, period: str) -> bool:
    value_years = [int(year) for year in re.findall(r"\b(?:19|20)\d{2}\b", value)]
    period_years = [int(year) for year in re.findall(r"\b(?:19|20)\d{2}\b", period)]
    if not value_years or not period_years:
        return contains_normalized_fact(normalize_fact(period), normalize_fact(value))
    if len(period_years) >= 2:
        lower, upper = min(period_years), max(period_years)
        return all(lower <= year <= upper for year in value_years)
    if CURRENT_PERIOD_PATTERN.search(period):
        return all(period_years[0] <= year <= date.today().year for year in value_years)
    return all(year == period_years[0] for year in value_years)


def has_negation(text: str) -> bool:
    return NEGATION_PATTERN.search(text) is not None


def modality_signature(text: str) -> frozenset[str]:
    without_month_may = re.sub(r"\bmay\s+(?=(?:19|20)\d{2}\b)", "", text, flags=re.IGNORECASE)
    return frozenset(
        label for label, pattern in MODALITY_PATTERNS.items() if pattern.search(without_month_may)
    )


def is_factual_assertion(sentence: str) -> bool:
    tokens = significant_tokens(sentence)
    return bool(
        tokens & CLAIM_VERBS
        or FACTUAL_ASSERTION_PATTERN.search(sentence)
        or extract_dates(sentence)
        or extract_numbers(sentence)
        or extract_technologies(sentence)
        or extract_companies(sentence)
        or extract_titles(sentence)
    )


def validate_visual_output(
    template_content: bytes,
    rendered_content: bytes,
    *,
    allowed_removed_text: str | None = None,
) -> tuple[dict[str, Any], list[str]]:
    issues: list[str] = []
    source_links = Counter(extract_hyperlink_targets(template_content))
    rendered_links = Counter(extract_hyperlink_targets(rendered_content))
    missing_links = source_links - rendered_links
    added_links = rendered_links - source_links
    if missing_links or added_links:
        issues.append(
            "hyperlink count or targets changed "
            f"({sum(source_links.values())} source, {sum(rendered_links.values())} rendered)"
        )

    table_issues = detect_table_overflow(template_content, rendered_content)
    issues.extend(table_issues)
    cell_overflow_count = sum(
        "possible table overflow" in issue for issue in table_issues
    )
    source_geometry, rendered_geometry = render_and_inspect_documents(
        template_content,
        rendered_content,
    )
    geometry_report, geometry_issues = compare_rendered_geometry(
        source_geometry,
        rendered_geometry,
        expected_rendered_text=extract_docx_rendered_text(rendered_content),
        allowed_removed_text=allowed_removed_text,
        source_image_digests=extract_docx_image_digests(template_content),
        rendered_image_digests=extract_docx_image_digests(rendered_content),
    )
    issues.extend(geometry_issues)
    links_preserved = (
        not missing_links
        and not added_links
        and geometry_report["missingPdfLinkCount"] == 0
        and geometry_report["addedPdfLinkCount"] == 0
        and geometry_report["linkLocationChangedCount"] == 0
    )
    return (
        {
            "status": "passed" if not issues else "failed",
            **geometry_report,
            "linksPreserved": links_preserved,
            "sourceLinkCount": sum(source_links.values()),
            "renderedLinkCount": sum(rendered_links.values()),
            "missingLinkCount": sum(missing_links.values()),
            "addedLinkCount": sum(added_links.values()),
            "tableOverflow": bool(table_issues),
            "cellOverflowCount": cell_overflow_count,
            "tableStructureIssueCount": len(table_issues) - cell_overflow_count,
            "issues": issues,
        },
        issues,
    )


def render_and_count_pages(source: bytes, rendered: bytes) -> tuple[int, int]:
    source_geometry, rendered_geometry = render_and_inspect_documents(source, rendered)
    return source_geometry["pageCount"], rendered_geometry["pageCount"]


def render_and_inspect_documents(
    source: bytes,
    rendered: bytes,
) -> tuple[dict[str, Any], dict[str, Any]]:
    executable = shutil.which("soffice") or shutil.which("libreoffice")
    if not executable:
        raise DocumentValidationError(
            "Document validation failed: LibreOffice is required for page validation"
        )
    stable_tmp = "/private/tmp" if Path("/private/tmp").is_dir() else None
    with tempfile.TemporaryDirectory(
        prefix="tasko-docx-validation-",
        dir=stable_tmp,
    ) as directory:
        workdir = Path(directory)
        source_path = workdir / "source.docx"
        rendered_path = workdir / "rendered.docx"
        source_path.write_bytes(source)
        rendered_path.write_bytes(rendered)
        runtime_home = workdir / "home"
        runtime_home.mkdir()
        runtime_tmp = Path("/private/tmp")
        if not runtime_tmp.is_dir():
            runtime_tmp = workdir / "tmp"
            runtime_tmp.mkdir()
        environment = {
            **os.environ,
            "HOME": str(runtime_home),
            "TMPDIR": str(runtime_tmp),
            "TEMP": str(runtime_tmp),
            "TMP": str(runtime_tmp),
            "XDG_CONFIG_HOME": str(runtime_home / "xdg-config"),
            "XDG_CACHE_HOME": str(runtime_home / "xdg-cache"),
        }
        Path(environment["XDG_CONFIG_HOME"]).mkdir()
        Path(environment["XDG_CACHE_HOME"]).mkdir()
        for input_path in (source_path, rendered_path):
            profile = workdir / f"libreoffice-profile-{input_path.stem}"
            profile.mkdir()
            try:
                result = subprocess.run(
                    [
                        executable,
                        "--invisible",
                        "--headless",
                        "--norestore",
                        "--nofirststartwizard",
                        f"-env:UserInstallation={profile.as_uri()}",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        str(workdir),
                        str(input_path),
                    ],
                    capture_output=True,
                    text=True,
                    timeout=RENDER_TIMEOUT_SECONDS,
                    check=False,
                    env=environment,
                )
            except subprocess.TimeoutExpired as exc:
                raise DocumentValidationError(
                    "Document validation failed: DOCX rendering timed out"
                ) from exc
            output_path = workdir / f"{input_path.stem}.pdf"
            if result.returncode != 0 or not output_path.exists():
                detail = (result.stderr or result.stdout or "unknown conversion error").strip()
                raise DocumentValidationError(
                    f"Document validation failed: DOCX rendering failed ({detail[:240]})"
                )
        try:
            source_geometry = inspect_pdf_geometry(workdir / "source.pdf", workdir, "source")
            rendered_geometry = inspect_pdf_geometry(
                workdir / "rendered.pdf",
                workdir,
                "rendered",
            )
        except Exception as exc:
            if isinstance(exc, DocumentValidationError):
                raise
            raise DocumentValidationError(
                "Document validation failed: rendered PDF could not be inspected"
            ) from exc
        rasterizer = shutil.which("pdftoppm")
        if not rasterizer:
            raise DocumentValidationError(
                "Document validation failed: pdftoppm is required for visual validation"
            )
        try:
            raster_result = subprocess.run(
                [
                    rasterizer,
                    "-png",
                    "-r",
                    "96",
                    str(workdir / "rendered.pdf"),
                    str(workdir / "page"),
                ],
                capture_output=True,
                text=True,
                timeout=RENDER_TIMEOUT_SECONDS,
                check=False,
                env=environment,
            )
        except subprocess.TimeoutExpired as exc:
            raise DocumentValidationError(
                "Document validation failed: rendered page inspection timed out"
            ) from exc
        rendered_images = sorted(workdir.glob("page-*.png"))
        if (
            raster_result.returncode != 0
            or len(rendered_images) != rendered_geometry["pageCount"]
            or any(image.stat().st_size == 0 for image in rendered_images)
        ):
            detail = (
                raster_result.stderr or raster_result.stdout or "incomplete page rasterization"
            ).strip()
            raise DocumentValidationError(
                f"Document validation failed: rendered page inspection failed ({detail[:240]})"
            )
        return source_geometry, rendered_geometry


def inspect_pdf_geometry(
    pdf_path: Path,
    workdir: Path,
    label: str,
) -> dict[str, Any]:
    converter = shutil.which("pdftohtml")
    if not converter:
        raise DocumentValidationError(
            "Document validation failed: pdftohtml is required for geometry validation"
        )
    xml_path = workdir / f"{label}-geometry.xml"
    try:
        result = subprocess.run(
            [
                converter,
                "-q",
                "-xml",
                "-hidden",
                "-noroundcoord",
                "-zoom",
                "1.0",
                str(pdf_path),
                str(xml_path),
            ],
            capture_output=True,
            text=True,
            timeout=RENDER_TIMEOUT_SECONDS,
            check=False,
        )
    except subprocess.TimeoutExpired as exc:
        raise DocumentValidationError(
            "Document validation failed: PDF geometry inspection timed out"
        ) from exc
    if result.returncode != 0 or not xml_path.exists():
        detail = (result.stderr or result.stdout or "missing geometry XML").strip()
        raise DocumentValidationError(
            f"Document validation failed: PDF geometry inspection failed ({detail[:240]})"
        )

    root = etree.parse(str(xml_path)).getroot()
    page_sizes: list[dict[str, float]] = []
    text_boxes: list[dict[str, Any]] = []
    image_boxes: list[dict[str, Any]] = []
    for page_index, page in enumerate(root.findall(".//page"), start=1):
        page_top = xml_float(page.get("top"))
        page_sizes.append(
            {
                "width": xml_float(page.get("width")),
                "height": xml_float(page.get("height")),
            }
        )
        for text_node in page.findall(".//text"):
            text = "".join(text_node.itertext()).strip()
            if not text:
                continue
            text_boxes.append(
                geometry_box(text_node, page_index, page_top, text=text)
            )
        for image_node in page.findall(".//image"):
            image_path = xml_path.parent / str(image_node.get("src") or "")
            digest = ""
            if image_path.is_file():
                digest = hashlib.sha256(image_path.read_bytes()).hexdigest()
            image_boxes.append(
                geometry_box(image_node, page_index, page_top, digest=digest)
            )

    reader = PdfReader(pdf_path)
    link_boxes = extract_pdf_link_boxes(reader)
    return {
        "pageCount": len(reader.pages),
        "pageSizes": page_sizes,
        "textBoxes": text_boxes,
        "imageBoxes": image_boxes,
        "linkBoxes": link_boxes,
        "text": "\n".join(box["text"] for box in text_boxes),
    }


def xml_float(value: str | None) -> float:
    try:
        return float(value or 0)
    except ValueError:
        return 0.0


def geometry_box(
    node: Any,
    page: int,
    page_top: float,
    **extra: Any,
) -> dict[str, Any]:
    return {
        "page": page,
        "x": xml_float(node.get("left")),
        "y": xml_float(node.get("top")) - page_top,
        "width": xml_float(node.get("width")),
        "height": xml_float(node.get("height")),
        **extra,
    }


def extract_pdf_link_boxes(reader: PdfReader) -> list[dict[str, Any]]:
    boxes: list[dict[str, Any]] = []
    for page_index, page in enumerate(reader.pages, start=1):
        page_width = float(page.mediabox.width)
        page_height = float(page.mediabox.height)
        for annotation_reference in page.get("/Annots", []):
            annotation = annotation_reference.get_object()
            if str(annotation.get("/Subtype")) != "/Link":
                continue
            rectangle = annotation.get("/Rect")
            if not rectangle or len(rectangle) != 4:
                continue
            action = annotation.get("/A")
            action = action.get_object() if hasattr(action, "get_object") else action
            uri = action.get("/URI") if isinstance(action, dict) else None
            destination = annotation.get("/Dest")
            if uri:
                target = f"external:{uri}"
            elif destination is not None:
                target = f"internal:{destination}"
            else:
                target = "unknown"
            left, bottom, right, top = (float(value) for value in rectangle)
            boxes.append(
                {
                    "page": page_index,
                    "x": min(left, right),
                    "y": page_height - max(bottom, top),
                    "width": abs(right - left),
                    "height": abs(top - bottom),
                    "target": target,
                    "pageWidth": page_width,
                    "pageHeight": page_height,
                }
            )
    return boxes


def compare_rendered_geometry(
    source: dict[str, Any],
    rendered: dict[str, Any],
    *,
    expected_rendered_text: str,
    allowed_removed_text: str | None = None,
    source_image_digests: Counter[str],
    rendered_image_digests: Counter[str],
) -> tuple[dict[str, Any], list[str]]:
    issues: list[str] = []
    source_pages = int(source["pageCount"])
    rendered_pages = int(rendered["pageCount"])
    page_count_changed = source_pages != rendered_pages
    if page_count_changed:
        issues.append(f"page count changed from {source_pages} to {rendered_pages}")

    expected_tokens = Counter(geometry_text_tokens(expected_rendered_text))
    rendered_tokens = Counter(geometry_text_tokens(str(rendered.get("text") or "")))
    missing_text = expected_tokens - rendered_tokens
    missing_text_count = sum(missing_text.values())
    missing_text_samples = list(missing_text.elements())[:8]
    if missing_text_count:
        issues.append(
            f"{missing_text_count} rendered text tokens are missing from PDF"
            + (f": {', '.join(missing_text_samples)}" if missing_text_samples else "")
        )

    disappeared_source_text = Counter()
    if allowed_removed_text is not None:
        source_pdf_tokens = Counter(geometry_text_tokens(str(source.get("text") or "")))
        disappeared_source_text = (
            source_pdf_tokens - rendered_tokens
        ) - Counter(geometry_text_tokens(allowed_removed_text))
        disappeared_source_text_count = sum(disappeared_source_text.values())
        if disappeared_source_text_count:
            issues.append(
                f"{disappeared_source_text_count} source text tokens disappeared unexpectedly"
            )
    else:
        disappeared_source_text_count = 0

    text_geometry_changed = count_changed_box_geometry(
        source,
        rendered,
        source.get("textBoxes", []),
        rendered.get("textBoxes", []),
        key="text",
    )
    if text_geometry_changed:
        issues.append(
            f"{text_geometry_changed} unchanged text boxes moved or resized significantly"
        )

    source_text_outside = count_boxes_outside_pages(source, source.get("textBoxes", []))
    rendered_text_outside = count_boxes_outside_pages(
        rendered,
        rendered.get("textBoxes", []),
    )
    if rendered_text_outside:
        issues.append(
            f"{rendered_text_outside} text boxes extend outside rendered page bounds"
        )

    missing_source_images = source_image_digests - rendered_image_digests
    missing_source_image_count = sum(missing_source_images.values())
    if missing_source_image_count:
        issues.append(f"{missing_source_image_count} source images are missing")

    source_pdf_images = Counter(
        box["digest"] for box in source.get("imageBoxes", []) if box.get("digest")
    )
    rendered_pdf_images = Counter(
        box["digest"] for box in rendered.get("imageBoxes", []) if box.get("digest")
    )
    missing_pdf_image_count = sum((source_pdf_images - rendered_pdf_images).values())
    if not source_pdf_images and not rendered_pdf_images:
        missing_pdf_image_count = max(
            0,
            len(source.get("imageBoxes", [])) - len(rendered.get("imageBoxes", [])),
        )
    if missing_pdf_image_count:
        issues.append(f"{missing_pdf_image_count} source image boxes are missing from PDF")

    image_geometry_changed = count_changed_box_geometry(
        source,
        rendered,
        source.get("imageBoxes", []),
        rendered.get("imageBoxes", []),
        key="digest",
    )
    if image_geometry_changed:
        issues.append(
            f"{image_geometry_changed} image boxes moved or resized significantly"
        )
    source_images_outside = count_boxes_outside_pages(source, source.get("imageBoxes", []))
    rendered_images_outside = count_boxes_outside_pages(
        rendered,
        rendered.get("imageBoxes", []),
    )
    if rendered_images_outside:
        issues.append(
            f"{rendered_images_outside} image boxes extend outside rendered page bounds"
        )

    source_pdf_links = Counter(box.get("target") for box in source.get("linkBoxes", []))
    rendered_pdf_links = Counter(box.get("target") for box in rendered.get("linkBoxes", []))
    missing_pdf_links = source_pdf_links - rendered_pdf_links
    added_pdf_links = rendered_pdf_links - source_pdf_links
    missing_pdf_link_count = sum(missing_pdf_links.values())
    added_pdf_link_count = sum(added_pdf_links.values())
    if missing_pdf_link_count or added_pdf_link_count:
        issues.append(
            "PDF hyperlink count or targets changed "
            f"({sum(source_pdf_links.values())} source, "
            f"{sum(rendered_pdf_links.values())} rendered)"
        )
    link_location_changed = count_changed_box_geometry(
        source,
        rendered,
        source.get("linkBoxes", []),
        rendered.get("linkBoxes", []),
        key="target",
    )
    if link_location_changed:
        issues.append(
            f"{link_location_changed} hyperlink bounding boxes moved significantly"
        )

    return (
        {
            "sourcePageCount": source_pages,
            "renderedPageCount": rendered_pages,
            "pageCountChanged": page_count_changed,
            "sourceTextBoxCount": len(source.get("textBoxes", [])),
            "renderedTextBoxCount": len(rendered.get("textBoxes", [])),
            "missingTextCount": missing_text_count,
            "missingTextSamples": missing_text_samples,
            "disappearedSourceTextCount": disappeared_source_text_count,
            "disappearedSourceTextSamples": list(disappeared_source_text.elements())[:8],
            "textGeometryChangedCount": text_geometry_changed,
            "sourceTextOutsidePageCount": source_text_outside,
            "textOutsidePageCount": rendered_text_outside,
            "sourceImageCount": sum(source_image_digests.values()),
            "renderedImageCount": sum(rendered_image_digests.values()),
            "sourceImageBoxCount": len(source.get("imageBoxes", [])),
            "renderedImageBoxCount": len(rendered.get("imageBoxes", [])),
            "missingSourceImageCount": missing_source_image_count,
            "missingPdfImageCount": missing_pdf_image_count,
            "imageGeometryChangedCount": image_geometry_changed,
            "sourceImageOutsidePageCount": source_images_outside,
            "imageOutsidePageCount": rendered_images_outside,
            "sourcePdfLinkCount": len(source.get("linkBoxes", [])),
            "renderedPdfLinkCount": len(rendered.get("linkBoxes", [])),
            "missingPdfLinkCount": missing_pdf_link_count,
            "addedPdfLinkCount": added_pdf_link_count,
            "linkLocationChangedCount": link_location_changed,
        },
        issues,
    )


def geometry_text_tokens(text: str) -> list[str]:
    return re.findall(r"[^\W_]+", text.casefold())


def count_changed_box_geometry(
    source_document: dict[str, Any],
    rendered_document: dict[str, Any],
    source_boxes: list[dict[str, Any]],
    rendered_boxes: list[dict[str, Any]],
    *,
    key: str,
) -> int:
    source_by_key: dict[str, list[dict[str, Any]]] = {}
    rendered_by_key: dict[str, list[dict[str, Any]]] = {}
    for box in source_boxes:
        value = normalize_geometry_key(box.get(key))
        if value:
            source_by_key.setdefault(value, []).append(box)
    for box in rendered_boxes:
        value = normalize_geometry_key(box.get(key))
        if value:
            rendered_by_key.setdefault(value, []).append(box)

    changed = 0
    for value in source_by_key.keys() & rendered_by_key.keys():
        source_matches = source_by_key[value]
        rendered_matches = rendered_by_key[value]
        for source_box, rendered_box in zip(source_matches, rendered_matches):
            if box_geometry_changed(
                source_document,
                rendered_document,
                source_box,
                rendered_box,
            ):
                changed += 1
    return changed


def normalize_geometry_key(value: Any) -> str:
    if not isinstance(value, str):
        return ""
    return re.sub(r"\s+", " ", value.casefold()).strip()


def box_geometry_changed(
    source_document: dict[str, Any],
    rendered_document: dict[str, Any],
    source_box: dict[str, Any],
    rendered_box: dict[str, Any],
) -> bool:
    if source_box["page"] != rendered_box["page"]:
        return True
    source_size = page_size(source_document, int(source_box["page"]))
    rendered_size = page_size(rendered_document, int(rendered_box["page"]))
    if not all((*source_size, *rendered_size)):
        return True
    source_center = (
        (source_box["x"] + source_box["width"] / 2) / source_size[0],
        (source_box["y"] + source_box["height"] / 2) / source_size[1],
    )
    rendered_center = (
        (rendered_box["x"] + rendered_box["width"] / 2) / rendered_size[0],
        (rendered_box["y"] + rendered_box["height"] / 2) / rendered_size[1],
    )
    moved = any(
        abs(source_value - rendered_value) > GEOMETRY_POSITION_TOLERANCE
        for source_value, rendered_value in zip(source_center, rendered_center, strict=True)
    )
    resized = any(
        relative_change(float(source_box[dimension]), float(rendered_box[dimension]))
        > GEOMETRY_SIZE_TOLERANCE
        for dimension in ("width", "height")
    )
    return moved or resized


def relative_change(source: float, rendered: float) -> float:
    if source == rendered:
        return 0.0
    return abs(rendered - source) / max(abs(source), 1.0)


def count_boxes_outside_pages(
    document: dict[str, Any],
    boxes: list[dict[str, Any]],
) -> int:
    outside = 0
    for box in boxes:
        width, height = page_size(document, int(box["page"]))
        if (
            box["x"] < -PAGE_BOUNDARY_TOLERANCE
            or box["y"] < -PAGE_BOUNDARY_TOLERANCE
            or box["x"] + box["width"] > width + PAGE_BOUNDARY_TOLERANCE
            or box["y"] + box["height"] > height + PAGE_BOUNDARY_TOLERANCE
        ):
            outside += 1
    return outside


def page_size(document: dict[str, Any], page: int) -> tuple[float, float]:
    sizes = document.get("pageSizes", [])
    if not 1 <= page <= len(sizes):
        return 0.0, 0.0
    size = sizes[page - 1]
    return float(size.get("width") or 0), float(size.get("height") or 0)


def extract_docx_image_digests(content: bytes) -> Counter[str]:
    with zipfile.ZipFile(BytesIO(content)) as archive:
        return Counter(
            hashlib.sha256(archive.read(name)).hexdigest()
            for name in archive.namelist()
            if name.startswith("word/media/") and not name.endswith("/")
        )


def detect_table_overflow(source: bytes, rendered: bytes) -> list[str]:
    source_document = Document(BytesIO(source))
    rendered_document = Document(BytesIO(rendered))
    issues: list[str] = []
    if len(source_document.tables) != len(rendered_document.tables):
        return ["table count changed"]
    for table_index, (source_table, rendered_table) in enumerate(
        zip(source_document.tables, rendered_document.tables, strict=True),
        start=1,
    ):
        rendered_width = table_width_twips(rendered_table)
        available_width = max(
            section_available_width_twips(section)
            for section in rendered_document.sections
        )
        if rendered_width and rendered_width > available_width + 20:
            issues.append(f"table {table_index} exceeds the available page width")
        if len(source_table.rows) != len(rendered_table.rows):
            issues.append(f"row count changed in table {table_index}")
            continue
        for row_index, (source_row, rendered_row) in enumerate(
            zip(source_table.rows, rendered_table.rows, strict=True),
            start=1,
        ):
            if len(source_row.cells) != len(rendered_row.cells):
                issues.append(f"cell count changed in table {table_index}, row {row_index}")
                continue
            for cell_index, (source_cell, rendered_cell) in enumerate(
                zip(source_row.cells, rendered_row.cells, strict=True),
                start=1,
            ):
                exact_height = rendered_row.height_rule == WD_ROW_HEIGHT_RULE.EXACTLY
                content_expanded_excessively = len(rendered_cell.text.strip()) > max(
                    180,
                    len(source_cell.text.strip()) * 3,
                )
                fixed_row_would_clip = exact_height and row_text_exceeds_exact_height(
                    source_cell.text,
                    rendered_cell.text,
                    rendered_cell,
                    rendered_row,
                )
                if content_expanded_excessively or fixed_row_would_clip:
                    issues.append(
                        f"possible table overflow at table {table_index}, row {row_index}, "
                        f"cell {cell_index}"
                    )
    return issues


def table_width_twips(table: Any) -> int:
    grid = table._tbl.tblGrid
    if grid is None:
        return 0
    return sum(
        int(column.get(qn("w:w"), "0"))
        for column in grid.gridCol_lst
    )


def section_available_width_twips(section: Any) -> int:
    width = int(section.page_width or 0)
    left = int(section.left_margin or 0)
    right = int(section.right_margin or 0)
    return max(0, (width - left - right) // 635)


def row_text_exceeds_exact_height(
    source_text: str,
    rendered_text: str,
    rendered_cell: Any,
    rendered_row: Any,
) -> bool:
    height = int(rendered_row.height or 0) // 635
    width_element = rendered_cell._tc.tcPr.tcW
    width = int(width_element.get(qn("w:w"), "0")) if width_element is not None else 0
    if height <= 0 or width <= 0:
        return len(rendered_text.strip()) > max(80, len(source_text.strip()) * 1.4)
    characters_per_line = max(8, width // 110)
    line_capacity = max(1, height // 240)
    source_lines = estimated_wrapped_lines(source_text, characters_per_line)
    rendered_lines = estimated_wrapped_lines(rendered_text, characters_per_line)
    return rendered_lines > max(source_lines, line_capacity)


def estimated_wrapped_lines(text: str, characters_per_line: int) -> int:
    lines = text.splitlines() or [""]
    return sum(
        max(1, (len(line) + characters_per_line - 1) // characters_per_line)
        for line in lines
    )


def extract_hyperlink_targets(content: bytes) -> list[str]:
    targets: list[str] = []
    with zipfile.ZipFile(BytesIO(content)) as archive:
        archive_names = archive.namelist()
        for name in archive_names:
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            part_root = etree.fromstring(archive.read(name))
            directory, file_name = name.rsplit("/", 1)
            relationships_name = f"{directory}/_rels/{file_name}.rels"
            relationships: dict[str, str] = {}
            if relationships_name in archive_names:
                relationships_root = etree.fromstring(archive.read(relationships_name))
                relationships = {
                    relationship.get("Id", ""): relationship.get("Target", "")
                    for relationship in relationships_root
                    if relationship.get("Type", "").endswith("/hyperlink")
                }
            for hyperlink in part_root.iter(qn("w:hyperlink")):
                relationship_id = hyperlink.get(qn("r:id"), "")
                anchor = hyperlink.get(qn("w:anchor"), "")
                if relationship_id in relationships:
                    targets.append(f"external:{relationships[relationship_id]}")
                elif anchor:
                    targets.append(f"internal:{anchor}")
            for instruction in part_root.iter(qn("w:instrText")):
                field_match = re.search(
                    r'\bHYPERLINK\s+(?:"([^"]+)"|(\S+))',
                    instruction.text or "",
                    flags=re.IGNORECASE,
                )
                if field_match:
                    targets.append(f"field:{field_match.group(1) or field_match.group(2)}")
    return sorted(targets)


def extract_docx_text(content: bytes) -> str:
    return "\n".join(extract_docx_paragraphs(content))


def extract_docx_rendered_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    texts = extract_container_text(document)
    seen_parts: set[str] = set()
    for section in document.sections:
        for container in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            part_name = str(container.part.partname)
            if part_name in seen_parts:
                continue
            seen_parts.add(part_name)
            texts.extend(extract_container_text(container))
    return "\n".join(texts)


def extract_container_text(container: Any) -> list[str]:
    texts = [paragraph.text.strip() for paragraph in container.paragraphs]
    for table in container.tables:
        for row in table.rows:
            texts.extend(cell.text.strip() for cell in row.cells)
    return [text for text in texts if text]


def extract_docx_paragraphs(content: bytes) -> list[str]:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            paragraphs.extend(cell.text.strip() for cell in row.cells)
    return [paragraph for paragraph in paragraphs if paragraph]


def flatten_evidence(value: Any) -> list[str]:
    if isinstance(value, str):
        if value.startswith("data:"):
            return []
        return [value]
    if isinstance(value, dict):
        return [text for nested in value.values() for text in flatten_evidence(nested)]
    if isinstance(value, list):
        return [text for nested in value for text in flatten_evidence(nested)]
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        return [str(value)]
    return []


def extract_dates(text: str) -> list[str]:
    return [match.group(0) for match in DATE_PATTERN.finditer(text)]


def extract_numbers(text: str) -> list[str]:
    return [match.group(0) for match in NUMBER_PATTERN.finditer(text)]


def extract_technologies(text: str) -> list[str]:
    normalized = text.casefold()
    return [
        technology
        for technology in TECHNOLOGIES
        if re.search(rf"(?<![\w]){re.escape(technology)}(?![\w])", normalized)
    ]


def extract_companies(text: str) -> list[str]:
    return [match.group(1) for match in COMPANY_PATTERN.finditer(text)]


def extract_titles(text: str) -> list[str]:
    return [match.group(1) or match.group(2) for match in TITLE_PATTERN.finditer(text)]


def split_sentences(text: str) -> list[str]:
    return [
        sentence.strip()
        for sentence in re.split(r"(?<=[.!?])\s+|\n+", text)
        if sentence.strip()
    ]


def significant_tokens(text: str) -> set[str]:
    return {
        cleaned
        for token in re.findall(r"[^\W\d_][\w+#.-]{2,}", text.casefold())
        if (cleaned := token.strip(".-")) and cleaned not in STOP_WORDS
    }


def semantic_tokens(text: str) -> set[str]:
    tokens: set[str] = set()
    for token in significant_tokens(text):
        if len(token) > 5 and token.endswith("ing"):
            token = token[:-3]
        elif len(token) > 4 and token.endswith("ed"):
            token = token[:-2]
        tokens.add(token)
    return tokens


def normalize_fact(value: str) -> str:
    return re.sub(r"\s+", " ", value.casefold().replace(",", ".")).strip()


def contains_normalized_fact(haystack: str, needle: str) -> bool:
    if not needle:
        return False
    return re.search(
        rf"(?<![\w]){re.escape(needle)}(?![\w])",
        haystack,
        flags=re.IGNORECASE,
    ) is not None
