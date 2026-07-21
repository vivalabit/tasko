import base64
import json
import zipfile
from io import BytesIO

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from app.services.cover_letter_blocks import extract_cover_letter_blocks_from_docx
from app.services.document_export import (
    build_document_from_template,
    ensure_cover_letter_date_replacement,
)
from app.services.document_validation import build_document_diff


def add_hyperlink(paragraph, text: str, url: str) -> str:
    relationship_id = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    properties.append(color)
    node = OxmlElement("w:t")
    node.text = text
    run.extend((properties, node))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return relationship_id


def cover_letter_template() -> tuple[bytes, str]:
    document = Document()
    document.sections[0].header.paragraphs[0].text = "EDUARD · ENGINEER"
    document.add_paragraph("Eduard · Zurich · eduard@example.com")
    document.add_paragraph("Dear Hiring Team,")
    body = document.add_paragraph()
    body.style = document.styles["Normal"]
    body.add_run("Original body").bold = True
    relationship_id = add_hyperlink(body, " portfolio", "https://example.com/work")
    controls = body.add_run()
    controls.add_tab()
    controls.add_text("Remote")
    controls.add_break()
    controls.add_text("Delivery")
    document.add_paragraph("Kind regards,")
    document.add_paragraph("Eduard")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Protected reference"
    document.sections[0].footer.paragraphs[0].text = "Private application"
    output = BytesIO()
    document.save(output)
    return output.getvalue(), relationship_id


def document_bytes(document: Document) -> bytes:
    output = BytesIO()
    document.save(output)
    return output.getvalue()


@pytest.mark.parametrize(
    ("greeting", "closing"),
    [
        ("Sehr geehrte Damen und Herren,", "Mit freundlichen Grüßen,"),
        ("Liebes pom+ Team,", "Freundliche Grüsse"),
        ("Bonjour Madame, Monsieur,", "Cordialement,"),
        ("Gentile Responsabile delle assunzioni,", "Cordiali saluti,"),
        ("Estimada responsable de selección,", "Saludos cordiales,"),
        ("Geachte heer/mevrouw,", "Met vriendelijke groet,"),
        ("Szanowni Państwo,", "Z poważaniem,"),
        ("Bästa rekryteringsteam,", "Med vänliga hälsningar,"),
    ],
)
def test_multilingual_greetings_are_editable_while_closings_stay_protected(
    greeting: str,
    closing: str,
) -> None:
    document = Document()
    document.add_paragraph(greeting)
    document.add_paragraph("Reusable body paragraph.")
    document.add_paragraph(closing)
    document.add_paragraph("Eduard")

    paragraphs = extract_cover_letter_blocks_from_docx(document_bytes(document))

    assert [(paragraph["type"], paragraph["editable"]) for paragraph in paragraphs] == [
        ("greeting", True),
        ("body", True),
        ("closing", False),
        ("signature", False),
    ]


def test_signature_drawing_before_closing_text_keeps_body_editable() -> None:
    document = Document()
    document.add_paragraph("Eduard Ishchenko")
    document.add_paragraph("Sehr geehrter Herr Duss")
    document.add_paragraph("Reusable body paragraph.")
    closing = document.add_paragraph()
    closing.add_run()._r.append(OxmlElement("w:drawing"))
    closing.add_run("Freundliche Grüsse")
    document.add_paragraph("Eduard Ishchenko")

    paragraphs = extract_cover_letter_blocks_from_docx(document_bytes(document))

    assert [(paragraph["type"], paragraph["editable"]) for paragraph in paragraphs] == [
        ("protected", False),
        ("greeting", True),
        ("body", True),
        ("closing", False),
        ("signature", False),
    ]


def test_signature_drawing_without_closing_text_ends_editable_body() -> None:
    document = Document()
    document.add_paragraph("Motivationsschreiben Application Developer")
    document.add_paragraph("Sehr geehrte Damen und Herren")
    document.add_paragraph("First reusable body paragraph.")
    document.add_paragraph("Second reusable body paragraph.")
    signature = document.add_paragraph()
    signature.add_run()._r.append(OxmlElement("w:drawing"))
    document.add_paragraph("Eduard Ishchenko")

    paragraphs = extract_cover_letter_blocks_from_docx(document_bytes(document))

    assert [(paragraph["type"], paragraph["editable"]) for paragraph in paragraphs] == [
        ("subject", True),
        ("greeting", True),
        ("body", True),
        ("body", True),
        ("closing", False),
        ("signature", False),
    ]


def test_cover_letter_renderer_updates_subject_and_greeting() -> None:
    document = Document()
    document.add_paragraph("Motivationsschreiben Application Developer")
    document.add_paragraph("Sehr geehrte Damen und Herren,")
    document.add_paragraph("Reusable body paragraph.")
    document.add_paragraph("Freundliche Grüsse")
    document.add_paragraph("Eduard Ishchenko")
    template = document_bytes(document)
    content = json.dumps(
        {
            "replacements": [
                {
                    "paragraphId": "paragraph-0001",
                    "spanId": "paragraph-0001-span-0001",
                    "original": "Motivationsschreiben Application Developer",
                    "replacement": "Bewerbung als Platform Engineer bei Acme",
                    "reason": "Uses the current vacancy title",
                    "evidenceIds": [
                        "source:paragraph-0001-span-0001",
                        "vacancy:title",
                        "vacancy:company",
                    ],
                },
                {
                    "paragraphId": "paragraph-0002",
                    "spanId": "paragraph-0002-span-0001",
                    "original": "Sehr geehrte Damen und Herren,",
                    "replacement": "Liebes Acme Recruiting-Team,",
                    "reason": "Uses the team fallback greeting",
                    "evidenceIds": [
                        "source:paragraph-0002-span-0001",
                        "vacancy:company",
                    ],
                },
            ]
        }
    )

    rendered = build_document_from_template(
        template_content=template,
        content=content,
        document_type="cover_letter",
    )

    assert [paragraph.text for paragraph in Document(BytesIO(rendered)).paragraphs[:2]] == [
        "Bewerbung als Platform Engineer bei Acme",
        "Liebes Acme Recruiting-Team,",
    ]


def test_cover_letter_renderer_updates_only_date_inside_contact_header() -> None:
    document = Document()
    header = document.add_paragraph()
    header.add_run("Eduard Ishchenko")
    header.add_run().add_break()
    header.add_run("Männedorf, 15. Juli")
    header.add_run().add_break()
    header.add_run("+41 79 516 84 31")
    header.add_run().add_break()
    header.add_run("eduard@example.com")
    document.add_paragraph("Sehr geehrte Damen und Herren,")
    document.add_paragraph("Reusable body paragraph.")
    document.add_paragraph("Freundliche Grüsse")
    document.add_paragraph("Eduard Ishchenko")
    template = document_bytes(document)

    paragraphs = extract_cover_letter_blocks_from_docx(template)
    header_spans = paragraphs[0]["spans"]
    editable_spans = [span for span in header_spans if span["editable"]]

    assert paragraphs[0]["type"] == "protected"
    assert [(span["spanId"], span["original"]) for span in editable_spans] == [
        ("paragraph-0001-span-0003", "Männedorf, 15. Juli")
    ]

    content = json.dumps(
        {
            "replacements": [
                {
                    "paragraphId": "paragraph-0001",
                    "spanId": "paragraph-0001-span-0003",
                    "original": "Männedorf, 15. Juli",
                    "replacement": "Männedorf, 21. Juli 2026",
                    "reason": "Updates the letter date",
                    "evidenceIds": [
                        "source:paragraph-0001-span-0003",
                        "generation:date",
                    ],
                }
            ]
        }
    )
    rendered = build_document_from_template(
        template_content=template,
        content=content,
        document_type="cover_letter",
    )

    assert Document(BytesIO(rendered)).paragraphs[0].text == (
        "Eduard Ishchenko\nMännedorf, 21. Juli 2026\n"
        "+41 79 516 84 31\neduard@example.com"
    )


def test_cover_letter_generation_injects_missing_current_date_replacement() -> None:
    document = Document()
    header = document.add_paragraph()
    header.add_run("Männedorf, 15. Juli")
    document.add_paragraph("Dear Hiring Team,")
    document.add_paragraph("Reusable body paragraph.")
    document.add_paragraph("Kind regards,")
    document.add_paragraph("Eduard Ishchenko")
    template = document_bytes(document)
    model_content = json.dumps(
        {
            "replacements": [
                {
                    "paragraphId": "paragraph-0003",
                    "spanId": "paragraph-0003-span-0001",
                    "original": "Reusable body paragraph.",
                    "replacement": "Motivation-led body paragraph.",
                    "reason": "Tailors the motivation",
                    "evidenceIds": ["source:paragraph-0003-span-0001"],
                }
            ]
        }
    )

    enriched = ensure_cover_letter_date_replacement(
        template_content=template,
        content=model_content,
        generation_date="2026-07-21",
        language="German",
    )
    replacements = json.loads(enriched)["replacements"]

    assert replacements[-1] == {
        "paragraphId": "paragraph-0001",
        "spanId": "paragraph-0001-span-0001",
        "original": "Männedorf, 15. Juli",
        "replacement": "Männedorf, 21. Juli 2026",
        "reason": "Updates the letter date to the generation date",
        "evidenceIds": ["source:paragraph-0001-span-0001", "generation:date"],
    }


def test_cover_letter_renderer_preserves_distinct_paragraph_and_run_styles() -> None:
    document = Document()
    greeting = document.add_paragraph("Sehr geehrte Damen und Herren,")
    greeting.style = document.styles["Subtitle"]
    greeting.alignment = WD_ALIGN_PARAGRAPH.CENTER
    greeting.runs[0].italic = True
    first_body = document.add_paragraph(style="Quote")
    first_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    first_body.add_run("First reusable paragraph.").bold = True
    second_body = document.add_paragraph(style="Intense Quote")
    second_body.alignment = WD_ALIGN_PARAGRAPH.LEFT
    second_body.add_run("Second reusable paragraph.").italic = True
    closing = document.add_paragraph("Mit freundlichen Grüßen,", style="Caption")
    closing.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature = document.add_paragraph("Eduard", style="Title")
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    template = document_bytes(document)
    content = json.dumps(
        {
            "replacements": [
                {
                    "paragraphId": "paragraph-0002",
                    "spanId": "paragraph-0002-span-0001",
                    "original": "First reusable paragraph.",
                    "replacement": "First targeted paragraph.",
                    "reason": "Tailors the first body paragraph",
                    "evidenceIds": [
                        "profile:experience:experience-acme:achievement-a1b2c3d4e5"
                    ],
                },
                {
                    "paragraphId": "paragraph-0003",
                    "spanId": "paragraph-0003-span-0001",
                    "original": "Second reusable paragraph.",
                    "replacement": "Second targeted paragraph.",
                    "reason": "Tailors the second body paragraph",
                    "evidenceIds": ["profile:skills"],
                },
            ]
        }
    )

    source_blocks = extract_cover_letter_blocks_from_docx(template)
    rendered = build_document_from_template(
        template_content=template,
        content=content,
        document_type="cover_letter",
    )
    rendered_document = Document(BytesIO(rendered))

    assert [paragraph["style"] for paragraph in source_blocks] == [
        {"paragraphStyle": "Subtitle", "alignment": "center"},
        {"paragraphStyle": "Quote", "alignment": "both"},
        {"paragraphStyle": "IntenseQuote", "alignment": "left"},
        {"paragraphStyle": "Caption", "alignment": "right"},
        {"paragraphStyle": "Title", "alignment": "right"},
    ]
    assert [paragraph.style.style_id for paragraph in rendered_document.paragraphs] == [
        "Subtitle",
        "Quote",
        "IntenseQuote",
        "Caption",
        "Title",
    ]
    assert [paragraph.alignment for paragraph in rendered_document.paragraphs] == [
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.RIGHT,
        WD_ALIGN_PARAGRAPH.RIGHT,
    ]
    assert [paragraph.text for paragraph in rendered_document.paragraphs] == [
        "Sehr geehrte Damen und Herren,",
        "First targeted paragraph.",
        "Second targeted paragraph.",
        "Mit freundlichen Grüßen,",
        "Eduard",
    ]
    assert rendered_document.paragraphs[0].runs[0].italic is True
    assert rendered_document.paragraphs[1].runs[0].bold is True
    assert rendered_document.paragraphs[2].runs[0].italic is True

    with zipfile.ZipFile(BytesIO(template)) as source, zipfile.ZipFile(
        BytesIO(rendered)
    ) as result:
        source_root = etree.fromstring(source.read("word/document.xml"))
        rendered_root = etree.fromstring(result.read("word/document.xml"))
    source_properties = [
        etree.tostring(paragraph.find(qn("w:pPr")))
        for paragraph in source_root.findall(".//" + qn("w:body") + "/" + qn("w:p"))
    ]
    rendered_properties = [
        etree.tostring(paragraph.find(qn("w:pPr")))
        for paragraph in rendered_root.findall(".//" + qn("w:body") + "/" + qn("w:p"))
    ]
    assert rendered_properties == source_properties


def test_cover_letter_preserves_immutable_drawing_symbol_and_supported_field() -> None:
    document = Document()
    document.add_paragraph("Dear Hiring Team,")
    paragraph = document.add_paragraph()
    paragraph.add_run("Original ")
    paragraph.add_run().add_picture(
        BytesIO(
            base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/"
                "x8AAusB9Y9ZJYQAAAAASUVORK5CYII="
            )
        )
    )
    symbol = OxmlElement("w:sym")
    symbol.set(qn("w:font"), "Wingdings")
    symbol.set(qn("w:char"), "F0B7")
    paragraph.add_run()._r.append(symbol)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), " PAGE ")
    field_run = OxmlElement("w:r")
    field_result = OxmlElement("w:t")
    field_result.text = "1"
    field_run.append(field_result)
    field.append(field_run)
    paragraph._p.append(field)
    paragraph.add_run("body backed by evidence.")
    document.add_paragraph("Kind regards,")
    template = document_bytes(document)

    paragraphs = extract_cover_letter_blocks_from_docx(template)
    body = paragraphs[1]
    assert [(span["type"], span["editable"]) for span in body["spans"]] == [
        ("text", True),
        ("drawing", False),
        ("symbol", False),
        ("field", False),
        ("text", True),
    ]
    assert [item["type"] for item in body["protectedElements"]] == [
        "drawing",
        "symbol",
        "field",
    ]

    rendered = build_document_from_template(
        template_content=template,
        content=json.dumps(
            {
                "replacements": [
                    {
                        "paragraphId": "paragraph-0002",
                        "spanId": "paragraph-0002-span-0001",
                        "original": "Original ",
                        "replacement": "Targeted ",
                        "reason": "Tailors only editable text",
                        "evidenceIds": ["source:paragraph-0002-span-0001"],
                    }
                ]
            }
        ),
        document_type="cover_letter",
    )

    with zipfile.ZipFile(BytesIO(template)) as source_archive, zipfile.ZipFile(
        BytesIO(rendered)
    ) as rendered_archive:
        source_root = etree.fromstring(source_archive.read("word/document.xml"))
        rendered_root = etree.fromstring(rendered_archive.read("word/document.xml"))
        for tag in (qn("w:drawing"), qn("w:sym"), qn("w:fldSimple")):
            assert etree.tostring(rendered_root.find(".//" + tag)) == etree.tostring(
                source_root.find(".//" + tag)
            )


def test_cover_letter_blocks_model_protected_paragraphs_and_editable_spans() -> None:
    template, _ = cover_letter_template()

    paragraphs = extract_cover_letter_blocks_from_docx(template)
    body = paragraphs[2]

    assert [(paragraph["type"], paragraph["editable"]) for paragraph in paragraphs] == [
        ("protected", False),
        ("greeting", True),
        ("body", True),
        ("closing", False),
        ("signature", False),
        ("tableCell", False),
    ]
    assert body["paragraphId"] == "paragraph-0003"
    assert body["original"] == "Original body portfolio\tRemote\nDelivery"
    assert body["style"]["paragraphStyle"] == "Normal"
    assert [(span["type"], span["editable"]) for span in body["spans"]] == [
        ("text", True),
        ("hyperlink", False),
        ("tab", False),
        ("text", True),
        ("lineBreak", False),
        ("text", True),
    ]
    assert body["spans"][0]["style"]["bold"] is True
    assert body["spans"][0]["evidenceId"] == "source:paragraph-0003-span-0001"
    assert body["hyperlinks"] == [
        {
            "hyperlinkId": "paragraph-0003-hyperlink-0001",
            "target": "https://example.com/work",
            "anchor": "",
        }
    ]
    assert [item["type"] for item in body["protectedElements"]] == [
        "hyperlink",
        "tab",
        "lineBreak",
    ]
    assert all(
        span["editable"] is False
        for paragraph in paragraphs
        if paragraph["type"] not in {"subject", "greeting", "body"}
        for span in paragraph["spans"]
    )


def test_cover_letter_renderer_updates_only_editable_spans_and_preserves_package() -> None:
    template, relationship_id = cover_letter_template()
    content = json.dumps(
        {
            "replacements": [
                {
                    "paragraphId": "paragraph-0003",
                    "spanId": "paragraph-0003-span-0001",
                    "original": "Original body",
                    "replacement": "Targeted introduction",
                    "reason": "Uses verified experience",
                    "evidenceIds": ["source:paragraph-0003-span-0001"],
                },
                {
                    "paragraphId": "paragraph-0003",
                    "spanId": "paragraph-0003-span-0004",
                    "original": "Remote",
                    "replacement": "Hybrid",
                    "reason": "Uses the confirmed work arrangement",
                    "evidenceIds": ["confirmation:work-arrangement"],
                },
            ]
        }
    )

    rendered = build_document_from_template(
        template_content=template,
        content=content,
        document_type="cover_letter",
    )
    rendered_document = Document(BytesIO(rendered))
    body = rendered_document.paragraphs[2]
    diff = build_document_diff(
        template,
        content,
        "cover_letter",
        rendered_content=rendered,
    )

    assert body.text == "Targeted introduction portfolio\tHybrid\nDelivery"
    assert rendered_document.paragraphs[0].text == "Eduard · Zurich · eduard@example.com"
    assert rendered_document.paragraphs[1].text == "Dear Hiring Team,"
    assert rendered_document.paragraphs[3].text == "Kind regards,"
    assert rendered_document.paragraphs[4].text == "Eduard"
    assert rendered_document.tables[0].cell(0, 0).text == "Protected reference"
    assert rendered_document.part.rels[relationship_id].target_ref == "https://example.com/work"
    assert diff == [
        {
            "blockId": "paragraph-0003",
            "paragraphId": "paragraph-0003",
            "spanId": "paragraph-0003-span-0001",
            "type": "body",
            "original": "Original body",
            "replacement": "Targeted introduction",
            "reason": "Uses verified experience",
            "evidenceIds": ["source:paragraph-0003-span-0001"],
        },
        {
            "blockId": "paragraph-0003",
            "paragraphId": "paragraph-0003",
            "spanId": "paragraph-0003-span-0004",
            "type": "body",
            "original": "Remote",
            "replacement": "Hybrid",
            "reason": "Uses the confirmed work arrangement",
            "evidenceIds": ["confirmation:work-arrangement"],
        },
    ]

    with zipfile.ZipFile(BytesIO(template)) as source, zipfile.ZipFile(BytesIO(rendered)) as result:
        assert source.read("word/styles.xml") == result.read("word/styles.xml")
        assert source.read("word/_rels/document.xml.rels") == result.read(
            "word/_rels/document.xml.rels"
        )
    root = etree.fromstring(zipfile.ZipFile(BytesIO(rendered)).read("word/document.xml"))
    body_paragraph = root.findall(".//" + qn("w:p"))[2]
    assert body_paragraph.find(".//" + qn("w:rPr") + "/" + qn("w:b")) is not None
    assert body_paragraph.find(".//" + qn("w:hyperlink")) is not None
    assert len(body_paragraph.findall(".//" + qn("w:tab"))) == 1
    assert len(body_paragraph.findall(".//" + qn("w:br"))) == 1


def test_cover_letter_renderer_edits_greeting_and_rejects_stale_original() -> None:
    template, _ = cover_letter_template()
    greeting = json.dumps(
        {
            "replacements": [
                {
                    "paragraphId": "paragraph-0002",
                    "spanId": "paragraph-0002-span-0001",
                    "original": "Dear Hiring Team,",
                    "replacement": "Dear Acme Team,",
                    "reason": "Addresses the company hiring team",
                    "evidenceIds": ["source:paragraph-0002-span-0001"],
                }
            ]
        }
    )
    stale = json.dumps(
        {
            "replacements": [
                {
                    "paragraphId": "paragraph-0003",
                    "spanId": "paragraph-0003-span-0001",
                    "original": "Stale original",
                    "replacement": "Targeted introduction",
                    "reason": "Stale model response",
                    "evidenceIds": [
                        "profile:experience:experience-acme:achievement-a1b2c3d4e5"
                    ],
                }
            ]
        }
    )

    rendered = build_document_from_template(
        template_content=template,
        content=greeting,
        document_type="cover_letter",
    )
    assert Document(BytesIO(rendered)).paragraphs[1].text == "Dear Acme Team,"

    with pytest.raises(ValueError, match="original text does not match"):
        build_document_from_template(
            template_content=template,
            content=stale,
            document_type="cover_letter",
        )
    with pytest.raises(ValueError, match="must contain a replacements array"):
        build_document_from_template(
            template_content=template,
            content='{"unexpected": []}',
            document_type="cover_letter",
        )
