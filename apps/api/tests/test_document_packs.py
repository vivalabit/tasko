from collections.abc import Generator
import base64
from datetime import UTC, datetime, timedelta
from io import BytesIO
import json

import pytest
from docx import Document
from fastapi import HTTPException
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, event, func, select
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

from app.api import documents as documents_api
from app.core.database import Base, get_db
from app.main import app
from app.models.applications import StoredApplicationRecord
from app.models.documents import (
    DocumentFileRecord,
    DocumentPackJobRecord,
    DocumentRecord,
    DocumentTemplateRecord,
    DocumentValidationArtifactRecord,
    DocumentVersionRecord,
)
from app.models.jobs import JobMatchRecord, StoredJobRecord
from app.models.profile import ProfilePayload, ProfileRecord
from app.services.ai_match import (
    DEFAULT_AI_MATCH_MODEL,
    MATCHER_VERSION,
    MATCH_PROMPT_VERSION,
    build_job_snapshot,
    build_job_snapshot_hash,
    build_profile_hash,
)
from app.services.document_validation import DocumentValidationError
from app.services.job_match_store import APPLICATION_GUIDE_STORAGE_KEY
from app.services.storage_cleanup import cleanup_expired_document_storage


@pytest.fixture
def pack_client() -> Generator[tuple[TestClient, sessionmaker[Session]], None, None]:
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    testing_session_local = sessionmaker(bind=engine, autoflush=False, autocommit=False)

    def override_get_db() -> Generator[Session, None, None]:
        with testing_session_local() as db:
            yield db

    vacancy_data = {
        "id": "vacancy-1",
        "title": "Backend Engineer",
        "company": "Acme",
        "overview": "Build Python services.",
    }
    profile_data = {
        "name": "Alex",
        "experience": "Python delivery at Acme in 2023",
        "skills": "Python",
    }
    with testing_session_local() as db:
        db.add_all(
            [
                StoredApplicationRecord(
                    id="application-pack",
                    data={
                        "id": "application-pack",
                        "status": "draft",
                        "job": {"id": "vacancy-1"},
                    },
                ),
                StoredApplicationRecord(
                    id="application-other",
                    data={
                        "id": "application-other",
                        "status": "draft",
                        "job": {"id": "vacancy-1"},
                    },
                ),
                StoredJobRecord(
                    id="vacancy-1",
                    data=vacancy_data,
                ),
                ProfileRecord(
                    id="default",
                    data=profile_data,
                ),
                JobMatchRecord(
                    id="match-pack",
                    job_id="vacancy-1",
                    profile_hash=build_profile_hash(ProfilePayload.model_validate(profile_data)),
                    vacancy_hash=build_job_snapshot_hash(build_job_snapshot(vacancy_data)),
                    model=DEFAULT_AI_MATCH_MODEL,
                    prompt_version=MATCH_PROMPT_VERSION,
                    matcher_version=MATCHER_VERSION,
                    cache_key="cache-pack",
                    score=90,
                    source="openclaw",
                    confidence="high",
                    breakdown={
                        APPLICATION_GUIDE_STORAGE_KEY: {
                            "language": "English",
                            "clarificationQuestions": [],
                            "evidenceMatrix": [
                                {
                                    "requirement": "Python delivery",
                                    "status": "verified",
                                    "evidence": "Python delivery at Acme in 2023",
                                }
                            ],
                        }
                    },
                    reasons=[],
                    gaps=[],
                    heuristic_score=90,
                    created_at=datetime.now(UTC),
                ),
            ]
        )
        db.commit()

    app.dependency_overrides[get_db] = override_get_db
    try:
        yield TestClient(app), testing_session_local
    finally:
        app.dependency_overrides.clear()


def document_data_url(document: Document) -> str:
    output = BytesIO()
    document.save(output)
    return (
        "data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;"
        "base64," + base64.b64encode(output.getvalue()).decode()
    )


def upload_pack_templates(client: TestClient) -> tuple[str, str]:
    resume = Document()
    resume.add_paragraph("Original resume summary.")
    cover = Document()
    cover.add_paragraph("Dear Hiring Team,")
    cover.add_paragraph("Original cover letter body.")
    cover.add_paragraph("Kind regards,")
    resume_response = client.post(
        "/documents/templates",
        json={
            "type": "tailored_resume",
            "name": "Pack CV",
            "fileName": "cv.docx",
            "dataUrl": document_data_url(resume),
        },
    )
    cover_response = client.post(
        "/documents/templates",
        json={
            "type": "cover_letter",
            "name": "Pack cover letter",
            "fileName": "cover.docx",
            "dataUrl": document_data_url(cover),
        },
    )
    assert resume_response.status_code == 201
    assert cover_response.status_code == 201
    return resume_response.json()["id"], cover_response.json()["id"]


def pack_request(
    resume_template_id: str,
    cover_template_id: str,
    *,
    persistence_mode: str = "atomic",
    include_cover: bool = True,
) -> dict[str, object]:
    resume = {
        "title": "Tailored CV",
        "content": json.dumps({"replacements": []}),
        "templateId": resume_template_id,
        "generationFingerprint": "a" * 64,
        "generationModel": "test-model",
        "inputVersions": {"profile": "profile-v1"},
    }
    cover = {
        "title": "Cover letter",
        "content": "Dear Hiring Team,\n\nPython delivery at Acme in 2023.\n\nKind regards,",
        "templateId": cover_template_id,
        "generationFingerprint": "b" * 64,
        "generationModel": "test-model",
        "inputVersions": {"profile": "profile-v1"},
    }
    return {
        "packJobId": "pack-job-1",
        "jobId": "vacancy-1",
        "applicationId": "application-pack",
        "persistenceMode": persistence_mode,
        "resume": resume,
        "coverLetter": cover if include_cover else None,
        "partialReason": None if include_cover else "Cover generation exhausted retries",
    }


def validation_report(document_type: str) -> dict[str, object]:
    return {
        "factual": {"status": "passed", "checkedChanges": 0},
        "visual": {
            "status": "passed",
            "sourcePageCount": 1,
            "renderedPageCount": 1,
            "linksPreserved": True,
            "tableOverflow": False,
        },
        "diff": [],
        "documentType": document_type,
    }


def test_atomic_pack_commits_both_documents_and_retry_is_idempotent(
    pack_client,
    monkeypatch,
) -> None:
    client, testing_session_local = pack_client
    resume_template_id, cover_template_id = upload_pack_templates(client)
    calls: list[str] = []

    def validate(**kwargs):
        calls.append(kwargs["document_type"])
        return validation_report(kwargs["document_type"])

    monkeypatch.setattr("app.api.documents.validate_generated_document", validate)
    request = pack_request(resume_template_id, cover_template_id)

    created = client.post("/documents/packs", json=request)
    retry_request = json.loads(json.dumps(request))
    retry_request["resume"]["generationFingerprint"] = "c" * 64
    retry_request["resume"]["inputVersions"] = {"profile": "client-retry-value"}
    retried = client.post("/documents/packs", json=retry_request)

    with testing_session_local() as db:
        document_count = db.scalar(select(func.count()).select_from(DocumentRecord))
        version_count = db.scalar(select(func.count()).select_from(DocumentVersionRecord))
        job_count = db.scalar(select(func.count()).select_from(DocumentPackJobRecord))

    assert created.status_code == 201
    assert created.json()["status"] == "completed"
    assert len(created.json()["documents"]) == 2
    assert created.json()["documents"][0]["generationFingerprint"] != "a" * 64
    assert created.json()["documents"][1]["generationFingerprint"] != "b" * 64
    assert all(
        document["generationFingerprint"] == document["currentGenerationFingerprint"]
        for document in created.json()["documents"]
    )
    assert all(
        document["inputVersions"]["fingerprintVersion"] == "generation-fingerprint-v3"
        for document in created.json()["documents"]
    )
    assert [stage["status"] for stage in created.json()["stages"]] == [
        "completed",
        "completed",
        "completed",
    ]
    assert retried.status_code == 201
    assert retried.json()["documents"] == created.json()["documents"]
    assert calls == ["tailored_resume", "cover_letter"]
    assert document_count == 2
    assert version_count == 2
    assert job_count == 1


def test_initial_document_list_paginates_relations_and_batches_freshness_context(
    pack_client,
    monkeypatch,
) -> None:
    client, testing_session_local = pack_client
    resume_template_id, cover_template_id = upload_pack_templates(client)
    monkeypatch.setattr(
        "app.api.documents.validate_generated_document",
        lambda **kwargs: validation_report(kwargs["document_type"]),
    )
    created = client.post(
        "/documents/packs",
        json=pack_request(resume_template_id, cover_template_id),
    )
    assert created.status_code == 201

    now = datetime.now(UTC)
    with testing_session_local() as db:
        records = db.scalars(select(DocumentRecord)).all()
        for record in records:
            template_id = db.scalar(
                select(DocumentFileRecord.template_id).where(
                    DocumentFileRecord.document_id == record.id,
                    DocumentFileRecord.version == 1,
                )
            )
            assert template_id is not None
            for version in range(2, 26):
                db.add(
                    DocumentVersionRecord(
                        id=f"{record.id}-version-{version}",
                        document_id=record.id,
                        version=version,
                        content=f"Version {version}",
                        created_at=now,
                    )
                )
                db.add(
                    DocumentFileRecord(
                        id=f"{record.id}-file-{version}",
                        document_id=record.id,
                        version=version,
                        template_id=template_id,
                        content=b"rendered-docx" * 10_000,
                        created_at=now,
                    )
                )
            record.current_version = 25
            record.updated_at = now
        db.commit()

    real_context_loader = documents_api.load_authoritative_application_generation_context
    context_calls = 0

    def counted_context_loader(*args, **kwargs):
        nonlocal context_calls
        context_calls += 1
        return real_context_loader(*args, **kwargs)

    monkeypatch.setattr(
        documents_api,
        "load_authoritative_application_generation_context",
        counted_context_loader,
    )
    statements: list[str] = []
    engine = testing_session_local.kw["bind"]

    def capture_statement(_connection, _cursor, statement, _parameters, _context, _many):
        statements.append(statement)

    event.listen(engine, "before_cursor_execute", capture_statement)
    try:
        response = client.get(
            "/documents",
            params={"applicationId": "application-pack"},
        )
    finally:
        event.remove(engine, "before_cursor_execute", capture_statement)

    assert response.status_code == 200
    assert len(response.json()) == 2
    assert context_calls == 1
    for document in response.json():
        assert document["versionsTotal"] == 25
        assert document["versionsHasMore"] is True
        assert [version["version"] for version in document["versions"]] == list(
            range(6, 26)
        )
        assert all(version["hasRenderedDocx"] is True for version in document["versions"])
        assert document["generationFingerprint"] == document["currentGenerationFingerprint"]

    normalized_statements = [" ".join(statement.lower().split()) for statement in statements]
    assert any(
        "row_number() over" in statement and "document_versions" in statement
        for statement in normalized_statements
    )
    file_queries = [
        statement
        for statement in normalized_statements
        if " from document_files" in statement
    ]
    assert file_queries
    assert all(
        "document_files.content" not in statement.partition(" from document_files")[0]
        for statement in file_queries
    )


def test_pack_status_can_be_recovered_after_response_loss(pack_client, monkeypatch) -> None:
    client, _ = pack_client
    resume_template_id, cover_template_id = upload_pack_templates(client)
    monkeypatch.setattr(
        "app.api.documents.validate_generated_document",
        lambda **kwargs: validation_report(kwargs["document_type"]),
    )
    request = pack_request(resume_template_id, cover_template_id)
    created = client.post("/documents/packs", json=request)

    recovered = client.get(
        f"/documents/packs/{request['packJobId']}",
        params={"applicationId": request["applicationId"]},
    )
    foreign = client.get(
        f"/documents/packs/{request['packJobId']}",
        params={"applicationId": "application-other"},
    )
    missing = client.get(
        "/documents/packs/missing-pack",
        params={"applicationId": request["applicationId"]},
    )

    assert created.status_code == 201
    assert recovered.status_code == 200
    assert recovered.json() == created.json()
    assert foreign.status_code == 409
    assert missing.status_code == 404


def test_expiration_cleanup_removes_expired_jobs_and_validation_artifacts(
    pack_client,
) -> None:
    client, testing_session_local = pack_client
    now = datetime.now(UTC)
    with testing_session_local() as db:
        db.add_all(
            [
                DocumentPackJobRecord(
                    id="expired-pack",
                    request_fingerprint="a" * 64,
                    application_id="application-pack",
                    persistence_mode="atomic",
                    status="completed",
                    document_ids=[],
                    stages=[],
                    message="expired",
                    created_at=now - timedelta(days=8),
                    updated_at=now - timedelta(days=8),
                    expires_at=now - timedelta(days=1),
                ),
                DocumentPackJobRecord(
                    id="active-pack",
                    request_fingerprint="b" * 64,
                    application_id="application-pack",
                    persistence_mode="atomic",
                    status="completed",
                    document_ids=[],
                    stages=[],
                    message="active",
                    created_at=now,
                    updated_at=now,
                    expires_at=now + timedelta(days=7),
                ),
                DocumentValidationArtifactRecord(
                    id="expired-artifact",
                    application_id="application-pack",
                    document_type="tailored_resume",
                    template_id="expired-template",
                    template_hash="c" * 64,
                    result_hash="d" * 64,
                    evidence_hash="e" * 64,
                    rendered_hash="f" * 64,
                    rendered_content=b"expired",
                    validation_report={},
                    consumed_at=None,
                    expires_at=now - timedelta(minutes=1),
                    created_at=now - timedelta(hours=1),
                ),
            ]
        )
        db.commit()

    deleted_artifacts, deleted_jobs = cleanup_expired_document_storage(
        testing_session_local,
        now=now,
    )
    response = client.get("/documents/packs/active-pack?applicationId=application-pack")

    with testing_session_local() as db:
        expired_job = db.get(DocumentPackJobRecord, "expired-pack")
        active_job = db.get(DocumentPackJobRecord, "active-pack")
        expired_artifact = db.get(
            DocumentValidationArtifactRecord,
            "expired-artifact",
        )

    assert response.status_code == 200
    assert response.json()["expiresAt"]
    assert (deleted_artifacts, deleted_jobs) == (1, 1)
    assert expired_job is None
    assert active_job is not None
    assert expired_artifact is None


def test_pack_storage_foreign_keys_cascade_application_and_template_deletes() -> None:
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )

    @event.listens_for(engine, "connect")
    def enable_foreign_keys(connection, _record) -> None:
        connection.execute("PRAGMA foreign_keys=ON")

    Base.metadata.create_all(bind=engine)
    testing_session_local = sessionmaker(bind=engine, autoflush=False, autocommit=False)
    now = datetime.now(UTC)

    with testing_session_local() as db:
        application = StoredApplicationRecord(id="cascade-application", data={})
        template = DocumentTemplateRecord(
            id="cascade-template",
            type="tailored_resume",
            name="Cascade template",
            file_name="cascade.docx",
            content_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
            content_sha256="a" * 64,
            content=b"template",
            extracted_text="",
            created_at=now,
            updated_at=now,
        )
        db.add_all([application, template])
        db.flush()
        db.add_all(
            [
                DocumentPackJobRecord(
                    id="cascade-pack",
                    request_fingerprint="b" * 64,
                    application_id=application.id,
                    persistence_mode="atomic",
                    status="completed",
                    document_ids=[],
                    stages=[],
                    message="",
                    created_at=now,
                    updated_at=now,
                    expires_at=now + timedelta(days=1),
                ),
                DocumentValidationArtifactRecord(
                    id="cascade-artifact",
                    application_id=application.id,
                    document_type="tailored_resume",
                    template_id=template.id,
                    template_hash="c" * 64,
                    result_hash="d" * 64,
                    evidence_hash="e" * 64,
                    rendered_hash="f" * 64,
                    rendered_content=b"rendered",
                    validation_report={},
                    consumed_at=None,
                    expires_at=now + timedelta(minutes=30),
                    created_at=now,
                ),
            ]
        )
        db.commit()

        db.delete(application)
        db.commit()

        assert db.get(DocumentPackJobRecord, "cascade-pack") is None
        assert db.get(DocumentValidationArtifactRecord, "cascade-artifact") is None

        replacement_application = StoredApplicationRecord(
            id="cascade-template-application",
            data={},
        )
        db.add(replacement_application)
        db.flush()
        db.add(
            DocumentValidationArtifactRecord(
                id="template-cascade-artifact",
                application_id=replacement_application.id,
                document_type="tailored_resume",
                template_id=template.id,
                template_hash="c" * 64,
                result_hash="d" * 64,
                evidence_hash="e" * 64,
                rendered_hash="f" * 64,
                rendered_content=b"rendered",
                validation_report={},
                consumed_at=None,
                expires_at=now + timedelta(minutes=30),
                created_at=now,
            )
        )
        db.commit()

        db.delete(template)
        db.commit()

        assert db.get(DocumentValidationArtifactRecord, "template-cascade-artifact") is None


def test_resume_preflight_validates_without_saving(pack_client, monkeypatch) -> None:
    client, testing_session_local = pack_client
    resume_template_id, cover_template_id = upload_pack_templates(client)
    monkeypatch.setattr(
        "app.api.documents.validate_generated_document",
        lambda **kwargs: validation_report(kwargs["document_type"]),
    )
    request = pack_request(resume_template_id, cover_template_id)

    response = client.post(
        "/documents/packs/validate-resume",
        json={
            "applicationId": request["applicationId"],
            "resume": request["resume"],
        },
    )

    with testing_session_local() as db:
        document_count = db.scalar(select(func.count()).select_from(DocumentRecord))
        job_count = db.scalar(select(func.count()).select_from(DocumentPackJobRecord))
        artifact_count = db.scalar(
            select(func.count()).select_from(DocumentValidationArtifactRecord)
        )

    assert response.status_code == 200
    assert response.json()["status"] == "passed"
    assert response.json()["validation"]["documentType"] == "tailored_resume"
    assert response.json()["validationArtifactId"]
    assert response.json()["expiresAt"]
    assert document_count == 0
    assert job_count == 0
    assert artifact_count == 1


def test_pack_reuses_and_consumes_resume_validation_artifact(
    pack_client,
    monkeypatch,
) -> None:
    client, testing_session_local = pack_client
    resume_template_id, cover_template_id = upload_pack_templates(client)
    calls: list[str] = []

    def validate(**kwargs):
        calls.append(kwargs["document_type"])
        return validation_report(kwargs["document_type"])

    monkeypatch.setattr("app.api.documents.validate_generated_document", validate)
    request = pack_request(resume_template_id, cover_template_id)
    preflight = client.post(
        "/documents/packs/validate-resume",
        json={
            "applicationId": request["applicationId"],
            "resume": request["resume"],
        },
    )
    artifact_id = preflight.json()["validationArtifactId"]
    request["resume"]["validationArtifactId"] = artifact_id

    saved = client.post("/documents/packs", json=request)
    retry = client.post("/documents/packs", json=request)

    with testing_session_local() as db:
        artifact = db.get(DocumentValidationArtifactRecord, artifact_id)
        assert artifact is not None
        assert artifact.consumed_at is not None

    assert saved.status_code == 201
    assert retry.status_code == 201
    assert retry.json() == saved.json()
    assert calls == ["tailored_resume", "cover_letter"]

    reused_request = json.loads(json.dumps(request))
    reused_request["packJobId"] = "pack-job-artifact-reuse"
    reused = client.post("/documents/packs", json=reused_request)
    assert reused.status_code == 422
    assert "already been used" in reused.json()["detail"]["message"]


def test_validation_artifact_is_not_consumed_when_atomic_pack_rolls_back(
    pack_client,
    monkeypatch,
) -> None:
    client, testing_session_local = pack_client
    resume_template_id, cover_template_id = upload_pack_templates(client)

    def validate(**kwargs):
        if kwargs["document_type"] == "cover_letter":
            raise DocumentValidationError("Cover validation failed")
        return validation_report(kwargs["document_type"])

    monkeypatch.setattr("app.api.documents.validate_generated_document", validate)
    request = pack_request(resume_template_id, cover_template_id)
    preflight = client.post(
        "/documents/packs/validate-resume",
        json={"applicationId": request["applicationId"], "resume": request["resume"]},
    )
    artifact_id = preflight.json()["validationArtifactId"]
    request["resume"]["validationArtifactId"] = artifact_id

    failed = client.post("/documents/packs", json=request)
    with testing_session_local() as db:
        artifact = db.get(DocumentValidationArtifactRecord, artifact_id)
        assert artifact is not None
        assert artifact.consumed_at is None

    monkeypatch.setattr(
        "app.api.documents.validate_generated_document",
        lambda **kwargs: validation_report(kwargs["document_type"]),
    )
    saved = client.post("/documents/packs", json=request)

    assert failed.status_code == 422
    assert saved.status_code == 201


@pytest.mark.parametrize("changed_input", ["template", "result", "evidence"])
def test_pack_rejects_validation_artifact_hash_mismatch(
    pack_client,
    monkeypatch,
    changed_input: str,
) -> None:
    client, testing_session_local = pack_client
    resume_template_id, cover_template_id = upload_pack_templates(client)
    monkeypatch.setattr(
        "app.api.documents.validate_generated_document",
        lambda **kwargs: validation_report(kwargs["document_type"]),
    )
    request = pack_request(resume_template_id, cover_template_id)
    preflight = client.post(
        "/documents/packs/validate-resume",
        json={"applicationId": request["applicationId"], "resume": request["resume"]},
    )
    artifact_id = preflight.json()["validationArtifactId"]
    request["resume"]["validationArtifactId"] = artifact_id

    if changed_input == "result":
        request["resume"]["content"] = json.dumps(
            {
                "replacements": [
                    {
                        "blockId": "block-0001",
                        "spanId": "block-0001-span-0001",
                        "original": "Original resume summary.",
                        "replacement": "Changed result",
                        "reason": "Hash mismatch fixture",
                        "evidenceIds": ["source:block-0001-span-0001"],
                    }
                ]
            }
        )
    elif changed_input == "evidence":
        with testing_session_local() as db:
            profile = db.get(ProfileRecord, "default")
            assert profile is not None
            profile.data = {**profile.data, "skills": "Python, FastAPI"}
            db.commit()
    else:
        changed_template = Document()
        changed_template.add_paragraph("Changed resume template.")
        changed_output = BytesIO()
        changed_template.save(changed_output)
        with testing_session_local() as db:
            template = db.get(DocumentTemplateRecord, resume_template_id)
            assert template is not None
            template.content = changed_output.getvalue()
            db.commit()

    response = client.post("/documents/packs", json=request)

    with testing_session_local() as db:
        artifact = db.get(DocumentValidationArtifactRecord, artifact_id)
        assert artifact is not None
        assert artifact.consumed_at is None
    if changed_input == "evidence":
        assert response.status_code == 409
        assert response.json()["detail"] == "analysis_stale"
    else:
        assert response.status_code == 422
        assert "hashes do not match" in response.json()["detail"]["message"]


def test_atomic_pack_updates_both_documents_as_one_idempotent_version_batch(
    pack_client,
    monkeypatch,
) -> None:
    client, testing_session_local = pack_client
    resume_template_id, cover_template_id = upload_pack_templates(client)
    monkeypatch.setattr(
        "app.api.documents.validate_generated_document",
        lambda **kwargs: validation_report(kwargs["document_type"]),
    )
    initial_request = pack_request(resume_template_id, cover_template_id)
    initial = client.post("/documents/packs", json=initial_request)
    documents_by_type = {
        document["type"]: document for document in initial.json()["documents"]
    }
    update_request = pack_request(resume_template_id, cover_template_id)
    update_request["packJobId"] = "pack-job-2"
    update_request["resume"]["documentId"] = documents_by_type["tailored_resume"]["id"]
    update_request["coverLetter"]["documentId"] = documents_by_type["cover_letter"]["id"]

    updated = client.post("/documents/packs", json=update_request)
    retried = client.post("/documents/packs", json=update_request)

    with testing_session_local() as db:
        version_count = db.scalar(select(func.count()).select_from(DocumentVersionRecord))

    assert initial.status_code == 201
    assert updated.status_code == 201
    assert {document["currentVersion"] for document in updated.json()["documents"]} == {2}
    assert retried.json()["documents"] == updated.json()["documents"]
    assert version_count == 4


def test_pack_rejects_document_from_another_application_or_wrong_type(
    pack_client,
    monkeypatch,
) -> None:
    client, testing_session_local = pack_client
    resume_template_id, cover_template_id = upload_pack_templates(client)
    monkeypatch.setattr(
        "app.api.documents.validate_generated_document",
        lambda **kwargs: validation_report(kwargs["document_type"]),
    )
    initial = client.post(
        "/documents/packs",
        json=pack_request(resume_template_id, cover_template_id),
    )
    documents_by_type = {
        document["type"]: document for document in initial.json()["documents"]
    }

    foreign_request = pack_request(resume_template_id, cover_template_id)
    foreign_request["packJobId"] = "foreign-pack"
    foreign_request["applicationId"] = "application-other"
    foreign_request["resume"]["documentId"] = documents_by_type["tailored_resume"]["id"]
    foreign = client.post("/documents/packs", json=foreign_request)

    wrong_type_request = pack_request(resume_template_id, cover_template_id)
    wrong_type_request["packJobId"] = "wrong-type-pack"
    wrong_type_request["resume"]["documentId"] = documents_by_type["cover_letter"]["id"]
    wrong_type = client.post("/documents/packs", json=wrong_type_request)

    with testing_session_local() as db:
        version_count = db.scalar(select(func.count()).select_from(DocumentVersionRecord))

    assert initial.status_code == 201
    assert foreign.status_code == 409
    assert foreign.json()["detail"] == (
        "Existing document is not attached to the application"
    )
    assert wrong_type.status_code == 422
    assert wrong_type.json()["detail"] == (
        "Existing document type does not match pack item type"
    )
    assert version_count == 2


def test_atomic_pack_stops_after_failed_cv_validation(pack_client, monkeypatch) -> None:
    client, testing_session_local = pack_client
    resume_template_id, cover_template_id = upload_pack_templates(client)
    calls: list[str] = []

    def validate(**kwargs):
        calls.append(kwargs["document_type"])
        raise DocumentValidationError("CV contains an unsupported claim")

    monkeypatch.setattr("app.api.documents.validate_generated_document", validate)

    response = client.post(
        "/documents/packs",
        json=pack_request(resume_template_id, cover_template_id),
    )

    with testing_session_local() as db:
        document_count = db.scalar(select(func.count()).select_from(DocumentRecord))
        job_count = db.scalar(select(func.count()).select_from(DocumentPackJobRecord))

    assert response.status_code == 422
    assert response.json()["detail"]["stage"] == "resume_validation"
    assert response.json()["detail"]["status"] == "rolled_back"
    assert calls == ["tailored_resume"]
    assert document_count == 0
    assert job_count == 0


def test_atomic_pack_rolls_back_when_cover_letter_validation_fails(
    pack_client,
    monkeypatch,
) -> None:
    client, testing_session_local = pack_client
    resume_template_id, cover_template_id = upload_pack_templates(client)

    def validate(**kwargs):
        if kwargs["document_type"] == "cover_letter":
            raise DocumentValidationError("Cover letter adds an unsupported company")
        return validation_report(kwargs["document_type"])

    monkeypatch.setattr("app.api.documents.validate_generated_document", validate)

    response = client.post(
        "/documents/packs",
        json=pack_request(resume_template_id, cover_template_id),
    )

    with testing_session_local() as db:
        document_count = db.scalar(select(func.count()).select_from(DocumentRecord))
        version_count = db.scalar(select(func.count()).select_from(DocumentVersionRecord))

    assert response.status_code == 422
    assert response.json()["detail"]["stage"] == "cover_letter_validation"
    assert response.json()["detail"]["status"] == "rolled_back"
    assert document_count == 0
    assert version_count == 0


def test_atomic_pack_rolls_back_database_mutations_when_commit_stage_fails(
    pack_client,
    monkeypatch,
) -> None:
    client, testing_session_local = pack_client
    resume_template_id, cover_template_id = upload_pack_templates(client)
    monkeypatch.setattr(
        "app.api.documents.validate_generated_document",
        lambda **kwargs: validation_report(kwargs["document_type"]),
    )
    persist = documents_api.persist_pack_document

    def fail_cover_persist(**kwargs):
        if kwargs["document_type"] == "cover_letter":
            raise HTTPException(status_code=409, detail="Simulated cover persistence failure")
        return persist(**kwargs)

    monkeypatch.setattr("app.api.documents.persist_pack_document", fail_cover_persist)

    response = client.post(
        "/documents/packs",
        json=pack_request(resume_template_id, cover_template_id),
    )

    with testing_session_local() as db:
        document_count = db.scalar(select(func.count()).select_from(DocumentRecord))
        version_count = db.scalar(select(func.count()).select_from(DocumentVersionRecord))
        job_count = db.scalar(select(func.count()).select_from(DocumentPackJobRecord))

    assert response.status_code == 409
    assert document_count == 0
    assert version_count == 0
    assert job_count == 0


@pytest.mark.parametrize("include_cover", [True, False])
def test_explicit_partial_pack_saves_only_validated_cv(
    pack_client,
    monkeypatch,
    include_cover: bool,
) -> None:
    client, testing_session_local = pack_client
    resume_template_id, cover_template_id = upload_pack_templates(client)

    def validate(**kwargs):
        if kwargs["document_type"] == "cover_letter":
            raise DocumentValidationError("Cover letter validation failed")
        return validation_report(kwargs["document_type"])

    monkeypatch.setattr("app.api.documents.validate_generated_document", validate)
    request = pack_request(
        resume_template_id,
        cover_template_id,
        persistence_mode="partial",
        include_cover=include_cover,
    )

    response = client.post("/documents/packs", json=request)

    with testing_session_local() as db:
        records = db.scalars(select(DocumentRecord)).all()
        job = db.get(DocumentPackJobRecord, "pack-job-1")

    assert response.status_code == 201
    assert response.json()["status"] == "partial"
    assert len(response.json()["documents"]) == 1
    assert response.json()["documents"][0]["type"] == "tailored_resume"
    assert [record.type for record in records] == ["tailored_resume"]
    assert job is not None
    assert job.status == "partial"
