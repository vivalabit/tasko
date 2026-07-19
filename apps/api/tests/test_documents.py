from collections.abc import Generator
import base64
from datetime import UTC, datetime
from io import BytesIO
import json
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, func, select
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

from app.core.database import Base, get_db
from app.main import app
from app.models.applications import CandidateConfirmationRecord, StoredApplicationRecord
from app.models.documents import (
    DocumentAttachmentRecord,
    DocumentFileRecord,
    DocumentGenerationProvenanceRecord,
    DocumentRecord,
    DocumentTemplateRecord,
    DocumentVersionGenerationProvenanceRecord,
    DocumentVersionRecord,
    DocumentVersionValidationRecord,
)
from app.models.jobs import JobMatchRecord, StoredJobRecord
from app.models.profile import ProfileRecord
from app.services.ai_match import MATCHER_VERSION
from app.services.job_match_store import APPLICATION_GUIDE_STORAGE_KEY


def test_template_preflight_reports_capabilities_and_rejections() -> None:
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    testing_session_local = sessionmaker(bind=engine, autoflush=False, autocommit=False)

    def override_get_db() -> Generator[Session, None, None]:
        with testing_session_local() as db:
            yield db

    supported_document = Document()
    supported_document.add_paragraph("PROFILE", style="Heading 1")
    supported_body = supported_document.add_paragraph(
        "Product designer building reliable B2B workflows with research evidence."
    )
    supported_body.add_run()._r.append(OxmlElement("w:drawing"))
    supported_output = BytesIO()
    supported_document.save(supported_output)
    supported_data_url = (
        "data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;"
        "base64," + base64.b64encode(supported_output.getvalue()).decode()
    )

    unsupported_document = Document()
    unsupported_paragraph = unsupported_document.add_paragraph("Unsupported source")
    unsupported_paragraph.add_run()._r.append(OxmlElement("w:object"))
    unsupported_output = BytesIO()
    unsupported_document.save(unsupported_output)
    unsupported_data_url = (
        "data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;"
        "base64," + base64.b64encode(unsupported_output.getvalue()).decode()
    )

    app.dependency_overrides[get_db] = override_get_db
    client = TestClient(app)
    try:
        supported = client.post(
            "/documents/templates/preflight",
            json={
                "type": "tailored_resume",
                "name": "Supported CV",
                "fileName": "resume.docx",
                "dataUrl": supported_data_url,
                "promptCharacters": 2200,
            },
        )
        rejected = client.post(
            "/documents/templates/preflight",
            json={
                "type": "tailored_resume",
                "name": "Unsupported CV",
                "fileName": "unsupported.docx",
                "dataUrl": unsupported_data_url,
            },
        )
        with testing_session_local() as db:
            template_count = db.scalar(
                select(func.count()).select_from(DocumentTemplateRecord)
            )
    finally:
        app.dependency_overrides.clear()

    assert supported.status_code == 200
    assert supported.json()["supported"] is True
    assert supported.json()["template"] is None
    assert supported.json()["editableCount"] == 1
    assert supported.json()["immutableCount"] == 2
    assert {item["type"] for item in supported.json()["immutableElements"]} == {
        "heading",
        "drawing",
    }
    assert supported.json()["aiContext"] is None
    assert rejected.status_code == 200
    assert rejected.json()["supported"] is False
    assert rejected.json()["template"] is None
    assert rejected.json()["rejectedElements"] == [
        {"element": "object", "description": "embedded objects"}
    ]
    assert template_count == 0


def test_workspace_source_documents_persist_by_application_and_can_be_deleted() -> None:
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    testing_session_local = sessionmaker(bind=engine, autoflush=False, autocommit=False)

    def override_get_db() -> Generator[Session, None, None]:
        with testing_session_local() as db:
            yield db

    with testing_session_local() as db:
        db.add_all(
            [
                StoredApplicationRecord(
                    id="application-sources",
                    data={"id": "application-sources", "status": "draft"},
                ),
                StoredApplicationRecord(
                    id="application-other",
                    data={"id": "application-other", "status": "draft"},
                ),
            ]
        )
        db.commit()

    document = Document()
    document.add_paragraph("Professional profile")
    output = BytesIO()
    document.save(output)
    original_content = output.getvalue()
    data_url = (
        "data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;"
        "base64," + base64.b64encode(original_content).decode()
    )

    app.dependency_overrides[get_db] = override_get_db
    client = TestClient(app)
    try:
        created = client.post(
            "/documents/workspace-sources",
            json={
                "applicationId": "application-sources",
                "category": "CV / Resume",
                "title": "Target CV",
                "language": "English",
                "fileName": "Résumé-Едуард.docx",
                "dataUrl": data_url,
            },
        )
        listed = client.get(
            "/documents/workspace-sources/library",
            params={"applicationId": "application-sources"},
        )
        other_workspace = client.get(
            "/documents/workspace-sources/library",
            params={"applicationId": "application-other"},
        )
        source_id = created.json()["id"]
        wrong_workspace_delete = client.delete(
            f"/documents/workspace-sources/{source_id}",
            params={"applicationId": "application-other"},
        )
        deleted = client.delete(
            f"/documents/workspace-sources/{source_id}",
            params={"applicationId": "application-sources"},
        )
        listed_after_delete = client.get(
            "/documents/workspace-sources/library",
            params={"applicationId": "application-sources"},
        )
    finally:
        app.dependency_overrides.clear()

    assert created.status_code == 201
    assert created.json()["fileName"] == "Résumé-Едуард.docx"
    assert created.json()["category"] == "CV / Resume"
    assert listed.status_code == 200
    assert [source["id"] for source in listed.json()] == [source_id]
    restored_content = base64.b64decode(listed.json()[0]["dataUrl"].partition(",")[2])
    assert restored_content == original_content
    assert other_workspace.json() == []
    assert wrong_workspace_delete.status_code == 404
    assert deleted.status_code == 204
    assert listed_after_delete.json() == []


def test_document_versions_download_and_application_attachments() -> None:
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    testing_session_local = sessionmaker(bind=engine, autoflush=False, autocommit=False)

    def override_get_db() -> Generator[Session, None, None]:
        with testing_session_local() as db:
            yield db

    with testing_session_local() as db:
        db.add(
            StoredApplicationRecord(
                id="application-one",
                data={"id": "application-one", "status": "applied"},
            )
        )
        db.commit()

    app.dependency_overrides[get_db] = override_get_db
    client = TestClient(app)

    try:
        created = client.post(
            "/documents",
            json={
                "type": "cover_letter",
                "title": "Figma cover letter",
                "content": "Dear Hiring Team,\n\nI am applying for the role.",
                "jobId": "job-figma",
                "applicationId": "application-one",
            },
        )
        document_id = created.json()["id"]
        updated = client.patch(
            f"/documents/{document_id}",
            json={
                "content": json.dumps({"replacements": []}),
            },
        )
        restored = client.post(
            f"/documents/{document_id}/restore",
            json={"version": 1},
        )
        listed = client.get("/documents?jobId=job-figma")
        downloaded = client.get(f"/documents/{document_id}/download?version=2")
        detached = client.delete(
            f"/documents/{document_id}/attachments/application-one"
        )
        deleted = client.delete(f"/documents/{document_id}")

        with testing_session_local() as db:
            version_count = db.scalar(select(func.count()).select_from(DocumentVersionRecord))
            attachment_count = db.scalar(
                select(func.count()).select_from(DocumentAttachmentRecord)
            )
            provenance_count = db.scalar(
                select(func.count()).select_from(DocumentGenerationProvenanceRecord)
            )
            version_provenance_count = db.scalar(
                select(func.count()).select_from(
                    DocumentVersionGenerationProvenanceRecord
                )
            )
    finally:
        app.dependency_overrides.clear()

    assert created.status_code == 201
    assert created.json()["currentVersion"] == 1
    assert created.json()["applicationIds"] == ["application-one"]
    assert created.json()["generationFingerprint"] is None
    assert created.json()["generationModel"] is None
    assert created.json()["inputVersions"] == {}
    assert created.json()["versions"][0]["hasRenderedDocx"] is False
    assert updated.status_code == 200
    assert updated.json()["id"] == document_id
    assert updated.json()["currentVersion"] == 2
    assert len(updated.json()["versions"]) == 2
    assert updated.json()["applicationIds"] == ["application-one"]
    assert updated.json()["generationFingerprint"] is None
    assert updated.json()["generationModel"] is None
    assert restored.status_code == 200
    assert restored.json()["id"] == document_id
    assert restored.json()["currentVersion"] == 3
    assert restored.json()["versions"][-1]["content"].startswith("Dear Hiring Team")
    assert restored.json()["applicationIds"] == ["application-one"]
    assert restored.json()["generationFingerprint"] is None
    assert listed.status_code == 200
    assert listed.json()[0]["jobId"] == "job-figma"
    assert listed.json()[0]["generationFingerprint"] is None
    assert downloaded.status_code == 410
    assert downloaded.json()["detail"] == (
        "Rendered DOCX is no longer available for recovery"
    )
    assert detached.status_code == 204
    assert deleted.status_code == 204
    assert version_count == 0
    assert attachment_count == 0
    assert provenance_count == 0
    assert version_provenance_count == 0


def test_document_download_preserves_unicode_filename_for_rendered_docx() -> None:
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    testing_session_local = sessionmaker(bind=engine, autoflush=False, autocommit=False)

    def override_get_db() -> Generator[Session, None, None]:
        with testing_session_local() as db:
            yield db

    rendered_document = Document()
    rendered_document.add_paragraph("Rendered content")
    rendered_output = BytesIO()
    rendered_document.save(rendered_output)

    with testing_session_local() as db:
        record = DocumentRecord(
            id="unicode-document",
            type="tailored_resume",
            title="Résumé Едуард",
            current_version=1,
        )
        record.versions.append(
            DocumentVersionRecord(
                id="unicode-version",
                document_id=record.id,
                version=1,
                content=json.dumps({"replacements": []}),
            )
        )
        record.files.append(
            DocumentFileRecord(
                id="unicode-file",
                document_id=record.id,
                version=1,
                template_id=None,
                content=rendered_output.getvalue(),
            )
        )
        db.add(record)
        db.commit()

    app.dependency_overrides[get_db] = override_get_db
    client = TestClient(app)
    try:
        downloaded = client.get("/documents/unicode-document/download")
    finally:
        app.dependency_overrides.clear()

    assert downloaded.status_code == 200
    assert downloaded.content == rendered_output.getvalue()
    assert downloaded.headers["content-disposition"] == (
        'attachment; filename="R-sum-v1.docx"; '
        "filename*=UTF-8''R%C3%A9sum%C3%A9-%D0%95%D0%B4%D1%83%D0%B0%D1%80%D0%B4-v1.docx"
    )


def test_document_version_history_is_paginated_from_newest_to_oldest() -> None:
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    testing_session_local = sessionmaker(bind=engine, autoflush=False, autocommit=False)
    now = datetime.now(UTC)
    with testing_session_local() as db:
        record = DocumentRecord(
            id="paginated-document",
            type="cover_letter",
            title="Long history",
            current_version=25,
            created_at=now,
            updated_at=now,
        )
        record.versions.extend(
            DocumentVersionRecord(
                id=f"paginated-version-{version}",
                document_id=record.id,
                version=version,
                content=f"Version {version}",
                created_at=now,
            )
            for version in range(1, 26)
        )
        db.add(record)
        db.commit()

    def override_get_db() -> Generator[Session, None, None]:
        with testing_session_local() as db:
            yield db

    app.dependency_overrides[get_db] = override_get_db
    client = TestClient(app)
    try:
        document = client.get("/documents/paginated-document")
        older_page = client.get(
            "/documents/paginated-document/versions?limit=20&offset=20"
        )
    finally:
        app.dependency_overrides.clear()

    assert document.status_code == 200
    assert document.json()["versionsTotal"] == 25
    assert document.json()["versionsHasMore"] is True
    assert [version["version"] for version in document.json()["versions"]] == list(
        range(6, 26)
    )
    assert older_page.status_code == 200
    assert older_page.json()["total"] == 25
    assert [version["version"] for version in older_page.json()["items"]] == list(
        range(1, 6)
    )


def test_document_attachment_requires_existing_application() -> None:
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    testing_session_local = sessionmaker(bind=engine, autoflush=False, autocommit=False)

    def override_get_db() -> Generator[Session, None, None]:
        with testing_session_local() as db:
            yield db

    app.dependency_overrides[get_db] = override_get_db
    client = TestClient(app)

    try:
        incomplete_provenance = client.post(
            "/documents",
            json={
                "type": "tailored_resume",
                "title": "Incomplete provenance",
                "content": "Verified achievements",
                "generationModel": "test-model",
            },
        )
        response = client.post(
            "/documents",
            json={
                "type": "tailored_resume",
                "title": "Resume for Stripe",
                "content": "# Experience\nVerified achievements",
                "applicationId": "missing-application",
            },
        )
    finally:
        app.dependency_overrides.clear()

    assert incomplete_provenance.status_code == 422
    assert incomplete_provenance.json()["detail"] == (
        "Application ID is required for generated documents"
    )
    assert response.status_code == 404
    assert response.json()["detail"] == "Application not found"


def test_cover_letter_template_preserves_visual_structure() -> None:
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    testing_session_local = sessionmaker(bind=engine, autoflush=False, autocommit=False)

    def override_get_db() -> Generator[Session, None, None]:
        with testing_session_local() as db:
            yield db

    with testing_session_local() as db:
        db.add(
            StoredApplicationRecord(
                id="application-template",
                data={"id": "application-template", "status": "draft"},
            )
        )
        db.commit()

    template = Document()
    template.sections[0].header.paragraphs[0].text = "EDUARD · SOFTWARE ENGINEER"
    template.add_paragraph("Dear Hiring Team,")
    body = template.add_paragraph("Original reusable body paragraph.")
    body.style = template.styles["Normal"]
    template.add_paragraph("Kind regards,")
    template.add_paragraph("Eduard")
    template.sections[0].footer.paragraphs[0].text = "Private application"
    template_output = BytesIO()
    template.save(template_output)
    data_url = "data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64," + base64.b64encode(
        template_output.getvalue()
    ).decode()

    app.dependency_overrides[get_db] = override_get_db
    client = TestClient(app)

    try:
        uploaded = client.post(
            "/documents/templates",
            json={
                "type": "cover_letter",
                "name": "My cover letter",
                "fileName": "cover-letter.docx",
                "dataUrl": data_url,
            },
        )
        template_id = uploaded.json()["id"]
        listed = client.get("/documents/templates/library")
        created = client.post(
            "/documents",
            json={
                "type": "cover_letter",
                "title": "Generated cover letter",
                "content": (
                    "Dear Acme Hiring Team,\n\n"
                    "My first tailored paragraph.\n\n"
                    "My second tailored paragraph.\n\n"
                    "Kind regards,\n\nEduard"
                ),
                "jobId": "job-acme",
                "applicationId": "application-template",
                "templateId": template_id,
            },
        )
        downloaded = client.get(f"/documents/{created.json()['id']}/download")
    finally:
        app.dependency_overrides.clear()

    rendered = Document(BytesIO(downloaded.content))
    body_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    header_text = rendered.sections[0].header.paragraphs[0].text
    footer_text = rendered.sections[0].footer.paragraphs[0].text

    assert uploaded.status_code == 201
    assert listed.status_code == 200
    assert listed.json()[0]["fileName"] == "cover-letter.docx"
    assert created.status_code == 201
    assert created.json()["generationFingerprint"] is None
    assert downloaded.status_code == 200
    assert header_text == "EDUARD · SOFTWARE ENGINEER"
    assert footer_text == "Private application"
    assert "My first tailored paragraph." in body_text
    assert "My second tailored paragraph." in body_text
    assert "Original reusable body paragraph." not in body_text


def test_resume_template_rewrites_blocks_without_rebuilding_design() -> None:
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    testing_session_local = sessionmaker(bind=engine, autoflush=False, autocommit=False)

    def override_get_db() -> Generator[Session, None, None]:
        with testing_session_local() as db:
            yield db

    with testing_session_local() as db:
        db.add(
            StoredApplicationRecord(
                id="application-resume-template",
                data={"id": "application-resume-template", "status": "draft"},
            )
        )
        db.commit()

    template = Document()
    template.sections[0].header.paragraphs[0].text = "EDUARD · CONTACT"
    name = template.add_paragraph("Eduard Ishchenko")
    name.style = template.styles["Heading 1"]
    template.add_paragraph("Original professional summary with verified delivery experience.")
    table = template.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "Original achievement"
    template.sections[0].footer.paragraphs[0].text = "Resume footer"
    template_output = BytesIO()
    template.save(template_output)
    data_url = "data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64," + base64.b64encode(
        template_output.getvalue()
    ).decode()

    app.dependency_overrides[get_db] = override_get_db
    client = TestClient(app)
    try:
        uploaded = client.post(
            "/documents/templates",
            json={
                "type": "tailored_resume",
                "name": "Main CV",
                "fileName": "resume.docx",
                "dataUrl": data_url,
            },
        )
        created = client.post(
            "/documents",
            json={
                "type": "tailored_resume",
                "title": "Tailored resume",
                "content": json.dumps(
                    {
                        "replacements": [
                            {
                                "blockId": "block-0002",
                                "spanId": "block-0002-span-0001",
                                "original": "Original professional summary with verified delivery experience.",
                                "replacement": "Backend engineer focused on FastAPI",
                                "reason": "Matches the target role with verified profile evidence",
                                "evidenceIds": ["source:block-0002-span-0001"],
                            }
                        ]
                    }
                ),
                "applicationId": "application-resume-template",
                "templateId": uploaded.json()["id"],
            },
        )
        downloaded = client.get(f"/documents/{created.json()['id']}/download")
    finally:
        app.dependency_overrides.clear()

    rendered = Document(BytesIO(downloaded.content))
    with zipfile.ZipFile(BytesIO(template_output.getvalue())) as source_package:
        with zipfile.ZipFile(BytesIO(downloaded.content)) as rendered_package:
            for preserved_part in (
                "word/styles.xml",
                "word/header1.xml",
                "word/footer1.xml",
            ):
                assert rendered_package.read(preserved_part) == source_package.read(preserved_part)
            assert rendered_package.read("word/document.xml") != source_package.read(
                "word/document.xml"
            )

    assert uploaded.status_code == 201
    assert created.status_code == 201
    assert downloaded.status_code == 200
    assert rendered.sections[0].header.paragraphs[0].text == "EDUARD · CONTACT"
    assert rendered.sections[0].footer.paragraphs[0].text == "Resume footer"
    assert rendered.paragraphs[0].style.name == "Heading 1"
    assert rendered.paragraphs[1].text == "Backend engineer focused on FastAPI"
    assert rendered.tables[0].style.name == "Table Grid"
    assert rendered.tables[0].cell(0, 0).text == "Python"
    assert rendered.tables[0].cell(0, 1).text == "Original achievement"


def test_generated_template_document_exposes_validation_and_diff(monkeypatch) -> None:
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    testing_session_local = sessionmaker(bind=engine, autoflush=False, autocommit=False)

    def override_get_db() -> Generator[Session, None, None]:
        with testing_session_local() as db:
            yield db

    expected_diff = [{
        "blockId": "paragraph-change-0001",
        "type": "paragraph",
        "original": "Original body.",
        "replacement": "Built a Python service at Acme in 2023.",
        "reason": "Generated cover-letter paragraph update",
    }]

    captured_evidence = {}

    def fake_validation(**kwargs):
        captured_evidence.update(kwargs["evidence"])
        return {
            "factual": {"status": "passed", "checkedChanges": 1},
            "visual": {
                "status": "passed",
                "sourcePageCount": 1,
                "renderedPageCount": 1,
                "linksPreserved": True,
                "tableOverflow": False,
            },
            "diff": expected_diff,
        }

    monkeypatch.setattr("app.api.documents.validate_generated_document", fake_validation)

    with testing_session_local() as db:
        db.add_all(
            [
                StoredApplicationRecord(
                    id="application-validated",
                    data={
                        "id": "application-validated",
                        "status": "draft",
                        "job": {
                            "id": "job-validated",
                            "aiMatch": {
                                "applicationGuide": {
                                    "language": "Client-controlled language"
                                }
                            },
                        },
                    },
                ),
                StoredJobRecord(
                    id="job-validated",
                    data={
                        "id": "job-validated",
                        "title": "Backend Engineer",
                        "company": "Acme",
                    },
                ),
                ProfileRecord(
                    id="default",
                    data={
                        "name": "Alex",
                        "skills": "Python",
                        "experience": "Built a Python service at Acme in 2023.",
                    },
                ),
                JobMatchRecord(
                    id="match-validated",
                    job_id="job-validated",
                    profile_hash="profile-validated",
                    matcher_version=MATCHER_VERSION,
                    cache_key="cache-validated",
                    score=90,
                    source="openclaw",
                    confidence="high",
                    breakdown={
                        APPLICATION_GUIDE_STORAGE_KEY: {
                            "language": "German",
                            "clarificationQuestions": [
                                {
                                    "id": "production-python",
                                    "requirement": "Production Python",
                                    "question": "Have you used Python in production?",
                                    "blocking": True,
                                }
                            ],
                            "evidenceMatrix": [],
                        }
                    },
                    reasons=[],
                    gaps=[],
                    heuristic_score=90,
                    created_at=datetime.now(UTC),
                ),
                CandidateConfirmationRecord(
                    application_id="application-validated",
                    question_id="production-python",
                    requirement="Client-controlled requirement",
                    response="yes",
                    example_text="Built a Python service at Acme in 2023.",
                    blocking=False,
                    updated_at=datetime.now(UTC),
                ),
            ]
        )
        db.commit()

    template = Document()
    template.add_paragraph("Dear Hiring Team,")
    template.add_paragraph("Original body.")
    template.add_paragraph("Kind regards,")
    template_output = BytesIO()
    template.save(template_output)
    data_url = (
        "data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;"
        "base64," + base64.b64encode(template_output.getvalue()).decode()
    )

    app.dependency_overrides[get_db] = override_get_db
    client = TestClient(app)
    try:
        uploaded = client.post(
            "/documents/templates",
            json={
                "type": "cover_letter",
                "name": "Validated source",
                "fileName": "source.docx",
                "dataUrl": data_url,
            },
        )
        created = client.post(
            "/documents",
            json={
                "type": "cover_letter",
                "title": "Validated letter",
                "content": (
                    "Dear Hiring Team,\n\n"
                    "Built a Python service at Acme in 2023.\n\n"
                    "Kind regards,"
                ),
                "applicationId": "application-validated",
                "jobId": "job-validated",
                "templateId": uploaded.json()["id"],
                "generationFingerprint": "f" * 64,
                "generationModel": "test-model",
                "inputVersions": {"profile": "profile-v1"},
                "validationEvidence": {"profile": "Client-controlled evidence"},
            },
        )
        listed = client.get("/documents")
        with testing_session_local() as db:
            profile = db.get(ProfileRecord, "default")
            assert profile is not None
            profile.data = {**profile.data, "skills": "Python, PostgreSQL"}
            db.commit()
        refreshed = client.get("/documents?applicationId=application-validated")
        foreign_update = client.patch(
            f"/documents/{created.json()['id']}",
            json={
                "title": "Must not be changed",
                "applicationId": "application-not-owner",
            },
        )
        manually_updated = client.patch(
            f"/documents/{created.json()['id']}",
            json={
                "content": (
                    "Dear Hiring Team,\n\n"
                    "This content was edited without regeneration.\n\n"
                    "Kind regards,"
                ),
                "applicationId": "application-validated",
            },
        )
        with testing_session_local() as db:
            validation_count = db.scalar(
                select(func.count()).select_from(DocumentVersionValidationRecord)
            )
            current_provenance_count = db.scalar(
                select(func.count()).select_from(DocumentGenerationProvenanceRecord)
            )
            version_provenance_count = db.scalar(
                select(func.count()).select_from(
                    DocumentVersionGenerationProvenanceRecord
                )
            )
    finally:
        app.dependency_overrides.clear()

    assert uploaded.status_code == 201
    assert created.status_code == 201
    assert created.json()["generationFingerprint"] != "f" * 64
    assert created.json()["generationFingerprint"] == created.json()[
        "currentGenerationFingerprint"
    ]
    assert created.json()["inputVersions"]["fingerprintVersion"] == (
        "generation-fingerprint-v3"
    )
    assert created.json()["inputVersions"]["sourceDocument"]["id"] == uploaded.json()["id"]
    assert created.json()["inputVersions"]["profile"] != "profile-v1"
    assert created.json()["versions"][0]["hasRenderedDocx"] is True
    assert created.json()["versions"][0]["factualValidation"]["status"] == "passed"
    assert created.json()["versions"][0]["visualValidation"]["renderedPageCount"] == 1
    assert created.json()["versions"][0]["diff"] == expected_diff
    assert listed.json()[0]["versions"][0]["diff"] == expected_diff
    assert refreshed.status_code == 200
    assert refreshed.json()[0]["generationFingerprint"] == created.json()[
        "generationFingerprint"
    ]
    assert refreshed.json()[0]["currentGenerationFingerprint"] != created.json()[
        "generationFingerprint"
    ]
    assert foreign_update.status_code == 409
    assert foreign_update.json()["detail"] == (
        "Existing document is not attached to the application"
    )
    assert manually_updated.status_code == 200
    assert manually_updated.json()["currentVersion"] == 2
    assert manually_updated.json()["generationFingerprint"] is None
    assert manually_updated.json()["currentGenerationFingerprint"] is None
    assert manually_updated.json()["versions"][0]["factualValidation"]["status"] == "passed"
    assert manually_updated.json()["versions"][1]["factualValidation"] == {}
    assert manually_updated.json()["versions"][1]["visualValidation"] == {}
    assert validation_count == 1
    assert current_provenance_count == 0
    assert version_provenance_count == 1
    assert captured_evidence["language"] == "German"
    assert captured_evidence["profile"]["experience"].startswith("Built a Python service")
    assert captured_evidence["confirmations"] == [
        {
            "requirement": "Production Python",
            "response": "yes",
            "exampleText": "Built a Python service at Acme in 2023.",
        }
    ]
    evidence_catalog = {
        item["id"]: item for item in captured_evidence["evidenceCatalog"]
    }
    experience_evidence = [
        item
        for item in evidence_catalog.values()
        if item["id"].startswith("profile:experience:")
    ]
    assert "profile:experience" not in evidence_catalog
    assert any(
        item["claimType"] == "achievement"
        and item["text"].startswith("Built a Python service")
        for item in experience_evidence
    )
    assert any(
        item["claimType"] == "technology" and item["text"] == "Python"
        for item in experience_evidence
    )
    assert evidence_catalog["confirmation:production-python"]["type"] == "confirmation"
