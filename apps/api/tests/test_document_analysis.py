import base64
import hashlib
import json
from dataclasses import FrozenInstanceError
from io import BytesIO

import pytest
from docx import Document

from app.models.assistant import AssistantSourceDocument
from app.services import document_analysis, document_security
from app.services.assistant import build_source_document_context
from app.services.document_analysis import (
    DocumentAnalysisResult,
    analyze_docx_source,
    clear_document_analysis_cache,
    document_analysis_cache_info,
)
from app.services.document_preflight import analyze_document_template
from app.services.resume_blocks import extract_resume_blocks_from_docx


def resume_content() -> bytes:
    document = Document()
    document.add_paragraph("SUMMARY", style="Heading 1")
    document.add_paragraph("Backend engineer delivering reliable API services.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_source_analysis_reuses_one_zip_pass_for_all_consumers(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    content = resume_content()
    encoded = base64.b64encode(content).decode()
    real_validate = document_analysis.validate_and_read_docx_package
    real_zip_file = document_security.zipfile.ZipFile
    validation_calls = 0
    zip_open_calls = 0

    def counted_validate(value: bytes):
        nonlocal validation_calls
        validation_calls += 1
        return real_validate(value)

    def counted_zip_file(*args, **kwargs):
        nonlocal zip_open_calls
        zip_open_calls += 1
        return real_zip_file(*args, **kwargs)

    clear_document_analysis_cache()
    monkeypatch.setattr(
        document_analysis,
        "validate_and_read_docx_package",
        counted_validate,
    )
    monkeypatch.setattr(document_security.zipfile, "ZipFile", counted_zip_file)
    try:
        preflight = analyze_document_template(content, "tailored_resume")
        blocks = extract_resume_blocks_from_docx(content)
        context = build_source_document_context(
            [
                AssistantSourceDocument(
                    id="source-resume",
                    title="Source resume",
                    category="CV / Resume",
                    fileName="resume.docx",
                    dataUrl=(
                        "data:application/vnd.openxmlformats-officedocument."
                        f"wordprocessingml.document;base64,{encoded}"
                    ),
                )
            ]
        )
        analysis = analyze_docx_source(bytes(content), "tailored_resume")
        cache_info = document_analysis_cache_info()
    finally:
        clear_document_analysis_cache()

    assert preflight["supported"] is True
    assert blocks[1]["type"] == "summary"
    assert context[0]["format"] == "resume-blocks-v2"
    assert context[0]["blocks"] == blocks
    assert validation_calls == 1
    assert zip_open_calls == 1
    assert cache_info.misses == 1
    assert cache_info.hits >= 3
    assert cache_info.size == 1
    assert analysis.content_sha256 == hashlib.sha256(content).hexdigest()


def test_cached_analysis_is_immutable_and_returns_independent_payloads() -> None:
    content = resume_content()
    clear_document_analysis_cache()
    try:
        analysis = analyze_docx_source(content, "tailored_resume")
        cached = analyze_docx_source(content, "tailored_resume")
        first_payload = analysis.structured_elements()
        first_payload[0]["original"] = "mutated outside cache"
        second_payload = cached.structured_elements()
    finally:
        clear_document_analysis_cache()

    assert cached is analysis
    assert second_payload[0]["original"] == "SUMMARY"
    with pytest.raises(FrozenInstanceError):
        analysis.extracted_text = "mutated"  # type: ignore[misc]


def test_ai_context_prioritizes_editable_blocks_over_large_immutable_content() -> None:
    elements = [
        {
            "blockId": "block-0001",
            "type": "immutable",
            "original": "x" * 2_000,
            "editable": False,
            "spans": [],
        },
        {
            "blockId": "block-0002",
            "type": "heading",
            "original": "Experience",
            "editable": False,
            "spans": [],
        },
        {
            "blockId": "block-0003",
            "type": "achievement",
            "original": "Built reliable Python services.",
            "editable": True,
            "spans": [
                {
                    "spanId": "block-0003-span-0001",
                    "type": "text",
                    "original": "Built reliable Python services.",
                    "editable": True,
                    "evidenceId": "source:block-0003-span-0001",
                }
            ],
        },
    ]
    analysis = DocumentAnalysisResult(
        content_sha256="a" * 64,
        document_type="tailored_resume",
        extracted_text="",
        format_name="resume-blocks-v2",
        elements_key="blocks",
        structured_elements_json=json.dumps(elements),
        preflight_json="{}",
    )

    context = analysis.build_ai_context(
        source_id="source-cv",
        title="CV",
        category="CV / Resume",
        file_name="cv.docx",
        max_characters=900,
    )

    assert [block["blockId"] for block in context["blocks"]] == [
        "block-0002",
        "block-0003",
    ]
