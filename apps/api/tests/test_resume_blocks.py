import base64
import json
import zipfile
from io import BytesIO

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from lxml import etree

from app.services.document_export import build_document_from_template
from app.services.document_validation import build_document_diff
from app.services.resume_blocks import extract_resume_blocks_from_docx, paragraph_text


def resume_template() -> bytes:
    document = Document()
    document.add_paragraph("ada@example.com · +41 44 555 12 34")
    document.add_paragraph("SUMMARY", style="Heading 1")
    document.add_paragraph("Original professional profile backed by delivery evidence.")
    document.add_paragraph("SKILLS", style="Heading 1")
    document.add_paragraph("Python, FastAPI, PostgreSQL")
    document.add_paragraph("EXPERIENCE", style="Heading 1")
    document.add_paragraph("Acme · 2024")
    document.add_paragraph("Built and operated a verified production API for customers.")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "German B2"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def document_bytes(document: Document) -> bytes:
    output = BytesIO()
    document.save(output)
    return output.getvalue()


@pytest.fixture
def german_resume_docx() -> bytes:
    document = Document()
    sections = [
        ("Profil", "Backend-Entwicklerin für zuverlässige APIs."),
        ("Kenntnisse", "Python und FastAPI"),
        ("Bildung", "MSc Informatik, Universität Zürich"),
        ("Sprachen", "Deutsch C2 und Englisch C1"),
        ("Projekte", "Entwickelte eine verteilte Plattform für internationale Kundenteams."),
        ("Berufserfahrung", "Betrieb produktiver API-Dienste für regulierte Geschäftskunden."),
    ]
    for heading, content in sections:
        document.add_paragraph(heading)
        document.add_paragraph(content)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


@pytest.fixture
def german_table_resume_docx() -> bytes:
    document = Document()
    table = document.add_table(rows=6, cols=2)
    table.cell(0, 0).text = "Profil"
    table.cell(0, 1).text = "Kenntnisse"
    table.cell(1, 0).text = "Backend-Entwicklerin mit Erfahrung in zuverlässigen Plattformen."
    table.cell(1, 1).text = "Python, FastAPI, PostgreSQL"
    table.cell(2, 0).text = "Berufserfahrung"
    table.cell(2, 1).text = "Projekte"
    table.cell(3, 0).text = "Betrieb produktiver API-Dienste für Geschäftskunden."
    table.cell(3, 1).text = "Entwickelte eine Plattform für internationale Teams."
    table.cell(4, 0).text = "Sprachen"
    table.cell(4, 1).text = "Deutsch C2, Englisch C1, Französisch B2"
    table.cell(5, 0).text = "Bildung"
    table.cell(5, 1).text = "MSc Informatik, Universität Zürich"
    languages_table = document.add_table(rows=2, cols=2)
    languages_table.cell(0, 0).merge(languages_table.cell(0, 1)).text = "Sprachen"
    languages_table.cell(1, 0).text = "Deutsch C2"
    languages_table.cell(1, 1).text = "Englisch C1"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_docx_is_parsed_into_stable_typed_resume_blocks() -> None:
    blocks = extract_resume_blocks_from_docx(resume_template())

    assert [block["blockId"] for block in blocks] == [
        f"block-{index:04d}" for index in range(1, 10)
    ]
    assert {block["type"] for block in blocks} == {
        "immutable",
        "summary",
        "skill",
        "achievement",
        "heading",
        "contact",
        "table cell",
    }
    for block in blocks:
        expected_editable = block["type"] in {"summary", "skill", "achievement"}
        assert block["editable"] is expected_editable
        assert all(
            span["editable"] is expected_editable
            for span in block["spans"]
            if span["type"] == "text"
        )
        if block["type"] in {"heading", "contact", "immutable", "table cell"}:
            assert all(span["editable"] is False for span in block["spans"])


def test_german_resume_sections_are_classified(german_resume_docx: bytes) -> None:
    blocks = extract_resume_blocks_from_docx(german_resume_docx)
    types_by_text = {block["original"]: block["type"] for block in blocks}

    assert {types_by_text[heading] for heading in (
        "Berufserfahrung",
        "Kenntnisse",
        "Bildung",
        "Sprachen",
        "Projekte",
        "Profil",
    )} == {"heading"}
    assert types_by_text["Backend-Entwicklerin für zuverlässige APIs."] == "summary"
    assert types_by_text["Python und FastAPI"] == "skill"
    assert types_by_text["Deutsch C2 und Englisch C1"] == "skill"
    assert types_by_text[
        "Entwickelte eine verteilte Plattform für internationale Kundenteams."
    ] == "achievement"
    assert types_by_text[
        "Betrieb produktiver API-Dienste für regulierte Geschäftskunden."
    ] == "achievement"
    assert types_by_text["MSc Informatik, Universität Zürich"] == "immutable"


def test_german_sections_inside_tables_follow_rows_and_columns(
    german_table_resume_docx: bytes,
) -> None:
    blocks = extract_resume_blocks_from_docx(german_table_resume_docx)
    types_by_text = {block["original"]: block["type"] for block in blocks}

    assert types_by_text[
        "Backend-Entwicklerin mit Erfahrung in zuverlässigen Plattformen."
    ] == "summary"
    assert types_by_text["Python, FastAPI, PostgreSQL"] == "skill"
    assert types_by_text[
        "Betrieb produktiver API-Dienste für Geschäftskunden."
    ] == "achievement"
    assert types_by_text[
        "Entwickelte eine Plattform für internationale Teams."
    ] == "achievement"
    assert types_by_text["Deutsch C2, Englisch C1, Französisch B2"] == "skill"
    assert types_by_text["Deutsch C2"] == "skill"
    assert types_by_text["Englisch C1"] == "skill"
    assert types_by_text["MSc Informatik, Universität Zürich"] == "table cell"


def test_resume_renderer_applies_editable_span_replacements() -> None:
    template = resume_template()
    content = json.dumps(
        {
            "replacements": [
                    {
                        "blockId": "block-0003",
                        "spanId": "block-0003-span-0001",
                        "original": "Original professional profile backed by delivery evidence.",
                    "replacement": "Backend engineer delivering verified FastAPI services.",
                    "reason": "Aligns verified delivery evidence with the vacancy",
                    "evidenceIds": ["source:block-0003-span-0001"],
                }
            ]
        }
    )

    rendered = build_document_from_template(
        template_content=template,
        content=content,
        document_type="tailored_resume",
    )
    document = Document(BytesIO(rendered))
    diff = build_document_diff(
        template,
        content,
        "tailored_resume",
        rendered_content=rendered,
    )

    assert document.paragraphs[2].text == "Backend engineer delivering verified FastAPI services."
    assert document.paragraphs[6].text == "Acme · 2024"
    assert document.tables[0].cell(0, 0).text == "German B2"
    assert diff == [
        {
            "blockId": "block-0003",
            "spanId": "block-0003-span-0001",
            "type": "summary",
            "original": "Original professional profile backed by delivery evidence.",
            "replacement": "Backend engineer delivering verified FastAPI services.",
            "reason": "Aligns verified delivery evidence with the vacancy",
            "evidenceIds": ["source:block-0003-span-0001"],
        }
    ]


def test_resume_preserves_immutable_drawing_symbol_and_supported_field() -> None:
    document = Document()
    document.add_paragraph("SUMMARY", style="Heading 1")
    paragraph = document.add_paragraph()
    paragraph.add_run("Original ")
    paragraph.add_run().add_picture(
        BytesIO(
            base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/"
                "x8AAusB9Y9ZJYQAAAAASUVORK5CYII="
            )
        )
    )
    symbol = OxmlElement("w:sym")
    symbol.set(qn("w:font"), "Wingdings")
    symbol.set(qn("w:char"), "F0B7")
    paragraph.add_run()._r.append(symbol)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    paragraph.add_run()._r.append(field_begin)
    instruction = OxmlElement("w:instrText")
    instruction.text = " PAGE "
    paragraph.add_run()._r.append(instruction)
    field_separator = OxmlElement("w:fldChar")
    field_separator.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._r.append(field_separator)
    paragraph.add_run("1")
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(field_end)
    paragraph.add_run("professional profile backed by evidence.")
    template = document_bytes(document)

    blocks = extract_resume_blocks_from_docx(template)
    body = blocks[1]
    assert [(span["type"], span["editable"]) for span in body["spans"]] == [
        ("text", True),
        ("drawing", False),
        ("symbol", False),
        ("field", False),
        ("text", True),
    ]

    rendered = build_document_from_template(
        template_content=template,
        content=json.dumps(
            {
                "replacements": [
                    {
                        "blockId": "block-0002",
                        "spanId": "block-0002-span-0001",
                        "original": "Original ",
                        "replacement": "Targeted ",
                        "reason": "Tailors only editable text",
                        "evidenceIds": ["source:block-0002-span-0001"],
                    }
                ]
            }
        ),
        document_type="tailored_resume",
    )

    with zipfile.ZipFile(BytesIO(template)) as source_archive, zipfile.ZipFile(
        BytesIO(rendered)
    ) as rendered_archive:
        source_root = etree.fromstring(source_archive.read("word/document.xml"))
        rendered_root = etree.fromstring(rendered_archive.read("word/document.xml"))
        for tag in (qn("w:drawing"), qn("w:sym")):
            assert etree.tostring(rendered_root.find(".//" + tag)) == etree.tostring(
                source_root.find(".//" + tag)
            )
        for tag in (qn("w:fldChar"), qn("w:instrText")):
            assert [etree.tostring(item) for item in rendered_root.findall(".//" + tag)] == [
                etree.tostring(item) for item in source_root.findall(".//" + tag)
            ]
        assert rendered_archive.read("word/media/image1.png") == source_archive.read(
            "word/media/image1.png"
        )


def test_resume_rejects_externally_linked_drawing() -> None:
    document = Document()
    paragraph = document.add_paragraph("Profile")
    drawing = OxmlElement("w:drawing")
    blip = OxmlElement("a:blip")
    blip.set(qn("r:link"), "rIdExternalImage")
    drawing.append(blip)
    paragraph.add_run()._r.append(drawing)

    with pytest.raises(ValueError, match="externally linked drawings"):
        extract_resume_blocks_from_docx(document_bytes(document))


def test_resume_renderer_rejects_immutable_and_mismatched_spans() -> None:
    template = resume_template()
    immutable_replacement = json.dumps(
        {
            "replacements": [
                    {
                        "blockId": "block-0007",
                        "spanId": "block-0007-span-0001",
                        "original": "Acme · 2024",
                    "replacement": "Different employer · 2025",
                    "reason": "Unsafe invented change",
                    "evidenceIds": ["source:block-0007-span-0001"],
                }
            ]
        }
    )
    mismatched_original = json.dumps(
        {
            "replacements": [
                    {
                        "blockId": "block-0003",
                        "spanId": "block-0003-span-0001",
                        "original": "A different original",
                    "replacement": "Backend engineer",
                    "reason": "Stale model response",
                    "evidenceIds": ["source:block-0003-span-0001"],
                }
            ]
        }
    )

    with pytest.raises(ValueError, match="Immutable resume span cannot be changed"):
        build_document_from_template(
            template_content=template,
            content=immutable_replacement,
            document_type="tailored_resume",
        )
    with pytest.raises(ValueError, match="Resume span original does not match template"):
        build_document_from_template(
            template_content=template,
            content=mismatched_original,
            document_type="tailored_resume",
        )

    missing_evidence = json.dumps(
        {
            "replacements": [
                {
                    "blockId": "block-0003",
                    "spanId": "block-0003-span-0001",
                    "original": "Original professional profile backed by delivery evidence.",
                    "replacement": "Backend engineer",
                    "reason": "Missing provenance",
                }
            ]
        }
    )
    with pytest.raises(ValueError, match="must contain 1-20 unique evidenceIds strings"):
        build_document_from_template(
            template_content=template,
            content=missing_evidence,
            document_type="tailored_resume",
        )

    legacy_block_replacement = json.dumps(
        {
            "replacements": [
                {
                    "blockId": "block-0003",
                    "original": "Original professional profile backed by delivery evidence.",
                    "replacement": "Legacy whole-block replacement",
                    "reason": "Old resume-blocks-v1 contract",
                }
            ]
        }
    )
    with pytest.raises(ValueError, match="blockId, spanId, original"):
        build_document_from_template(
            template_content=template,
            content=legacy_block_replacement,
            document_type="tailored_resume",
        )


def add_hyperlink(paragraph, text: str, url: str) -> str:
    relationship_id = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    node = OxmlElement("w:t")
    node.text = text
    run.extend((properties, node))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return relationship_id


def add_plain_content_control(document: Document, text: str) -> None:
    control = OxmlElement("w:sdt")
    properties = OxmlElement("w:sdtPr")
    tag = OxmlElement("w:tag")
    tag.set(qn("w:val"), "tasko-editable")
    plain_text = OxmlElement("w:text")
    properties.extend((tag, plain_text))
    content = OxmlElement("w:sdtContent")
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    italic = OxmlElement("w:i")
    run_properties.append(italic)
    node = OxmlElement("w:t")
    node.text = text
    run.extend((run_properties, node))
    paragraph.append(run)
    content.append(paragraph)
    control.extend((properties, content))
    document._body._element.insert(-1, control)


def complex_resume_template() -> tuple[bytes, str]:
    document = Document()
    document.add_paragraph("SUMMARY", style="Heading 1")
    paragraph = document.add_paragraph()
    first = paragraph.add_run("Original ")
    first.bold = True
    second = paragraph.add_run("profile")
    second.bold = True
    relationship_id = add_hyperlink(paragraph, " site", "https://example.com/profile")
    final = paragraph.add_run()
    final.add_tab()
    final.add_text("Remote")
    final.add_break()
    final.add_text("Delivery")
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Original table achievement"
    add_plain_content_control(document, "Original controlled summary")
    output = BytesIO()
    document.save(output)
    return output.getvalue(), relationship_id


def document_root(content: bytes):
    with zipfile.ZipFile(BytesIO(content)) as archive:
        return etree.fromstring(archive.read("word/document.xml"))


def test_v2_spans_preserve_runs_hyperlinks_controls_and_structural_cells() -> None:
    template, relationship_id = complex_resume_template()
    source_blocks = extract_resume_blocks_from_docx(template)
    source_spans = source_blocks[1]["spans"]
    assert [(span["type"], span["editable"]) for span in source_spans] == [
        ("text", True),
        ("hyperlink", False),
        ("tab", False),
        ("text", True),
        ("lineBreak", False),
        ("text", True),
    ]
    assert source_blocks[2]["type"] == "table cell"
    assert source_blocks[2]["editable"] is False
    assert all(span["editable"] is False for span in source_blocks[2]["spans"])
    template_root = document_root(template)
    template_paragraph = template_root.findall(".//" + qn("w:p"))[1]
    original_run_properties = [
        etree.tostring(properties)
        for properties in template_paragraph.findall(".//" + qn("w:rPr"))
    ]
    content = json.dumps(
        {
            "replacements": [
                    {
                        "blockId": "block-0002",
                        "spanId": "block-0002-span-0001",
                        "original": "Original profile",
                        "replacement": "Targeted summary",
                        "reason": "Tailors the editable text before the protected hyperlink",
                        "evidenceIds": ["source:block-0002-span-0001"],
                    },
                    {
                        "blockId": "block-0002",
                        "spanId": "block-0002-span-0004",
                        "original": "Remote",
                        "replacement": "Hybrid",
                        "reason": "Updates an editable text span between protected controls",
                        "evidenceIds": ["source:block-0002-span-0004"],
                    },
                    {
                        "blockId": "block-0002",
                        "spanId": "block-0002-span-0006",
                        "original": "Delivery",
                        "replacement": "Delivery proof",
                        "reason": "Updates the final editable text span",
                        "evidenceIds": ["source:block-0002-span-0006"],
                    },
                    {
                        "blockId": "block-0004",
                        "spanId": "block-0004-span-0001",
                        "original": "Original controlled summary",
                    "replacement": "Verified controlled summary",
                    "reason": "Updates a supported plain-text content control",
                    "evidenceIds": ["source:block-0004-span-0001"],
                },
            ]
        }
    )

    rendered = build_document_from_template(
        template_content=template,
        content=content,
        document_type="tailored_resume",
    )
    rendered_root = document_root(rendered)
    rendered_paragraph = rendered_root.findall(".//" + qn("w:p"))[1]
    rendered_run_properties = [
        etree.tostring(properties)
        for properties in rendered_paragraph.findall(".//" + qn("w:rPr"))
    ]
    text_nodes = rendered_paragraph.findall(".//" + qn("w:t"))
    hyperlink = rendered_paragraph.find(".//" + qn("w:hyperlink"))
    controlled_text = rendered_root.find(
        ".//" + qn("w:sdt") + "/" + qn("w:sdtContent") + "//" + qn("w:t")
    )
    controlled_italic = rendered_root.find(
        ".//" + qn("w:sdt") + "/" + qn("w:sdtContent") + "//" + qn("w:rPr") + "/" + qn("w:i")
    )

    assert paragraph_text(rendered_paragraph) == "Targeted summary site\tHybrid\nDelivery proof"
    assert original_run_properties == rendered_run_properties
    assert hyperlink is not None
    assert hyperlink.get(qn("r:id")) == relationship_id
    assert text_nodes[0].text == "Targeted summary"
    assert text_nodes[1].text in {None, ""}
    assert len(rendered_paragraph.findall(".//" + qn("w:tab"))) == 1
    assert len(rendered_paragraph.findall(".//" + qn("w:br"))) == 1
    assert controlled_text is not None and controlled_text.text == "Verified controlled summary"
    assert controlled_italic is not None
    rendered_document = Document(BytesIO(rendered))
    assert rendered_document.part.rels[relationship_id].target_ref == "https://example.com/profile"
    assert rendered_document.tables[0].style.name == "Table Grid"
    assert rendered_document.tables[0].cell(0, 0).text == "Original table achievement"


def test_v2_rejects_protected_spans_ambiguous_formatting_and_word_fields() -> None:
    template, _ = complex_resume_template()
    for block_id, span_id, original in (
        ("block-0002", "block-0002-span-0002", " site"),
        ("block-0002", "block-0002-span-0003", "\t"),
        ("block-0002", "block-0002-span-0005", "\n"),
        ("block-0003", "block-0003-span-0001", "Original table achievement"),
    ):
        protected_replacement = json.dumps(
            {
                "replacements": [
                    {
                        "blockId": block_id,
                        "spanId": span_id,
                        "original": original,
                        "replacement": "Changed protected content",
                        "reason": "Invalid protected span change",
                        "evidenceIds": [f"source:{span_id}"],
                    }
                ]
            }
        )
        with pytest.raises(ValueError, match="Immutable resume span cannot be changed"):
            build_document_from_template(
                template_content=template,
                content=protected_replacement,
                document_type="tailored_resume",
            )

    control_in_text = json.dumps(
        {
            "replacements": [
                {
                    "blockId": "block-0002",
                    "spanId": "block-0002-span-0001",
                    "original": "Original profile",
                    "replacement": "Targeted\nsummary",
                    "reason": "Invalid control insertion",
                    "evidenceIds": ["source:block-0002-span-0001"],
                }
            ]
        }
    )
    with pytest.raises(ValueError, match="cannot contain tabs or line breaks"):
        build_document_from_template(
            template_content=template,
            content=control_in_text,
            document_type="tailored_resume",
        )

    mixed = Document()
    mixed.add_paragraph("SUMMARY", style="Heading 1")
    mixed_paragraph = mixed.add_paragraph()
    mixed_paragraph.add_run("Bold fragment").bold = True
    mixed_paragraph.add_run(" and italic fragment").italic = True
    mixed_output = BytesIO()
    mixed.save(mixed_output)
    with pytest.raises(ValueError, match="ambiguous mixed formatting"):
        extract_resume_blocks_from_docx(mixed_output.getvalue())

    unsupported = Document()
    paragraph = unsupported.add_paragraph("Total ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), ' INCLUDETEXT "https://example.com/source.docx" ')
    paragraph._p.append(field)
    output = BytesIO()
    unsupported.save(output)

    with pytest.raises(ValueError, match=r"unsupported Word field INCLUDETEXT"):
        extract_resume_blocks_from_docx(output.getvalue())
