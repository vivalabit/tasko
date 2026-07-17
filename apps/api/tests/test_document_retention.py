import base64
from collections.abc import Generator
from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, func, select
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

from app.core.database import Base, get_db
from app.main import app
from app.models.documents import DocumentFileRecord, DocumentTemplateRecord


def cover_template_data_url() -> str:
    document = Document()
    document.add_paragraph("Dear Hiring Team,")
    document.add_paragraph("Original body")
    document.add_paragraph("Kind regards,")
    output = BytesIO()
    document.save(output)
    encoded = base64.b64encode(output.getvalue()).decode()
    return (
        "data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;"
        f"base64,{encoded}"
    )


def test_templates_are_deduplicated_and_binary_retention_follows_owner_deletion() -> None:
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    testing_session_local = sessionmaker(bind=engine, autoflush=False, autocommit=False)

    def override_get_db() -> Generator[Session, None, None]:
        with testing_session_local() as db:
            yield db

    app.dependency_overrides[get_db] = override_get_db
    client = TestClient(app)
    data_url = cover_template_data_url()

    try:
        first = client.post(
            "/documents/templates",
            json={
                "type": "cover_letter",
                "name": "Original source",
                "fileName": "source.docx",
                "dataUrl": data_url,
            },
        )
        duplicate = client.post(
            "/documents/templates",
            json={
                "type": "cover_letter",
                "name": "Same bytes, another label",
                "fileName": "copy.docx",
                "dataUrl": data_url,
            },
        )
        listed = client.get("/documents/templates/library")
        created = client.post(
            "/documents",
            json={
                "type": "cover_letter",
                "title": "Retained output",
                "content": "Dear Acme Team,\n\nTailored body.\n\nKind regards,",
                "templateId": first.json()["id"],
            },
        )
        document_id = created.json()["id"]
        deleted_template = client.delete(f"/documents/templates/{first.json()['id']}")
        retained_download = client.get(f"/documents/{document_id}/download")

        with testing_session_local() as db:
            template_count = db.scalar(select(func.count()).select_from(DocumentTemplateRecord))
            retained_file = db.scalar(select(DocumentFileRecord))

        deleted_document = client.delete(f"/documents/{document_id}")
        with testing_session_local() as db:
            file_count_after_document_delete = db.scalar(
                select(func.count()).select_from(DocumentFileRecord)
            )
    finally:
        app.dependency_overrides.clear()

    assert first.status_code == 201
    assert duplicate.status_code == 201
    assert duplicate.json()["id"] == first.json()["id"]
    assert len(listed.json()) == 1
    assert "extractedText" not in listed.json()[0]
    assert created.status_code == 201
    assert deleted_template.status_code == 204
    assert template_count == 0
    assert retained_file is not None
    assert retained_file.template_id is None
    assert retained_download.status_code == 200
    assert deleted_document.status_code == 204
    assert file_count_after_document_delete == 0
