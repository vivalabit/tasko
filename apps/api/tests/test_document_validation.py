import base64
import json
import shutil
import time
from collections import Counter
from concurrent.futures import ThreadPoolExecutor
from io import BytesIO
from threading import Lock

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches

from app.services import document_validation as document_validation_service
from app.services.document_validation import (
    DocumentRenderArtifact,
    DocumentValidationError,
    build_authoritative_evidence_catalog,
    build_document_diff,
    clear_source_render_cache,
    compare_rendered_geometry,
    detect_table_overflow,
    render_and_inspect_documents,
    render_and_count_pages,
    source_render_cache_info,
    validate_evidence_id_references,
    validate_factual_changes,
    validate_referenced_factual_changes,
    validate_generated_document,
    validate_visual_output,
)


def document_bytes(document: Document) -> bytes:
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def pdf_geometry(
    *,
    pages: int = 1,
    text: str = "Portfolio Python",
    link_target: str = "external:https://example.com/profile",
) -> dict[str, object]:
    return {
        "pageCount": pages,
        "pageSizes": [{"width": 600.0, "height": 800.0} for _ in range(pages)],
        "textBoxes": [
            {
                "page": 1,
                "x": 50.0,
                "y": 50.0,
                "width": 120.0,
                "height": 14.0,
                "text": text,
            }
        ],
        "imageBoxes": [],
        "linkBoxes": [
            {
                "page": 1,
                "x": 50.0,
                "y": 70.0,
                "width": 80.0,
                "height": 12.0,
                "target": link_target,
            }
        ],
        "text": text,
    }


def atomic_evidence(
    experience_id: str,
    claim_type: str,
    text: str,
) -> tuple[str, dict[str, str]]:
    evidence_id = f"profile:experience:{experience_id}:{claim_type}"
    return evidence_id, {
        "type": "profile",
        "claimType": claim_type,
        "experienceId": experience_id,
        "text": text,
    }


def add_hyperlink(document: Document, target: str) -> None:
    paragraph = document.add_paragraph("Portfolio: ")
    relationship_id = paragraph.part.relate_to(
        target,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "profile"
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def test_factual_validation_checks_entities_technologies_and_claims() -> None:
    allowed = (
        "Built a Python service at Acme in 2023 for 12 clients. "
        "Profile evidence confirms FastAPI delivery and API migrations."
    )
    supported = [{
        "blockId": "block-0002",
        "original": "Built a Python service at Acme in 2023 for 12 clients.",
        "replacement": "Built a FastAPI service at Acme in 2023 for 12 clients.",
    }]
    unsupported = [{
        "blockId": "block-0002",
        "original": "Built a Python service at Acme in 2023 for 12 clients.",
        "replacement": "Led a Kubernetes migration at Globex in 2025 for 900 clients.",
    }]

    assert validate_factual_changes(supported, allowed) == []
    issues = validate_factual_changes(unsupported, allowed)
    assert any("unsupported date" in issue and "2025" in issue for issue in issues)
    assert any("unsupported number" in issue and "900" in issue for issue in issues)
    assert any("unsupported technology" in issue and "kubernetes" in issue for issue in issues)
    assert any("unsupported company" in issue and "Globex" in issue for issue in issues)
    assert any("unsupported claim" in issue for issue in issues)


def test_factual_validation_uses_fact_boundaries_and_checks_job_titles() -> None:
    allowed = "Worked as a Senior Software Engineer at Acme for 120 clients."
    changes = [{
        "blockId": "block-0003",
        "original": "Software Engineer at Acme",
        "replacement": "Principal Software Engineer at Acme for 12 clients.",
    }]

    issues = validate_factual_changes(changes, allowed)

    assert any("unsupported number" in issue and '"12 clients"' in issue for issue in issues)
    assert any(
        "unsupported job title" in issue and "Principal Software Engineer" in issue
        for issue in issues
    )


def test_authoritative_evidence_catalog_contains_source_profile_and_confirmation() -> None:
    source = Document()
    source.add_paragraph("SUMMARY", style="Heading 1")
    source.add_paragraph("Built a Python service at Acme.")

    catalog = build_authoritative_evidence_catalog(
        document_bytes(source),
        {
            "evidenceCatalog": [
                {"id": "profile:skills", "type": "profile", "text": "FastAPI"},
                {
                    "id": "confirmation:production",
                    "type": "confirmation",
                    "text": "Deployed Kubernetes in 2025.",
                },
                {"id": "vacancy:title", "type": "vacancy", "text": "Backend Engineer"},
            ]
        },
    )

    assert catalog["source:block-0002-span-0001"] == {
        "type": "source",
        "text": "Built a Python service at Acme.",
    }
    assert catalog["profile:skills"]["text"] == "FastAPI"
    assert catalog["confirmation:production"]["text"] == "Deployed Kubernetes in 2025."
    assert "vacancy:title" not in catalog


def test_authoritative_evidence_catalog_preserves_atomic_claim_metadata() -> None:
    source = Document()
    source.add_paragraph("Original")

    catalog = build_authoritative_evidence_catalog(
        document_bytes(source),
        {
            "evidenceCatalog": [
                {
                    "id": "profile:experience:acme:employer",
                    "type": "profile",
                    "claimType": "employer",
                    "experienceId": "acme",
                    "text": "Acme",
                }
            ]
        },
    )

    assert catalog["profile:experience:acme:employer"] == {
        "type": "profile",
        "claimType": "employer",
        "experienceId": "acme",
        "text": "Acme",
    }


def test_referenced_validation_rejects_unknown_and_uncited_evidence() -> None:
    catalog = {
        "source:block-0002-span-0001": {
            "type": "source",
            "text": "Built a Python service at Acme in 2023 for 12 clients.",
        },
        "confirmation:principal-role": {
            "type": "confirmation",
            "text": (
                "Worked as a Principal Software Engineer at Globex in 2025 for "
                "900 clients using Kubernetes."
            ),
        },
    }
    replacements: list[dict[str, object]] = [
        {
            "blockId": "block-0002",
            "spanId": "block-0002-span-0001",
            "original": "Built a Python service at Acme in 2023 for 12 clients.",
            "replacement": (
                "Worked as a Principal Software Engineer at Globex in 2025 for "
                "900 clients using Kubernetes."
            ),
            "reason": "Targets the confirmed role",
            "evidenceIds": ["source:block-0002-span-0001", "profile:missing"],
        }
    ]
    diff = [dict(replacements[0])]

    reference_issues = validate_evidence_id_references(replacements, catalog)
    factual_issues = validate_referenced_factual_changes(diff, catalog)

    assert reference_issues == [
        'block-0002-span-0001 references unknown evidence "profile:missing"'
    ]
    for unsupported in ("2025", "900 clients", "kubernetes", "Globex", "Principal"):
        assert any(
            unsupported.casefold() in issue.casefold()
            and "referenced evidence" in issue
            for issue in factual_issues
        )

    diff[0]["evidenceIds"] = ["confirmation:principal-role"]
    assert validate_referenced_factual_changes(diff, catalog) == []


def test_atomic_validation_accepts_one_attributed_experience_and_period() -> None:
    catalog = dict(
        [
            atomic_evidence("acme", "employer", "Acme"),
            atomic_evidence("acme", "title", "Platform Engineer"),
            atomic_evidence("acme", "period", "2022 — 2024"),
            atomic_evidence("acme", "technology", "Python"),
            atomic_evidence(
                "acme",
                "achievement",
                "Built a Python service reducing latency.",
            ),
        ]
    )
    change = {
        "blockId": "block-0002",
        "spanId": "block-0002-span-0001",
        "original": "Software engineering experience.",
        "replacement": (
            "As a Platform Engineer at Acme in 2023, built a Python service reducing latency."
        ),
        "evidenceIds": list(catalog),
    }

    assert validate_referenced_factual_changes([change], catalog) == []


def test_atomic_validation_rejects_cross_employer_claim_laundering() -> None:
    catalog = dict(
        [
            atomic_evidence("acme", "employer", "Acme"),
            atomic_evidence("globex", "period", "2021 — 2024"),
            atomic_evidence("globex", "technology", "Kubernetes"),
            atomic_evidence(
                "globex",
                "achievement",
                "Built a Kubernetes deployment platform.",
            ),
        ]
    )
    change = {
        "blockId": "block-0002",
        "spanId": "block-0002-span-0001",
        "original": "Engineering experience.",
        "replacement": "At Acme in 2023, built a Kubernetes deployment platform.",
        "evidenceIds": list(catalog),
    }

    issues = validate_referenced_factual_changes([change], catalog)

    assert any("different experience records" in issue for issue in issues)


@pytest.mark.parametrize(
    "replacement",
    [
        "At Acme in 2025, built a Python service.",
        "Currently at Acme, built a Python service.",
    ],
)
def test_atomic_validation_rejects_unsupported_chronology(replacement: str) -> None:
    catalog = dict(
        [
            atomic_evidence("acme", "employer", "Acme"),
            atomic_evidence("acme", "period", "2021 — 2024"),
            atomic_evidence("acme", "technology", "Python"),
            atomic_evidence("acme", "achievement", "Built a Python service."),
        ]
    )
    change = {
        "blockId": "block-0002",
        "spanId": "block-0002-span-0001",
        "original": "Engineering experience.",
        "replacement": replacement,
        "evidenceIds": list(catalog),
    }

    issues = validate_referenced_factual_changes([change], catalog)

    assert any("chronology or employment-period" in issue for issue in issues)


@pytest.mark.parametrize(
    ("evidence_text", "replacement", "expected_issue"),
    [
        (
            "Deployed Kubernetes to production.",
            "Did not deploy Kubernetes to production.",
            "changes negation",
        ),
        (
            "Could deploy Kubernetes to production.",
            "Deployed Kubernetes to production.",
            "changes modality",
        ),
    ],
)
def test_atomic_validation_preserves_negation_and_modality(
    evidence_text: str,
    replacement: str,
    expected_issue: str,
) -> None:
    catalog = dict(
        [
            atomic_evidence("acme", "technology", "Kubernetes"),
            atomic_evidence("acme", "achievement", evidence_text),
        ]
    )
    change = {
        "blockId": "block-0002",
        "spanId": "block-0002-span-0001",
        "original": "Platform operations experience.",
        "replacement": replacement,
        "evidenceIds": list(catalog),
    }

    issues = validate_referenced_factual_changes([change], catalog)

    assert any(expected_issue in issue for issue in issues)


def test_atomic_validation_rejects_unknown_assertions_fail_closed() -> None:
    catalog = dict(
        [
            atomic_evidence("acme", "technology", "Python"),
            atomic_evidence("acme", "achievement", "Built a Python API."),
        ]
    )
    change = {
        "blockId": "block-0002",
        "spanId": "block-0002-span-0001",
        "original": "Engineering experience.",
        "replacement": "Certified quantum computing expert.",
        "evidenceIds": list(catalog),
    }

    issues = validate_referenced_factual_changes([change], catalog)

    assert any("fail-closed" in issue for issue in issues)


def test_referenced_validation_preserves_confirmation_modality() -> None:
    catalog = {
        "confirmation:kubernetes": {
            "type": "confirmation",
            "text": "Could deploy Kubernetes to production.",
        }
    }
    change = {
        "blockId": "block-0002",
        "spanId": "block-0002-span-0001",
        "original": "Platform operations experience.",
        "replacement": "Deployed Kubernetes to production.",
        "evidenceIds": list(catalog),
    }

    issues = validate_referenced_factual_changes([change], catalog)

    assert any("changes modality relative to referenced evidence" in issue for issue in issues)


def test_referenced_validation_rejects_unknown_profile_assertion_fail_closed() -> None:
    catalog = {
        "profile:skills": {
            "type": "profile",
            "text": "Python",
        }
    }
    change = {
        "blockId": "block-0002",
        "spanId": "block-0002-span-0001",
        "original": "Engineering experience.",
        "replacement": "Certified Python expert.",
        "evidenceIds": list(catalog),
    }

    issues = validate_referenced_factual_changes([change], catalog)

    assert any("not supported by referenced evidence" in issue for issue in issues)


def test_referenced_validation_accepts_translated_current_period_marker() -> None:
    catalog = {
        "source:period": {
            "type": "source",
            "text": "06/2026 – Present",
        }
    }
    change = {
        "blockId": "block-0013",
        "spanId": "block-0013-span-0005",
        "original": "06/2026 – Present",
        "replacement": "06/2026 – heute",
        "evidenceIds": ["source:period"],
    }

    assert validate_referenced_factual_changes([change], catalog) == []


def test_referenced_validation_still_rejects_unsupported_job_title_specialization() -> None:
    catalog = {
        "source:title": {
            "type": "source",
            "text": "Web Engineer",
        }
    }
    change = {
        "blockId": "block-0015",
        "spanId": "block-0015-span-0003",
        "original": "Web Engineer",
        "replacement": "Web Engineer (Python, Backend & Automatisierung)",
        "evidenceIds": ["source:title"],
    }

    issues = validate_referenced_factual_changes([change], catalog)

    assert any("not supported by referenced evidence" in issue for issue in issues)


def test_period_does_not_authorize_a_same_number_of_clients() -> None:
    catalog = dict(
        [
            atomic_evidence("acme", "period", "2021 — 2024"),
            atomic_evidence(
                "acme",
                "achievement",
                "Managed a customer migration.",
            ),
        ]
    )
    change = {
        "blockId": "block-0002",
        "spanId": "block-0002-span-0001",
        "original": "Customer migration experience.",
        "replacement": "Managed a migration for 2023 clients.",
        "evidenceIds": list(catalog),
    }

    issues = validate_referenced_factual_changes([change], catalog)

    assert any('number "2023 clients"' in issue for issue in issues)


def test_generated_resume_validation_rejects_unknown_evidence_before_success(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = Document()
    source.add_paragraph("SUMMARY", style="Heading 1")
    source.add_paragraph("Built a Python service at Acme in 2023.")
    source_content = document_bytes(source)
    replacement = {
        "blockId": "block-0002",
        "spanId": "block-0002-span-0001",
        "original": "Built a Python service at Acme in 2023.",
        "replacement": "Built a FastAPI service at Acme in 2023.",
        "reason": "Uses the verified profile skill",
        "evidenceIds": ["source:block-0002-span-0001", "profile:missing"],
    }
    monkeypatch.setattr(
        "app.services.document_validation.validate_visual_output",
        lambda _source, _rendered, **_kwargs: ({"status": "passed"}, []),
    )

    with pytest.raises(
        DocumentValidationError,
        match='references unknown evidence "profile:missing"',
    ):
        validate_generated_document(
            template_content=source_content,
            rendered_content=source_content,
            generated_content=json.dumps({"replacements": [replacement]}),
            document_type="tailored_resume",
            evidence={"evidenceCatalog": []},
        )

    replacement["evidenceIds"] = [
        "source:block-0002-span-0001",
        "profile:skills",
    ]
    report = validate_generated_document(
        template_content=source_content,
        rendered_content=source_content,
        generated_content=json.dumps({"replacements": [replacement]}),
        document_type="tailored_resume",
        evidence={
            "evidenceCatalog": [
                {"id": "profile:skills", "type": "profile", "text": "FastAPI"}
            ]
        },
    )

    assert report["factual"] == {
        "status": "passed",
        "checkedChanges": 1,
        "checkedEvidenceCharacters": len("Built a Python service at Acme in 2023.FastAPI"),
    }


def test_cover_letter_diff_compares_the_rendered_docx() -> None:
    source = Document()
    source.add_paragraph("Dear Hiring Team,")
    source.add_paragraph("Original body.")
    source.add_paragraph("Kind regards,")
    source.add_table(rows=1, cols=1).cell(0, 0).text = "Preserved table text"
    rendered = Document(BytesIO(document_bytes(source)))
    rendered.paragraphs[1].text = "Validated replacement body."

    diff = build_document_diff(
        document_bytes(source),
        "This raw assistant text is not the comparison authority.",
        "cover_letter",
        rendered_content=document_bytes(rendered),
    )

    assert len(diff) == 1
    assert diff[0]["original"] == "Original body."
    assert diff[0]["replacement"] == "Validated replacement body."


def test_visual_validation_preserves_links_and_detects_table_overflow(monkeypatch) -> None:
    source = Document()
    add_hyperlink(source, "https://example.com/profile")
    source_table = source.add_table(rows=1, cols=1)
    source_table.cell(0, 0).text = "Python"
    rendered = Document(BytesIO(document_bytes(source)))
    rendered.tables[0].cell(0, 0).text = "Python"

    monkeypatch.setattr(
        "app.services.document_validation.render_and_inspect_documents",
        lambda _source, _rendered: (
            pdf_geometry(text="Portfolio profile Python"),
            pdf_geometry(text="Portfolio profile Python"),
        ),
    )
    report, issues = validate_visual_output(document_bytes(source), document_bytes(rendered))

    assert issues == []
    assert report["status"] == "passed"
    assert report["sourcePageCount"] == report["renderedPageCount"] == 1
    assert report["missingTextCount"] == 0
    assert report["textGeometryChangedCount"] == 0
    assert report["missingSourceImageCount"] == 0
    assert report["linkLocationChangedCount"] == 0
    assert report["linksPreserved"] is True
    assert report["issues"] == []

    rendered.tables[0].cell(0, 0).text = "X" * 220
    overflow = detect_table_overflow(document_bytes(source), document_bytes(rendered))
    assert "possible table overflow" in overflow[0]

    removed_link = Document(BytesIO(document_bytes(source)))
    hyperlink = removed_link.paragraphs[0]._p.find(qn("w:hyperlink"))
    removed_link.paragraphs[0]._p.remove(hyperlink)
    _, removed_link_issues = validate_visual_output(
        document_bytes(source),
        document_bytes(removed_link),
    )
    assert any("hyperlink count or targets changed" in issue for issue in removed_link_issues)

    changed_links = Document()
    add_hyperlink(changed_links, "https://example.com/different")
    changed_links.add_table(rows=1, cols=1).cell(0, 0).text = "Python"
    _, link_issues = validate_visual_output(
        document_bytes(source),
        document_bytes(changed_links),
    )
    assert any("hyperlink count or targets changed" in issue for issue in link_issues)


def test_visual_validation_rejects_link_count_changes_and_page_changes(monkeypatch) -> None:
    source = Document()
    add_hyperlink(source, "https://example.com/original")
    rendered = Document(BytesIO(document_bytes(source)))
    add_hyperlink(rendered, "https://example.com/new")
    monkeypatch.setattr(
        "app.services.document_validation.render_and_inspect_documents",
        lambda _source, _rendered: (pdf_geometry(), pdf_geometry(pages=2)),
    )

    report, issues = validate_visual_output(document_bytes(source), document_bytes(rendered))

    assert report["linksPreserved"] is False
    assert report["sourceLinkCount"] == 1
    assert report["renderedLinkCount"] == 2
    assert report["addedLinkCount"] == 1
    assert "page count changed from 1 to 2" in issues
    assert any("hyperlink count or targets changed" in issue for issue in issues)

    monkeypatch.setattr(
        "app.services.document_validation.render_and_inspect_documents",
        lambda _source, _rendered: (pdf_geometry(pages=2), pdf_geometry()),
    )
    _, fewer_page_issues = validate_visual_output(
        document_bytes(source),
        document_bytes(source),
    )
    assert "page count changed from 2 to 1" in fewer_page_issues


def test_geometry_comparison_reports_missing_content_boxes_links_and_overflow() -> None:
    source = pdf_geometry(text="Stable source")
    source["imageBoxes"] = [
        {
            "page": 1,
            "x": 40.0,
            "y": 100.0,
            "width": 80.0,
            "height": 60.0,
            "digest": "image-a",
        }
    ]
    rendered = pdf_geometry(text="Stable source")
    rendered["textBoxes"][0]["x"] = 300.0
    rendered["textBoxes"].append(
        {
            "page": 1,
            "x": 590.0,
            "y": 200.0,
            "width": 40.0,
            "height": 14.0,
            "text": "Outside",
        }
    )
    rendered["text"] = "Stable source Outside"
    rendered["imageBoxes"] = [
        {
            "page": 1,
            "x": 300.0,
            "y": 100.0,
            "width": 80.0,
            "height": 60.0,
            "digest": "image-a",
        }
    ]
    rendered["linkBoxes"][0]["x"] = 300.0

    report, issues = compare_rendered_geometry(
        source,
        rendered,
        expected_rendered_text="Stable source text Outside",
        source_image_digests=Counter({"source-image": 1}),
        rendered_image_digests=Counter(),
    )

    assert report["missingTextCount"] == 1
    assert report["missingTextSamples"] == ["text"]
    assert report["textGeometryChangedCount"] == 1
    assert report["textOutsidePageCount"] == 1
    assert report["missingSourceImageCount"] == 1
    assert report["imageGeometryChangedCount"] == 1
    assert report["linkLocationChangedCount"] == 1
    assert any("rendered text tokens are missing" in issue for issue in issues)
    assert any("text boxes extend outside" in issue for issue in issues)
    assert any("source images are missing" in issue for issue in issues)
    assert any("image boxes moved" in issue for issue in issues)
    assert any("hyperlink bounding boxes moved" in issue for issue in issues)


def test_geometry_comparison_allows_changed_text_but_rejects_unexpected_disappearance() -> None:
    source = pdf_geometry(text="Protected header Original body")
    rendered = pdf_geometry(text="Replacement body")

    report, issues = compare_rendered_geometry(
        source,
        rendered,
        expected_rendered_text="Replacement body",
        allowed_removed_text="Original body",
        source_image_digests=Counter(),
        rendered_image_digests=Counter(),
    )

    assert report["missingTextCount"] == 0
    assert report["disappearedSourceTextCount"] == 2
    assert report["disappearedSourceTextSamples"] == ["protected", "header"]
    assert any("source text tokens disappeared unexpectedly" in issue for issue in issues)


def test_geometry_comparison_allows_box_reflow_for_valid_text_replacements() -> None:
    source = pdf_geometry(text="Original body Stable footer")
    source["imageBoxes"] = [
        {
            "page": 1,
            "x": 40.0,
            "y": 100.0,
            "width": 80.0,
            "height": 60.0,
            "digest": "signature-image",
        }
    ]
    rendered = pdf_geometry(text="Short body Stable footer")
    source["textBoxes"][0]["text"] = "Stable footer"
    rendered["textBoxes"][0]["text"] = "Stable footer"
    rendered["textBoxes"][0]["x"] = 300.0
    rendered["imageBoxes"] = [
        {
            "page": 1,
            "x": 300.0,
            "y": 100.0,
            "width": 80.0,
            "height": 60.0,
            "digest": "signature-image",
        }
    ]
    rendered["linkBoxes"][0]["x"] = 300.0

    report, issues = compare_rendered_geometry(
        source,
        rendered,
        expected_rendered_text="Short body Stable footer",
        allowed_removed_text="Original body",
        source_image_digests=Counter({"signature-image": 1}),
        rendered_image_digests=Counter({"signature-image": 1}),
    )

    assert report["expectedTextReflow"] is True
    assert report["textGeometryChangedCount"] == 1
    assert report["imageGeometryChangedCount"] == 1
    assert report["linkLocationChangedCount"] == 1
    assert not any("moved or resized" in issue for issue in issues)
    assert not any("bounding boxes moved" in issue for issue in issues)


def test_source_pdf_and_geometry_are_cached_by_content_hash(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    clear_source_render_cache()
    calls: list[tuple[bytes, str, bool]] = []

    def fake_render(
        content: bytes,
        *,
        executable: str,
        label: str,
        rasterize: bool,
    ) -> DocumentRenderArtifact:
        assert executable == "/fake/soffice"
        calls.append((content, label, rasterize))
        return DocumentRenderArtifact(
            pdf_content=b"%PDF-1.7\n" + content,
            geometry_json=json.dumps(pdf_geometry(text=content.decode())),
        )

    monkeypatch.setattr(document_validation_service.shutil, "which", lambda _name: "/fake/soffice")
    monkeypatch.setattr(document_validation_service, "render_document_artifact", fake_render)
    try:
        first_source, _ = render_and_inspect_documents(b"shared-source", b"output-one")
        first_source["text"] = "mutated outside cache"
        second_source, _ = render_and_inspect_documents(b"shared-source", b"output-two")
        cache_info = source_render_cache_info()
    finally:
        clear_source_render_cache()

    assert [call for call in calls if call[1] == "source"] == [
        (b"shared-source", "source", False)
    ]
    assert [call[0] for call in calls if call[1] == "rendered"] == [
        b"output-one",
        b"output-two",
    ]
    assert second_source["text"] == "shared-source"
    assert cache_info.hits == 1
    assert cache_info.misses == 1
    assert cache_info.size == 1
    assert cache_info.in_flight == 0


def test_independent_conversions_run_with_bounded_parallelism(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    clear_source_render_cache()
    state_lock = Lock()
    active = 0
    max_active = 0
    source_calls = 0

    def fake_render(
        content: bytes,
        *,
        executable: str,
        label: str,
        rasterize: bool,
    ) -> DocumentRenderArtifact:
        nonlocal active, max_active, source_calls
        assert executable == "/fake/soffice"
        assert rasterize is (label == "rendered")
        with state_lock:
            active += 1
            max_active = max(max_active, active)
            source_calls += label == "source"
        time.sleep(0.03)
        with state_lock:
            active -= 1
        return DocumentRenderArtifact(
            pdf_content=b"%PDF-1.7\n" + content,
            geometry_json=json.dumps(pdf_geometry(text=content.decode())),
        )

    monkeypatch.setattr(document_validation_service.shutil, "which", lambda _name: "/fake/soffice")
    monkeypatch.setattr(document_validation_service, "render_document_artifact", fake_render)
    try:
        with ThreadPoolExecutor(max_workers=4) as callers:
            results = list(
                callers.map(
                    lambda index: render_and_inspect_documents(
                        b"shared-source",
                        f"output-{index}".encode(),
                    ),
                    range(4),
                )
            )
    finally:
        clear_source_render_cache()

    assert len(results) == 4
    assert source_calls == 1
    assert max_active == document_validation_service.VALIDATION_CONVERSION_MAX_WORKERS


@pytest.mark.skipif(
    not (shutil.which("soffice") or shutil.which("libreoffice")),
    reason="LibreOffice is not installed",
)
def test_docx_is_really_rendered_to_count_pages() -> None:
    source = Document()
    source.add_paragraph("One-page validation fixture")
    source_content = document_bytes(source)
    rendered = Document(BytesIO(source_content))
    rendered.add_page_break()
    rendered.add_paragraph("Second rendered page")

    assert render_and_count_pages(source_content, document_bytes(rendered)) == (1, 2)


@pytest.mark.skipif(
    not (shutil.which("soffice") or shutil.which("libreoffice")),
    reason="LibreOffice is not installed",
)
def test_identical_docx_geometry_is_reported_without_regressions() -> None:
    document = Document()
    document.add_heading("Geometry fixture", level=1)
    document.add_paragraph("Stable text inside the page.")
    add_hyperlink(document, "https://example.com/profile")
    document.add_picture(
        BytesIO(
            base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/"
                "x8AAusB9Y9ZJYQAAAAASUVORK5CYII="
            )
        ),
        width=Inches(0.5),
    )
    document.add_table(rows=1, cols=1).cell(0, 0).text = "Stable table cell"
    content = document_bytes(document)

    report, issues = validate_visual_output(content, content)

    assert issues == []
    assert report["pageCountChanged"] is False
    assert report["missingTextCount"] == 0
    assert report["textGeometryChangedCount"] == 0
    assert report["textOutsidePageCount"] == 0
    assert report["sourceImageCount"] == report["renderedImageCount"] == 1
    assert report["missingSourceImageCount"] == 0
    assert report["missingPdfImageCount"] == 0
    assert report["imageGeometryChangedCount"] == 0
    assert report["linkLocationChangedCount"] == 0

    rendered = Document(BytesIO(content))
    rendered.paragraphs[1].text = "Updated text inside the page."
    changed_report, changed_issues = validate_visual_output(
        content,
        document_bytes(rendered),
    )
    assert changed_issues == []
    assert changed_report["missingTextCount"] == 0
